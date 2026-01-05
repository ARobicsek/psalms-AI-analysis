"""
Script to convert the Readers Guide markdown to a well-formatted DOCX.
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def create_readers_guide_docx():
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading('How the Psalms Reader\'s Guide Works', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Updated date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Updated: January 2026')
    run.bold = True
    run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # Overview
    doc.add_heading('Overview', level=1)
    doc.add_paragraph(
        'The Psalms Reader\'s Guide is generated through an AI-powered pipeline that combines '
        'the reasoning capabilities of multiple advanced language models with ten specialized '
        'digital "librarians" that retrieve and curate source material. This system produces '
        'publication-quality biblical commentary that rivals traditional scholarly work.'
    )
    doc.add_paragraph(
        'The key innovation is a "telescopic analysis" approach—breaking complex tasks into '
        'specialized passes, each building on previous work while maintaining focus on specific '
        'aspects of analysis.'
    )
    
    # Stage 1
    doc.add_heading('The Six-Stage Pipeline', level=1)
    doc.add_heading('Stage 1: Macro Analysis — Setting the Stage', level=2)
    doc.add_paragraph(
        'The first AI agent (Claude Sonnet 4.5) examines the entire psalm with access to '
        'multiple source texts:'
    )
    bullets = [
        ('Hebrew (Masoretic Text)', 'The authoritative Jewish text with vowels and cantillation'),
        ('English translation', 'Jewish Publication Society (1985) rendering'),
        ('Septuagint (LXX)', 'Ancient Greek translation for textual criticism and early interpretation insights'),
    ]
    for term, desc in bullets:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(term + ': ')
        run.bold = True
        p.add_run(desc)
    
    doc.add_paragraph('The agent establishes:')
    items = [
        ('Genre identification', 'Is this a lament, praise, thanksgiving, wisdom, or royal psalm?'),
        ('Structural framework', 'How is the psalm organized? Where are the turns and transitions?'),
        ('Thematic thesis', 'What is the psalm\'s central argument or spiritual insight?'),
        ('Research questions', 'What puzzles or patterns require deeper investigation?'),
    ]
    for term, desc in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(term + ': ')
        run.bold = True
        p.add_run(desc)
    
    doc.add_paragraph(
        'This high-level view prevents the pipeline from getting lost in details before '
        'understanding the whole.'
    )
    
    # Stage 2
    doc.add_heading('Stage 2: Micro Analysis — Verse-by-Verse Discovery', level=2)
    doc.add_paragraph(
        'A second AI agent (Claude Sonnet 4.5) performs discovery-driven analysis of each verse, '
        'looking for:'
    )
    items = [
        ('Hebrew wordplay', 'Puns, alliteration, and sound patterns'),
        ('Unusual vocabulary', 'Rare words that may carry special significance'),
        ('Intertextual echoes', 'Connections to other biblical passages'),
        ('Syntactic features', 'Word order inversions, emphatic constructions'),
        ('LXX variants', 'Where the Greek translation differs, suggesting early textual traditions'),
    ]
    for term, desc in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(term + ': ')
        run.bold = True
        p.add_run(desc)
    
    doc.add_paragraph(
        'The agent works with phonetic transcriptions of the Hebrew text, enabling accurate '
        'analysis of sound patterns, alliteration, and wordplay without requiring knowledge of '
        'Hebrew pronunciation.'
    )
    doc.add_paragraph(
        'Crucially, this agent generates specific research requests for the librarians—asking '
        'for lexicon entries, concordance searches, figurative language parallels, and '
        'traditional commentary on particular terms and phrases.'
    )
    
    # Stage 3
    doc.add_heading('Stage 3: Research Assembly — The Ten Librarians', level=2)
    doc.add_paragraph('Ten specialized Python-based "librarians" retrieve and curate source material:')
    
    librarians = [
        ('BDB Librarian', 'Hebrew lexicon lookups from the Brown-Driver-Briggs and Klein dictionaries via Sefaria API'),
        ('Concordance Librarian', 'Searches a Hebrew concordance database with morphological variations, finding where specific words and phrases appear elsewhere in the Hebrew Bible'),
        ('Figurative Language Librarian', 'Queries a pre-analyzed database of metaphors, similes, and figurative expressions across Biblical Hebrew literature'),
        ('Figurative Curator (NEW)', 'An LLM-enhanced agent (Gemini 3 Pro) that transforms raw figurative concordance results into curated scholarly insights. Using a 3-iteration refinement process, it curates 5-15 examples per vehicle with full Hebrew text and synthesizes 4-5 prose insights (100-150 words each) connecting the examples to the psalm\'s themes.'),
        ('Commentary Librarian', 'Fetches traditional Jewish commentaries—Rashi, Ibn Ezra, Radak, Metzudat David, Malbim, Meiri, and Torah Temimah'),
        ('Liturgical Librarian', 'Identifies where psalm passages appear in Jewish liturgy across three traditions (Ashkenaz, Sefard, Edot HaMizrach), generating intelligent summaries using Gemini 2.5 Pro'),
        ('Related Psalms Librarian', 'Uses statistical analysis (V6 scoring) to identify the five most related psalms, showing shared roots, contiguous phrases, and skipgram patterns'),
        ('Sacks Librarian', 'Retrieves Rabbi Jonathan Sacks\' references and insights on the psalm'),
        ('Deep Web Research Librarian', 'Loads manually prepared research on cultural and artistic afterlife—how the psalm has been used in paintings, music, literature, and political discourse'),
        ('Research Bundle Assembler', 'Coordinates all librarians and formats results into a comprehensive research bundle (up to 700,000 characters)'),
    ]
    for i, (name, desc) in enumerate(librarians, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'{i}. {name}: ')
        run.bold = True
        p.add_run(desc)
    
    # Stage 3b
    doc.add_heading('Stage 3b: Figurative Language Curation', level=2)
    doc.add_paragraph(
        'The Figurative Curator agent (Gemini 3 Pro with high reasoning) transforms raw '
        'figurative concordance results into scholarly interpretive insights:'
    )
    items = [
        'Analyzes figurative vehicle requests from the Micro Analyst',
        'Executes and refines searches against the figurative language database',
        'Curates 5-15 examples per vehicle with full Hebrew text',
        'Synthesizes 4-5 prose insights connecting examples to the psalm\'s themes',
        'Adapts analysis structure to the psalm\'s pattern (journey, descent_ascent, lament_structure, etc.)',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')
    
    # Stage 4
    doc.add_heading('Stage 4: Synthesis Writing — Creating the Commentary', level=2)
    doc.add_paragraph(
        'A third AI agent (Claude Sonnet 4.5, or Gemini 2.5 Pro for large psalms) integrates '
        'all analysis into coherent commentary:'
    )
    items = [
        ('Introduction essay', '800-1200 words situating the psalm in its genre, historical context, and theological significance'),
        ('Verse-by-verse commentary', '150-400+ words per verse with generous quotation of sources in Hebrew and English'),
        ('Enhanced quotation emphasis', 'Biblical parallels are quoted, not just cited'),
        ('Poetic punctuation', 'LLM-generated verses include semicolons, periods, and commas showing structural divisions'),
    ]
    for term, desc in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(term + ': ')
        run.bold = True
        p.add_run(desc)
    
    doc.add_paragraph(
        'For large psalms (51+ verses), the system automatically switches to Gemini 2.5 Pro, '
        'leveraging its 1 million token context window to prevent content loss from research '
        'bundle trimming.'
    )
    
    # Stage 5
    doc.add_heading('Stage 5: Editorial Review — The Master Editor', level=2)
    doc.add_paragraph(
        'A fourth AI agent (GPT-5.1 with high reasoning effort) receives all prior outputs—'
        'typically 350,000 characters—for critical review using the restructured V2 prompt '
        'with explicit Deep Research guidance:'
    )
    items = [
        ('Factual accuracy', 'Biblical, historical, and grammatical verification'),
        ('Source verification', 'Ensuring claims are supported by retrieved sources'),
        ('Technical term definitions', 'Making scholarly vocabulary accessible'),
        ('Integration of cultural afterlife', 'Weaving in reception history, art, music references'),
        ('"Aha! moment" insights', 'Highlights discoveries only possible through comprehensive LLM analysis'),
        ('Style refinement', 'Eliminating AI tendencies toward breathlessness or jargon'),
    ]
    for term, desc in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(term + ': ')
        run.bold = True
        p.add_run(desc)
    
    doc.add_paragraph('The Master Editor produces two editions:')
    p = doc.add_paragraph(style='List Number')
    run = p.add_run('Main Edition: ')
    run.bold = True
    p.add_run('For sophisticated lay readers (New Yorker/Atlantic audience)')
    p = doc.add_paragraph(style='List Number')
    run = p.add_run('College Edition: ')
    run.bold = True
    p.add_run('More accessible version for undergraduate students')
    
    # Stage 6
    doc.add_heading('Stage 6: Document Generation', level=2)
    doc.add_paragraph('Python scripts format the commentary for publication:')
    items = [
        'Divine name modifications for study texts (יהוה → ה׳)',
        'Professional Word document formatting with Hebrew fonts',
        'Phonetic transcription italicization',
        'Bidirectional text handling',
        'Three output files: Main commentary, College edition, Combined document',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')
    
    # Sources Consulted
    doc.add_heading('Sources Consulted', level=1)
    
    doc.add_heading('Traditional Commentaries', level=2)
    commentaries = [
        ('Rashi', '11th century France — concise, focused on plain meaning'),
        ('Ibn Ezra', '12th century Spain — grammatical precision, philosophical depth'),
        ('Radak', '13th century Provence — expanded explanations, linguistic analysis'),
        ('Metzudat David', '18th century — accessible verse-by-verse commentary'),
        ('Malbim', '19th century — systematic analysis of synonyms and structure'),
        ('Meiri', '13th century Provence — philosophical and ethical interpretation'),
        ('Torah Temimah', '19th century — connects verses to rabbinic literature'),
    ]
    for name, desc in commentaries:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(name + ' ')
        run.bold = True
        p.add_run('(' + desc + ')')
    
    doc.add_heading('Modern Scholarship', level=2)
    items = [
        ('Rabbi Jonathan Sacks', 'contemporary British Orthodox philosopher'),
        ('Deep Web Research', 'cultural afterlife (art, music, literature, politics), scholarly debates, and reception history'),
    ]
    for name, desc in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(name + ' — ')
        run.bold = True
        p.add_run(desc)
    
    doc.add_heading('Linguistic Resources', level=2)
    items = [
        ('BDB Dictionary', 'Brown-Driver-Briggs Hebrew Lexicon'),
        ('Klein Dictionary', 'Comprehensive etymological dictionary'),
        ('Hebrew Concordance', '4-layer normalization (exact, voweled, consonantal, root)'),
        ('Figurative Language Database', 'Pre-analyzed metaphors and similes with curated insights'),
        ('Phonetic Transcription System', 'Reconstructed Biblical Hebrew phonology with dagesh distinction (b/v, k/kh, p/f)'),
    ]
    for name, desc in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(name + ' — ')
        run.bold = True
        p.add_run(desc)
    
    doc.add_heading('Ancient Translations', level=2)
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run('Septuagint (LXX) — ')
    run.bold = True
    p.add_run('Ancient Greek translation, crucial for textual criticism and understanding early interpretive traditions')
    
    doc.add_heading('Statistical Analysis', level=2)
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run('V6 Related Psalms Analysis — ')
    run.bold = True
    p.add_run('11,170 psalm pairs analyzed for shared roots, contiguous phrases, and skipgram patterns')
    
    # Key Technical Features
    doc.add_heading('Key Technical Features', level=1)
    
    doc.add_heading('Multi-Model Architecture', level=2)
    # Create a simple table
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Model'
    hdr_cells[1].text = 'Purpose'
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    models = [
        ('Claude Sonnet 4.5', 'Macro Analysis, Micro Analysis, Synthesis Writing'),
        ('Gemini 2.5 Pro', 'Synthesis fallback for large psalms (1M token context)'),
        ('Gemini 3 Pro', 'Figurative Curator (curated insights)'),
        ('GPT-5.1', 'Master Editorial Review (main and college editions)'),
    ]
    for i, (model, purpose) in enumerate(models, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = model
        row_cells[1].text = purpose
    
    doc.add_paragraph()
    
    doc.add_heading('Quality Assurance', level=2)
    items = [
        'Multi-pass validation with cross-checking between agents',
        'Enhanced quotation verification (sources quoted, not just cited)',
        'Phonetic accuracy verification against authoritative transcriptions',
        'Poetic punctuation ensuring verses show structural divisions',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Advanced Features', level=2)
    items = [
        ('Resume capability', 'Pipeline can restart from any step'),
        ('Special Instruction Pipeline', 'Author-directed revisions without altering standard outputs'),
        ('Strategic Verse Grouping', 'Prevents truncation in long psalms'),
        ('Priority-Based Figurative Trimming', 'LLM-decided term priority with lowest-priority content trimmed first'),
    ]
    for term, desc in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(term + ': ')
        run.bold = True
        p.add_run(desc)
    
    # The Result
    doc.add_heading('The Result', level=1)
    doc.add_paragraph('Each psalm commentary represents a synthesis of:')
    items = [
        'Traditional Jewish interpretation spanning a millennium',
        'Modern scholarly analysis and linguistic precision',
        'Cultural reception and artistic legacy',
        'AI-enhanced pattern recognition and cross-referencing',
        'Careful editorial refinement for accessibility and accuracy',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph(
        'The goal: scholarly depth made accessible to thoughtful readers, preserving the '
        'richness of Hebrew while opening it to English-language audiences.'
    )
    
    # Save
    output_path = r'c:\Users\ariro\OneDrive\Documents\Psalms\Documents\How Psalms Readers Guide works - January 2026.docx'
    doc.save(output_path)
    print(f'DOCX saved to: {output_path}')

if __name__ == '__main__':
    create_readers_guide_docx()
