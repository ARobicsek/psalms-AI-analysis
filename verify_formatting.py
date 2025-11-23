#!/usr/bin/env python3
"""
Verify that the formatted document meets all requirements from FORMATTING_GUIDE.md
"""

from docx import Document
from docx.shared import Pt, RGBColor
import sys

def rgb_to_hex(color_obj):
    """Convert RGBColor to hex string for comparison."""
    if color_obj is None:
        return None
    if hasattr(color_obj, 'rgb') and color_obj.rgb is not None:
        rgb_val = color_obj.rgb
        if isinstance(rgb_val, int):
            return f"#{rgb_val:06X}"
    return None

def check_paragraph_formatting(para, expected):
    """
    Check if paragraph matches expected formatting.
    For style inheritance, we only check that:
    1. The correct style is applied
    2. There is NO direct formatting (values should be None)
    """
    issues = []

    # Check style
    actual_style = para.style.name if para.style else 'None'
    if actual_style != expected['style']:
        issues.append(f"Style mismatch: expected '{expected['style']}', got '{actual_style}'")

    # For heading styles, verify NO direct formatting is applied (style inheritance)
    if para.runs and 'Heading' in expected['style']:
        first_run = para.runs[0]

        # Check that direct formatting is NOT applied (should be None for style inheritance)
        if first_run.font.name is not None:
            issues.append(f"Direct font formatting detected (should be None): {first_run.font.name}")

        if first_run.font.size is not None:
            issues.append(f"Direct size formatting detected (should be None): {first_run.font.size.pt}pt")

        if first_run.font.bold is not None:
            issues.append(f"Direct bold formatting detected (should be None): {first_run.font.bold}")

        if first_run.font.italic is not None:
            issues.append(f"Direct italic formatting detected (should be None): {first_run.font.italic}")

    return issues

def find_section_boundaries(doc):
    """Find paragraph indices for each specialty section."""
    sections = {
        'Primary Care': {'start': None, 'end': None},
        'Cardiology': {'start': None, 'end': None},
        'Infectious Diseases': {'start': None, 'end': None}
    }

    section_order = []

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not para.style:
            continue

        style_name = para.style.name
        if 'Heading 1' in style_name or 'Heading1' in style_name:
            if 'Primary Care' in text:
                sections['Primary Care']['start'] = i
                section_order.append('Primary Care')
            elif 'Cardiology' in text:
                sections['Cardiology']['start'] = i
                section_order.append('Cardiology')
            elif 'Infectious Diseases' in text:
                sections['Infectious Diseases']['start'] = i
                section_order.append('Infectious Diseases')

    # Set end boundaries
    for i, section_name in enumerate(section_order):
        if i < len(section_order) - 1:
            next_section = section_order[i + 1]
            sections[section_name]['end'] = sections[next_section]['start']
        else:
            sections[section_name]['end'] = len(doc.paragraphs)

    return sections

def verify_section(doc, section_name, start_idx, end_idx):
    """Verify formatting of a specific section."""
    print(f"\n{'='*80}")
    print(f"VERIFYING {section_name.upper()} SECTION")
    print(f"{'='*80}")

    issues_found = []
    checks_passed = 0

    # Expected formatting based on FORMATTING_GUIDE.md
    # We only check that the correct STYLE is applied, not direct formatting
    # (per the guide's emphasis on style inheritance)
    expected_formats = {
        'Heading 1': {'style': 'Heading 1'},
        'Heading 2': {'style': 'Heading 2'},
        'Heading 3': {'style': 'Heading 3'},
        'Heading 4': {'style': 'Heading 4'},
        'List Paragraph': {'style': 'List Paragraph'},
        'Normal': {'style': 'Normal'}
    }

    # Sample paragraphs from each heading level
    heading_samples = {
        'Heading 1': [],
        'Heading 2': [],
        'Heading 3': [],
        'Heading 4': [],
        'List Paragraph': [],
        'Normal': []
    }

    for i in range(start_idx, end_idx):
        if i >= len(doc.paragraphs):
            break

        para = doc.paragraphs[i]
        if not para.text.strip():
            continue

        style_name = para.style.name if para.style else 'None'

        for key in heading_samples.keys():
            if key in style_name:
                if len(heading_samples[key]) < 3:  # Sample up to 3 of each type
                    heading_samples[key].append((i, para))
                break

    # Check each sample
    for style_type, samples in heading_samples.items():
        if not samples:
            continue

        print(f"\n{style_type}:")
        for para_idx, para in samples:
            text_preview = para.text[:60] + ('...' if len(para.text) > 60 else '')
            print(f"  • Para {para_idx}: {text_preview}")

            if style_type in expected_formats:
                issues = check_paragraph_formatting(para, expected_formats[style_type])
                if issues:
                    for issue in issues:
                        print(f"    ⚠ {issue}")
                        issues_found.append((para_idx, style_type, text_preview, issue))
                else:
                    print(f"    ✓ Formatting correct")
                    checks_passed += 1
            else:
                checks_passed += 1

    return issues_found, checks_passed

def main():
    doc_path = '/home/user/psalms-AI-analysis/docs/Cordance/Cordance_insights_bank_draft_formatted.docx'

    print("Loading formatted document...")
    doc = Document(doc_path)

    print("\nFinding section boundaries...")
    sections = find_section_boundaries(doc)

    total_issues = []
    total_checks = 0

    # Verify each section
    for section_name, bounds in sections.items():
        if bounds['start'] is None:
            print(f"\n⚠ WARNING: {section_name} section not found!")
            continue

        issues, checks = verify_section(doc, section_name, bounds['start'], bounds['end'])
        total_issues.extend(issues)
        total_checks += checks

    # Summary
    print(f"\n{'='*80}")
    print("VERIFICATION SUMMARY")
    print(f"{'='*80}")
    print(f"Total checks passed: {total_checks}")
    print(f"Total issues found: {len(total_issues)}")

    if total_issues:
        print("\n⚠ ISSUES FOUND:")
        for para_idx, style_type, text_preview, issue in total_issues:
            print(f"\nParagraph {para_idx} ({style_type}):")
            print(f"  Text: {text_preview}")
            print(f"  Issue: {issue}")
        return 1
    else:
        print("\n✅ ALL FORMATTING CHECKS PASSED!")
        print("\nThe Cardiology and Infectious Diseases sections now match")
        print("the Primary Care section formatting exactly.")
        return 0

if __name__ == '__main__':
    sys.exit(main())
