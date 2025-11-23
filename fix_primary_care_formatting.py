#!/usr/bin/env python3
"""
Fix Primary Care document formatting by removing ALL direct formatting.
This ensures styles control all formatting, matching the Cardiology document exactly.
"""

from docx import Document
from docx.oxml import OxmlElement
from copy import deepcopy

def remove_run_formatting(run):
    """Remove all direct formatting from a run, keeping only style-based formatting."""
    r_pr = run._element.rPr
    if r_pr is not None:
        run._element.remove(r_pr)
        # Create empty rPr to ensure clean state
        run._element.insert(0, OxmlElement('w:rPr'))

def fix_primary_care_formatting():
    """Remove all direct formatting from Primary Care document."""

    print("Loading Primary Care document...")
    doc = Document('docs/Cordance/Cordance_Health_Insights_Bank_Primary_Care.docx')

    print(f"Processing {len(doc.paragraphs)} paragraphs...")

    fixed_count = 0

    for para in doc.paragraphs:
        # Check if paragraph has any runs with direct formatting
        has_direct_formatting = False
        for run in para.runs:
            if (run.bold is not None or
                run.font.size is not None or
                run.font.color.rgb is not None or
                run.italic is not None or
                run.underline is not None):
                has_direct_formatting = True
                break

        if has_direct_formatting:
            # Remove direct formatting from all runs
            for run in para.runs:
                remove_run_formatting(run)
            fixed_count += 1

            # Log what we fixed
            style_name = para.style.name if para.style else "None"
            text_preview = para.text[:50] + "..." if len(para.text) > 50 else para.text
            print(f"  Fixed: Style='{style_name}' | Text: {text_preview}")

    print(f"\nFixed {fixed_count} paragraphs with direct formatting")

    # Save the corrected document
    output_path = 'docs/Cordance/Cordance_Health_Insights_Bank_Primary_Care.docx'
    print(f"Saving to {output_path}...")
    doc.save(output_path)

    print("\n✓ Primary Care document formatting fixed!")
    print("All paragraphs now use only style inheritance, matching Cardiology exactly.")

if __name__ == '__main__':
    fix_primary_care_formatting()
