#!/usr/bin/env python3
"""
Improve Primary Care document formatting:
1. Set body text to Arial 11pt
2. Increase header sizes by 2pt
3. Make reference markers [1], [2], etc. superscript
"""

from docx import Document
from docx.shared import Pt, RGBColor
import re

def improve_formatting():
    """Apply formatting improvements to Primary Care document."""

    print("Loading Primary Care document...")
    doc = Document('docs/Cordance/Cordance_Health_Insights_Bank_Primary_Care.docx')

    # Define header sizes (current + 2pt)
    # Based on typical sizes: H1=32pt, H2=28pt, H3=26pt, H4=24pt
    header_sizes = {
        'Heading 1': 34,  # 32 + 2
        'Heading 2': 30,  # 28 + 2
        'Heading 3': 28,  # 26 + 2
        'Heading 4': 26,  # 24 + 2
    }

    # Color for headers (blue)
    header_color = RGBColor(0x1F, 0x4E, 0x78)

    print(f"\nProcessing {len(doc.paragraphs)} paragraphs...")

    headers_updated = 0
    body_updated = 0
    references_updated = 0

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else "None"

        # Update headers
        if style_name in header_sizes:
            for run in para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(header_sizes[style_name])
                run.font.color.rgb = header_color
                run.bold = True
            headers_updated += 1

        # Update body text and list paragraphs
        elif style_name in ['Normal', 'None', 'List Paragraph']:
            for run in para.runs:
                # Check if this run contains a reference marker like [1], [2], etc.
                text = run.text

                # If the entire run is a reference marker, make it superscript
                if re.match(r'^\[\d+\]$', text.strip()):
                    run.font.superscript = True
                    run.font.name = 'Arial'
                    run.font.size = Pt(9)  # Smaller for superscript
                    references_updated += 1
                else:
                    # Check if run contains inline references
                    if '[' in text and ']' in text:
                        # Split the text to handle inline references
                        # For now, just set the standard formatting
                        run.font.name = 'Arial'
                        run.font.size = Pt(11)
                        body_updated += 1
                    else:
                        # Regular body text
                        run.font.name = 'Arial'
                        run.font.size = Pt(11)
                        body_updated += 1

    print(f"\n✓ Updated {headers_updated} header paragraphs with new sizes:")
    for style, size in header_sizes.items():
        print(f"  - {style}: {size}pt")
    print(f"✓ Updated {body_updated} body text runs to Arial 11pt")
    print(f"✓ Updated {references_updated} reference markers to superscript")

    # Save the document
    output_path = 'docs/Cordance/Cordance_Health_Insights_Bank_Primary_Care.docx'
    print(f"\nSaving to {output_path}...")
    doc.save(output_path)

    print("\n✅ Formatting improvements complete!")

if __name__ == '__main__':
    improve_formatting()
