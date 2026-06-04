"""
Generates a print-ready Microsoft Word (.docx) document from the pipeline outputs.
This is designed to mirror the structure and content of the print_ready.md file.

This script takes the final edited introduction and verse commentaries,
parses them, and uses the python-docx library to create a well-formatted
Word document that preserves bidirectional text formatting for Hebrew and English.
"""

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns 

# Handle imports for both module and script usage
if __name__ == '__main__' and __package__ is None:
    sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent))
    from src.data_sources.tanakh_database import TanakhDatabase
    from src.utils.divine_names_modifier import DivineNamesModifier
    from src.data_sources.sefaria_client import strip_sefaria_footnotes
else:
    from ..data_sources.tanakh_database import TanakhDatabase
    from .divine_names_modifier import DivineNamesModifier
    from ..data_sources.sefaria_client import strip_sefaria_footnotes


def add_page_number(paragraph):
    """
    Adds a page number field to the given paragraph.
    The paragraph should be in a footer.
    """
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Create the run that will contain the page number field
    run = paragraph.add_run()
    
    # Create the complex field for the page number
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(ns.qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(ns.qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(ns.qn('w:fldCharType'), 'end')

    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_end)

class DocumentGenerator:
    """Generates a .docx commentary from pipeline artifacts."""

    # Complex-script font used for native-RTL Hebrew runs. Times New Roman ships with
    # full Hebrew + nikud coverage and matches the surrounding serif body text.
    HEBREW_FONT = 'Times New Roman'

    def __init__(self, psalm_num: int, intro_path: Path, verses_path: Path, stats_path: Path, output_path: Path, reader_questions_path: Optional[Path] = None):
        self.psalm_num = psalm_num
        self.intro_path = intro_path
        self.verses_path = verses_path
        self.stats_path = stats_path
        self.output_path = output_path
        self.reader_questions_path = reader_questions_path
        self.document = Document()
        self._set_default_styles()
        self.modifier = DivineNamesModifier()



    @staticmethod
    def _is_primarily_hebrew(text: str) -> bool:
        """
        Detect if a line is a standalone Hebrew verse line (not mixed Hebrew/English).
        
        A line is a standalone Hebrew verse if:
        - Contains sof-pasuq (׃) - the Hebrew end-of-verse marker
        - Has very few or no ASCII letters (< 5% of total letters)
        
        This is intentionally strict to avoid triggering on English commentary
        paragraphs that contain many Hebrew quotations.
        """
        if not text.strip():
            return False
        
        # Must have sof-pasuq to be considered a complete Hebrew verse line
        has_sof_pasuq = '׃' in text
        if not has_sof_pasuq:
            return False
        
        # Count Hebrew letters (U+05D0-U+05EA)
        hebrew_letters = len(re.findall(r'[\u05D0-\u05EA]', text))
        # Count ASCII letters
        ascii_letters = len(re.findall(r'[a-zA-Z]', text))
        
        total_letters = hebrew_letters + ascii_letters
        if total_letters == 0:
            return False
        
        # Require very low ASCII ratio (< 5%) to be considered primarily Hebrew
        # This excludes mixed paragraphs like "This is the psalm's theological center..."
        ascii_ratio = ascii_letters / total_letters
        
        return ascii_ratio < 0.05





    # ------------------------------------------------------------------
    # Native bidirectional rendering (logical order + w:rtl). This replaces
    # the legacy "reverse the Hebrew + LEFT-TO-RIGHT OVERRIDE" approach, which
    # garbled word order across markdown spans and broke line-wrapping. Word's
    # own bidi engine handles ordering, nikud, and wrapping correctly when each
    # Hebrew run carries w:rtl — exactly how the verse table and Arabic already work.
    # ------------------------------------------------------------------
    @staticmethod
    def _segment_by_script(text: str) -> List[tuple]:
        """Split text (kept in LOGICAL order — NOT reversed) into [(segment, is_hebrew), ...].

        Whitespace / maqqef / ASCII-hyphen that sits *between* two Hebrew words is glued
        into the Hebrew segment, so a multi-word Hebrew phrase becomes a single RTL run
        (and a trailing space before English correctly stays with the English run)."""
        def is_heb(c: str) -> bool:
            return '֐' <= c <= '׿'
        GLUE = {' ', ' ', '־', '-'}  # space, nbsp, maqqef, ASCII hyphen
        APOS = {"'", "’"}  # ASCII / curly apostrophe used as a Hebrew abbreviation mark
        n = len(text)
        segs: List[tuple] = []
        cur: List[str] = []
        cur_heb: Optional[bool] = None
        for i, c in enumerate(text):
            if is_heb(c):
                ch = True
            elif c in APOS and i > 0 and is_heb(text[i - 1]):
                ch = True  # geresh-style abbreviation mark stays with its Hebrew letter
            elif c in GLUE and cur_heb:
                j = i + 1
                while j < n and (text[j] in GLUE or text[j] in APOS):
                    j += 1
                ch = (j < n and is_heb(text[j]))
            else:
                ch = False
            if cur_heb is None:
                cur_heb = ch
            if ch == cur_heb:
                cur.append(c)
            else:
                segs.append((''.join(cur), bool(cur_heb)))
                cur, cur_heb = [c], ch
        if cur:
            segs.append((''.join(cur), bool(cur_heb)))
        return segs

    def _mark_run_hebrew(self, run, font_size: Optional[int] = None):
        """Mark a run as native RTL Hebrew: complex-script font + w:rtl (no reversal)."""
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(ns.qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(ns.qn('w:cs'), self.HEBREW_FONT)
        if rPr.find(ns.qn('w:rtl')) is None:
            rPr.append(OxmlElement('w:rtl'))
        if font_size is not None:
            szCs = rPr.find(ns.qn('w:szCs'))
            if szCs is None:
                szCs = OxmlElement('w:szCs')
                rPr.append(szCs)
            szCs.set(ns.qn('w:val'), str(int(font_size * 2)))

    def _add_inline_runs(self, paragraph, text: str, *, bold: bool = False, italic: bool = False,
                          set_font: bool = False, font_name: str = 'Aptos', font_size: int = 12):
        """Add `text` to `paragraph`, emitting Hebrew as native RTL runs (logical order)
        and everything else as ordinary LTR runs.

        This is the single funnel that all prose/markdown code paths use for mixed
        Hebrew/English text. `text` must already have divine-name modification applied.
        """
        for seg_text, is_heb in self._segment_by_script(text):
            if not seg_text:
                continue
            run = paragraph.add_run(seg_text)
            run.bold = bold
            run.italic = italic
            if is_heb:
                self._mark_run_hebrew(run, font_size=font_size if set_font else None)
            elif set_font:
                run.font.name = font_name
                run.font.size = Pt(font_size)

    @staticmethod
    def _set_style_complex_size(style, pt: int):
        """Set the complex-script size (w:szCs) on a paragraph style. python-docx's
        font.size only sets w:sz (Latin); without a matching szCs, Hebrew (a complex
        script) falls back to Word's default size and renders too small."""
        rPr = style.element.get_or_add_rPr()
        szCs = rPr.find(ns.qn('w:szCs'))
        if szCs is None:
            szCs = OxmlElement('w:szCs')
            rPr.append(szCs)
        szCs.set(ns.qn('w:val'), str(int(pt * 2)))

    def _add_primarily_hebrew_line(self, text: str, style: str = 'Normal', font_size: int = 13):
        """Render a standalone primarily-Hebrew line (e.g. a verse header) as native RTL
        run(s) — one per soft-break line — at `font_size`. The whole line is a single RTL
        run so internal punctuation (';', sof-pasuq) between cola keeps correct R->L order;
        segmenting it in an LTR paragraph would split the punctuation into its own LTR run
        and make Word reorder the cola as separate islands."""
        p = self.document.add_paragraph(style=style)
        self._set_paragraph_ltr(p)
        lines = text.split('\n')
        for i, line in enumerate(lines):
            if line.strip():
                run = p.add_run(line)
                self._mark_run_hebrew(run, font_size=font_size)
                run.font.size = Pt(font_size)
            if i < len(lines) - 1:
                p.add_run().add_break()
        return p

    def _set_default_styles(self):
        """Set default font for the document."""
        # Set default font
        font = self.document.styles['Normal'].font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # Set complex-script fallback font on Normal style so Arabic/Persian text
        # renders correctly (Aptos and Times New Roman lack Arabic glyphs in some installs)
        normal_rPr = self.document.styles['Normal'].element.get_or_add_rPr()
        normal_rFonts = normal_rPr.find(ns.qn('w:rFonts'))
        if normal_rFonts is None:
            normal_rFonts = OxmlElement('w:rFonts')
            normal_rPr.append(normal_rFonts)
        normal_rFonts.set(ns.qn('w:cs'), 'Times New Roman')
        # Hebrew (Times New Roman) reads visually smaller than the Aptos body text at the
        # same point size, so set the complex-script size 1pt larger (13 vs 12pt body) to
        # match. python-docx's font.size sets w:sz (Latin) but not w:szCs (complex script).
        self._set_style_complex_size(self.document.styles['Normal'], 13)

        # Create a new style for the sans-serif body text (intro and commentary)
        style = self.document.styles.add_style('BodySans', 1) # 1 for paragraph style
        style.base_style = self.document.styles['Normal']
        style.font.name = 'Aptos'
        style.font.size = Pt(12) # Aptos 12pt for intro and commentary

        # Set complex-script fallback on BodySans too (Aptos is Latin-only)
        body_rPr = style.element.get_or_add_rPr()
        body_rFonts = body_rPr.find(ns.qn('w:rFonts'))
        if body_rFonts is None:
            body_rFonts = OxmlElement('w:rFonts')
            body_rPr.append(body_rFonts)
        body_rFonts.set(ns.qn('w:cs'), 'Times New Roman')
        self._set_style_complex_size(style, 13)  # Hebrew in BodySans (intro/commentary) = 13pt (1pt over 12pt Latin)

        # Create a new style for the methodological summary
        summary_style = self.document.styles.add_style('SummaryText', 1) # 1 for paragraph style
        summary_style.base_style = self.document.styles['Normal']
        summary_style.font.name = 'Times New Roman'
        summary_style.font.size = Pt(9) # Times New Roman 9pt for summary
        self._set_style_complex_size(summary_style, 10)  # Hebrew in summary = 10pt (1pt over 9pt Latin)
        summary_style.paragraph_format.line_spacing = 0.8 # 70% line spacing

        # Set spacing for all heading levels to be tighter
        for i in range(1, 5):
            style = self.document.styles[f'Heading {i}']
            style.paragraph_format.space_before = Pt(12)
            style.paragraph_format.space_after = Pt(4)

        # Set paragraph spacing to be tighter (soft breaks)
        style = self.document.styles['Normal']
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(8)
        
        # Set document margins to 1 inch on all sides
        for section in self.document.sections:
            section.top_margin = Pt(72)  # 1 inch = 72 points
            section.bottom_margin = Pt(72)
            section.left_margin = Pt(72)
            section.right_margin = Pt(72)

    @staticmethod
    def _split_text_by_script(text: str, arabic_pattern) -> List[dict]:
        """Split text into segments of Arabic vs non-Arabic content.

        Returns a list of dicts: [{"text": "...", "is_arabic": bool}, ...]
        Adjacent characters of the same type are grouped together.
        """
        segments = []
        current_text = []
        current_is_arabic = None

        for char in text:
            is_arabic = bool(arabic_pattern.search(char))
            if current_is_arabic is None:
                current_is_arabic = is_arabic

            if is_arabic == current_is_arabic:
                current_text.append(char)
            else:
                segments.append({"text": ''.join(current_text), "is_arabic": current_is_arabic})
                current_text = [char]
                current_is_arabic = is_arabic

        if current_text:
            segments.append({"text": ''.join(current_text), "is_arabic": current_is_arabic})

        return segments

    def _join_rtl_runs_across_whitespace(self):
        """Post-pass: a *neutral-only* run (whitespace + punctuation such as the ';' or
        ':' between two Hebrew cola, with no Latin letters or digits) sitting between two
        RTL (Hebrew) runs must itself be RTL. Otherwise Word treats the two Hebrew runs as
        separate bidi islands and reorders them, garbling word order — e.g. when a
        bold/italic marker splits a Hebrew phrase, or a semicolon separates two cola.
        Runs containing Latin letters/digits, and Hebrew<->English boundary punctuation
        (only one neighbour is RTL), are left LTR (correct)."""
        def is_rtl(run):
            rPr = run._element.find(ns.qn('w:rPr'))
            return rPr is not None and rPr.find(ns.qn('w:rtl')) is not None

        def process(paragraph):
            runs = [r for r in paragraph.runs if r.text]
            for i, run in enumerate(runs):
                if re.search(r'[A-Za-z0-9]', run.text):
                    continue  # skip runs with strong-LTR content; only neutral runs qualify
                prev_rtl = i > 0 and is_rtl(runs[i - 1])
                next_rtl = i < len(runs) - 1 and is_rtl(runs[i + 1])
                if prev_rtl and next_rtl:
                    rPr = run._element.get_or_add_rPr()
                    if rPr.find(ns.qn('w:rtl')) is None:
                        rPr.append(OxmlElement('w:rtl'))

        for paragraph in self.document.paragraphs:
            process(paragraph)
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        process(paragraph)

    def _fix_complex_script_fonts(self):
        """Post-processing pass: fix fonts for Arabic/Persian/Urdu text in all runs.

        python-docx's font.name setter overrides the CS (complex script) font attribute,
        causing Arabic text to render as boxes when the Latin font (e.g. Aptos) lacks
        Arabic glyphs. This method iterates all runs and explicitly sets the CS font
        to 'Times New Roman' on any run containing Arabic-range characters.

        For runs containing ONLY Arabic text, we also set w:rtl to trigger correct
        shaping. For MIXED runs (Arabic + Latin/Hebrew), we split the run into
        separate runs so that w:rtl applies only to the Arabic segments — preventing
        Word's bidi algorithm from jumbling the non-Arabic text.
        """
        import copy
        # Arabic/Persian/Urdu range
        arabic_pattern = re.compile(r'[\u0600-\u06FF\u0750-\u077F\uFB50-\uFDFF\uFE70-\uFEFF]')
        # CJK Unified Ideographs (4E00-9FFF), Hiragana (3040-309F), Katakana (30A0-30FF)
        cjk_pattern = re.compile(r'[\u4E00-\u9FFF\u3040-\u309F\u30A0-\u30FF]')

        fixed_count = 0
        split_count = 0

        def _has_non_arabic_letters(text, arabic_pat):
            """Check if text contains any non-Arabic letter characters."""
            for ch in text:
                if ch.isalpha() and not arabic_pat.search(ch):
                    return True
            return False

        def process_paragraph_runs(paragraph):
            """Process all runs in a paragraph, splitting mixed Arabic/non-Arabic runs."""
            nonlocal fixed_count, split_count

            # Collect runs to process (iterate over a snapshot since we may modify the XML)
            runs_to_process = list(paragraph.runs)

            for run in runs_to_process:
                if not run.text:
                    continue

                text = run.text
                has_arabic = bool(arabic_pattern.search(text))
                has_non_arabic = _has_non_arabic_letters(text, arabic_pattern)

                # CJK processing (independent of Arabic handling)
                if cjk_pattern.search(text):
                    rPr = run._element.get_or_add_rPr()
                    rFonts = rPr.find(ns.qn('w:rFonts'))
                    if rFonts is None:
                        rFonts = OxmlElement('w:rFonts')
                        rPr.insert(0, rFonts)
                    rFonts.set(ns.qn('w:eastAsia'), 'DengXian')
                    fixed_count += 1

                if not has_arabic:
                    continue

                if has_arabic and not has_non_arabic:
                    # Pure Arabic run — apply CS font + RTL directly
                    rPr = run._element.get_or_add_rPr()
                    rFonts = rPr.find(ns.qn('w:rFonts'))
                    if rFonts is None:
                        rFonts = OxmlElement('w:rFonts')
                        rPr.insert(0, rFonts)
                    rFonts.set(ns.qn('w:cs'), 'Times New Roman')
                    if rPr.find(ns.qn('w:rtl')) is None:
                        rtl_elem = OxmlElement('w:rtl')
                        rPr.append(rtl_elem)
                    fixed_count += 1
                else:
                    # MIXED run — split into Arabic vs non-Arabic segments
                    segments = self._split_text_by_script(text, arabic_pattern)
                    if len(segments) <= 1:
                        # Shouldn't happen given the checks above, but handle gracefully
                        rPr = run._element.get_or_add_rPr()
                        rFonts = rPr.find(ns.qn('w:rFonts'))
                        if rFonts is None:
                            rFonts = OxmlElement('w:rFonts')
                            rPr.insert(0, rFonts)
                        rFonts.set(ns.qn('w:cs'), 'Times New Roman')
                        fixed_count += 1
                        continue

                    # Get the original run's XML element and its parent
                    original_r = run._element
                    parent = original_r.getparent()

                    # Save the original run properties (deep copy)
                    original_rPr = original_r.find(ns.qn('w:rPr'))

                    # Insert new runs BEFORE the original, then remove the original
                    for seg in segments:
                        new_r = OxmlElement('w:r')

                        # Copy run properties from original
                        if original_rPr is not None:
                            new_rPr = copy.deepcopy(original_rPr)
                        else:
                            new_rPr = OxmlElement('w:rPr')

                        if seg["is_arabic"]:
                            # Set CS font on Arabic segment
                            rFonts = new_rPr.find(ns.qn('w:rFonts'))
                            if rFonts is None:
                                rFonts = OxmlElement('w:rFonts')
                                new_rPr.insert(0, rFonts)
                            rFonts.set(ns.qn('w:cs'), 'Times New Roman')
                            # Add w:rtl for Arabic segment
                            if new_rPr.find(ns.qn('w:rtl')) is None:
                                rtl_elem = OxmlElement('w:rtl')
                                new_rPr.append(rtl_elem)
                        else:
                            # Non-Arabic segment: ensure NO w:rtl
                            existing_rtl = new_rPr.find(ns.qn('w:rtl'))
                            if existing_rtl is not None:
                                new_rPr.remove(existing_rtl)

                        new_r.append(new_rPr)

                        # Add the text element
                        new_t = OxmlElement('w:t')
                        new_t.text = seg["text"]
                        # Preserve whitespace
                        new_t.set(ns.qn('xml:space'), 'preserve')
                        new_r.append(new_t)

                        parent.insert(list(parent).index(original_r), new_r)

                    # Remove the original run
                    parent.remove(original_r)
                    split_count += 1
                    fixed_count += 1

        for paragraph in self.document.paragraphs:
            process_paragraph_runs(paragraph)

        # Also check table cells
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        process_paragraph_runs(paragraph)

        if fixed_count > 0:
            print(f"  Fixed complex-script fonts on {fixed_count} run(s) containing Arabic/CJK text.")
        if split_count > 0:
            print(f"  Split {split_count} mixed-script run(s) into separate Arabic/non-Arabic segments.")

    def _set_paragraph_ltr(self, paragraph):
        """
        Explicitly set paragraph direction to LTR at the XML level.
        This prevents Word's bidi algorithm from reordering runs when there's mixed Hebrew/English.
        """
        pPr = paragraph._element.get_or_add_pPr()

        # Set bidi to 0 (LTR) - if the element exists and is 1, it means RTL
        bidi_elem = pPr.find(ns.qn('w:bidi'))
        if bidi_elem is not None:
            pPr.remove(bidi_elem)
        # Not setting w:bidi or setting it to 0 means LTR
        # We'll explicitly set it to 0 for clarity
        bidi_elem = OxmlElement('w:bidi')
        bidi_elem.set(ns.qn('w:val'), '0')
        pPr.append(bidi_elem)

    def _set_run_font_xml(self, run, font_name='Aptos', font_size=12):
        """
        Set font and size at the XML level for comprehensive coverage.
        This ensures the font applies to all character ranges (ASCII, complex script, etc.).
        This approach is more reliable than the high-level API for Hebrew text.
        """
        rPr = run._element.get_or_add_rPr()

        # Set rFonts element with all font attributes
        rFonts = rPr.find(ns.qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)

        # Set font for all ranges: ascii, hAnsi (high ANSI), and cs (complex scripts)
        rFonts.set(ns.qn('w:ascii'), font_name)
        rFonts.set(ns.qn('w:hAnsi'), font_name)
        rFonts.set(ns.qn('w:cs'), font_name)

        # Set size at XML level
        sz = rPr.find(ns.qn('w:sz'))
        if sz is None:
            sz = OxmlElement('w:sz')
            rPr.append(sz)
        sz.set(ns.qn('w:val'), str(font_size * 2))  # Word uses half-points

        # Set complex script size
        szCs = rPr.find(ns.qn('w:szCs'))
        if szCs is None:
            szCs = OxmlElement('w:szCs')
            rPr.append(szCs)
        szCs.set(ns.qn('w:val'), str(font_size * 2))  # Word uses half-points

    def _split_long_hebrew_block(self, text: str):
        """
        Check if text contains a bare Hebrew segment of 6+ words.

        Long inline Hebrew segments wrap awkwardly in LTR paragraphs because
        the LRO-reversed text wraps top-to-bottom (English style), making the
        beginning of the Hebrew quote appear on the LOWER line. For these long
        segments, it's better to extract them into standalone RTL paragraphs
        where Word handles Hebrew line-wrapping natively.

        Returns (before, hebrew, after) tuple if found, or None.
        """
        heb_word = r'[\u05D0-\u05EA\u05F3\u05F4][\u0590-\u05FF]*'
        separator = r'(?:[\s:;,./?!()\[\]"\'\-–—\u2026\u05BE\u05C0\u05C3]|\*{1,2})+'
        # Match 6+ consecutive Hebrew words
        bare_hebrew = rf'(?<![\u05D0-\u05EA\u0590-\u05FF])({heb_word}(?:{separator}{heb_word}){{5,}})(?![\u05D0-\u05EA])'

        match = re.search(bare_hebrew, text)
        if not match:
            return None

        hebrew = match.group(1).strip()

        # Skip verse quotations — these contain sof-pasuq (׃) and are rendered inline
        # as native RTL runs (Aptos 12pt) by _add_inline_runs, not as standalone blocks.
        if '׃' in hebrew:
            return None

        before = text[:match.start()].rstrip()
        after = text[match.end():].lstrip()
        # Strip orphaned leading comma/punctuation from the continuation text
        # (the Hebrew was inline, so the original often had ", translation...")
        after = re.sub(r'^[,;:]\s*', '', after)
        return (before, hebrew, after)

    def _add_hebrew_block_paragraph(self, hebrew_text: str, style: str = 'Normal', font_name: str = 'Times New Roman'):
        """
        Add a standalone RTL paragraph for long Hebrew text.

        Uses native RTL paragraph direction (w:bidi=1) so Word handles
        line-wrapping correctly — first line on the right, continuation
        wrapping to the next line from the right, matching natural Hebrew
        reading direction.
        """
        from docx.shared import Inches

        p = self.document.add_paragraph(style=style)

        # Set paragraph to RTL for native Hebrew rendering
        pPr = p._element.get_or_add_pPr()
        bidi_elem = pPr.find(ns.qn('w:bidi'))
        if bidi_elem is not None:
            pPr.remove(bidi_elem)
        bidi_elem = OxmlElement('w:bidi')
        bidi_elem.set(ns.qn('w:val'), '1')  # RTL
        pPr.append(bidi_elem)

        # Right-align for RTL block
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Indent slightly to distinguish as a quotation
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.right_indent = Inches(0.3)

        # Add RLM if the block ends in a period, colon, or semicolon 
        # so Word's native bidi engine doesn't incorrectly send it to the right side
        hebrew_count = len(re.findall(r'[\u05D0-\u05EA]', hebrew_text))
        if hebrew_count >= 3 and re.search(r'[.;:,!?]\s*$', hebrew_text):
            hebrew_text = hebrew_text.rstrip() + '\u200F'

        # Parse **bold** markers and create separate runs
        parts = re.split(r'\*\*(.+?)\*\*', hebrew_text)
        for i, part in enumerate(parts):
            if not part:
                continue
            run = p.add_run(part)
            self._set_run_font_xml(run, font_name=font_name, font_size=13)
            if i % 2 == 1:  # Odd indices are the bold-captured groups
                run.bold = True

    def _add_paragraph_with_markdown(self, text: str, style: str = 'Normal'):
        """Adds a paragraph, parsing basic markdown for bold/italics, including nested formatting."""
        # Apply divine name modification to the entire paragraph text first.
        modified_text = self.modifier.modify_text(text)

        # Handle markdown headers first (check from most specific to least specific)
        if modified_text.startswith('####'):
            self.document.add_heading(modified_text.replace('####', '').strip(), level=4)
            return
        elif modified_text.startswith('###'):
            self.document.add_heading(modified_text.replace('###', '').strip(), level=3)
            return
        elif modified_text.startswith('##'):
            self.document.add_heading(modified_text.replace('##', '').strip(), level=2)
            return

        # Check for long bare Hebrew segments — render as standalone RTL block
        split = self._split_long_hebrew_block(modified_text)
        if split:
            before, hebrew, after = split
            
            # If 'before' or 'after' are merely orphaned punctuation, whitespace, 
            # or verse references with no English words, merge them back into the Hebrew 
            # block to prevent empty vertical space and orphaned dots on their own lines.
            if before and not re.search(r'[a-zA-Z]', before):
                hebrew = before + hebrew
                before = ""
                
            if after and not re.search(r'[a-zA-Z]', after):
                hebrew = hebrew + after
                after = ""
                
            # If the Hebrew block is at the very beginning of the string and spans a whole line, it's a verse
            font_to_use = 'Aptos' if not before and (not after or after.startswith('\n')) else 'Times New Roman'
                
            if before:
                self._add_paragraph_with_markdown(before, style=style)
            self._add_hebrew_block_paragraph(hebrew, style=style, font_name=font_to_use)
            if after:
                self._add_paragraph_with_markdown(after, style=style)
            return

        # A standalone primarily-Hebrew line renders as native RTL run(s) at 13pt so its
        # cola stay ordered (internal ';'/punctuation would otherwise split the run and
        # make Word reorder the cola as separate islands).
        if self._is_primarily_hebrew(modified_text):
            self._add_primarily_hebrew_line(modified_text, style=style, font_size=13)
            return

        # Handle bullet lists (lines starting with "- ")
        is_bullet = False
        if modified_text.startswith('- '):
            modified_text = modified_text[2:]  # Remove "- " prefix
            is_bullet = True

        # Handle block quotes (lines starting with ">")
        is_quote = False
        if modified_text.startswith('>'):
            # Check if it's an empty quote line (just ">" with optional spaces)
            quote_text = modified_text[1:].lstrip()  # Remove ">" and any leading spaces
            if not quote_text:
                # Empty quote line - just add a blank paragraph and return
                self.document.add_paragraph()
                return
            modified_text = quote_text
            is_quote = True

        p = self.document.add_paragraph(style=style)

        # Apply bullet formatting if this is a list item
        if is_bullet:
            p.style = 'List Bullet'

        # Apply quote formatting if this is a block quote
        if is_quote:
            from docx.shared import Inches
            p.paragraph_format.left_indent = Inches(0.5)

        # Explicitly set paragraph to LTR to prevent Word's bidi algorithm from reordering runs
        self._set_paragraph_ltr(p)

        # Use the centralized formatting method with font setting for bullets
        self._process_markdown_formatting(p, modified_text, set_font=is_bullet)

        # Make block quotes italic
        if is_quote:
            for run in p.runs:
                run.italic = True

    def _process_introduction_content(self, text: str, style: str = 'Normal'):
        """
        Process introduction content, preserving headers while properly handling blockquotes.
        This method processes line by line, treating headers specially but using
        _add_commentary_with_bullets logic for blockquotes.
        """
        lines = text.split('\n')
        i = 0

        while i < len(lines):
            line = lines[i].strip()

            # Skip empty lines
            if not line:
                i += 1
                continue

            # Handle liturgical section marker - convert to proper heading
            if line == '---LITURGICAL-SECTION-START---':
                self.document.add_heading('Modern Jewish Liturgical Use', level=2)
                i += 1
                continue

            # Handle headers
            if line.startswith('####'):
                self.document.add_heading(line.replace('####', '').strip(), level=4)
                i += 1
                continue
            elif line.startswith('###'):
                self.document.add_heading(line.replace('###', '').strip(), level=3)
                i += 1
                continue
            elif line.startswith('##'):
                self.document.add_heading(line.replace('##', '').strip(), level=2)
                i += 1
                continue

            # Handle blockquotes - collect consecutive blockquote lines
            if line.startswith('>'):
                quote_block = []
                while i < len(lines) and lines[i].strip() and lines[i].strip().startswith('>'):
                    quote_text = lines[i].strip()[1:].lstrip()  # Remove ">" and leading spaces
                    if quote_text:  # Only add non-empty quote lines
                        quote_block.append(quote_text)
                    i += 1

                # Add quote block as indented italic paragraphs
                from docx.shared import Inches
                for quote_text in quote_block:
                    p = self.document.add_paragraph(style=style)
                    p.paragraph_format.left_indent = Inches(0.5)
                    self._set_paragraph_ltr(p)
                    self._process_markdown_formatting(p, quote_text, set_font=True)
                    # Make the entire quote italic
                    for run in p.runs:
                        run.italic = True
                continue

            # Handle regular paragraphs
            # Only treat as bullet if it's a true markdown list item with clear list markers
            # and not regular text that happens to start with dash
            # True bullets typically are short, start with lowercase, or are part of a clear list pattern
            original_line = lines[i]  # Keep original line for indentation detection
            is_true_bullet = (
                original_line.lstrip().startswith('- ') and
                len(original_line.strip()) > 5 and
                # Don't treat as bullet if it starts with uppercase letter followed by common sentence patterns
                not (
                    original_line.lstrip()[2:3].isupper() and
                    (':' in original_line[:50] or 'In ' in original_line[:10] or 'As ' in original_line[:10] or 'For ' in original_line[:10])
                )
            )

            if is_true_bullet:
                # Bullet point - detect leading indentation
                leading_spaces = len(original_line) - len(original_line.lstrip())
                p = self.document.add_paragraph(style='List Bullet')
                self._set_paragraph_ltr(p)

                # Apply indentation if present (convert spaces to inches)
                if leading_spaces > 0:
                    from docx.shared import Inches
                    # Approximately 6 spaces = 0.25 inch, scale accordingly
                    indent_inches = leading_spaces * 0.04  # Rough conversion: 1 space ≈ 0.04 inches
                    p.paragraph_format.left_indent = Inches(indent_inches)

                bullet_text = original_line.lstrip()[2:]  # Remove leading spaces and "- " prefix
                self._process_markdown_formatting(p, bullet_text, set_font=True)
            else:
                # Regular paragraph - use _add_paragraph_with_markdown which handles dash characters properly
                self._add_paragraph_with_markdown(line, style=style)

            i += 1

    def _add_commentary_with_bullets(self, text: str, style: str = 'Normal', is_verse_commentary: bool = False):
        """
        Adds commentary text, intelligently handling bullet lists and regular text.
        Bullet list items (lines starting with "- ") are converted to proper Word bullets.
        Regular text blocks use soft breaks, with empty lines creating paragraph breaks.
        """
        lines = text.split('\n')
        i = 0

        while i < len(lines):
            line = lines[i]

            # Check if this is an empty line (paragraph break)
            if not line.strip():
                i += 1
                continue

            # Check if this is a true bullet item (not just regular text starting with dash)
            original_line = line  # Keep original line for indentation detection
            line_stripped = line.strip()
            is_true_bullet = (
                line_stripped.startswith('- ') and
                len(line_stripped) > 5 and
                # Don't treat as bullet if it starts with uppercase letter followed by common sentence patterns
                not (
                    line_stripped[2:3].isupper() and
                    (':' in line_stripped[:50] or 'In ' in line_stripped[:10] or 'As ' in line_stripped[:10] or 'For ' in line_stripped[:10])
                )
            )

            if is_true_bullet:
                # Collect consecutive bullet items with indentation info
                bullet_block = []
                while i < len(lines) and lines[i].strip() and lines[i].lstrip().startswith('- '):
                    # Apply same logic to consecutive lines
                    bullet_line = lines[i]
                    bullet_stripped = bullet_line.strip()
                    is_consecutive_bullet = (
                        len(bullet_stripped) > 5 and
                        not (
                            bullet_stripped[2:3].isupper() and
                            (':' in bullet_stripped[:50] or 'In ' in bullet_stripped[:10] or 'As ' in bullet_stripped[:10] or 'For ' in bullet_stripped[:10])
                        )
                    )
                    if is_consecutive_bullet:
                        # Store both the bullet text and indentation info
                        leading_spaces = len(bullet_line) - len(bullet_line.lstrip())
                        bullet_text = bullet_line.lstrip()[2:]  # Remove "- " prefix
                        bullet_block.append((bullet_text, leading_spaces))
                        i += 1
                    else:
                        break

                # Add each bullet as a separate paragraph with proper indentation
                for bullet_text, leading_spaces in bullet_block:
                    p = self.document.add_paragraph(style='List Bullet')
                    # Explicitly set paragraph to LTR to prevent Word's bidi algorithm from reordering runs
                    self._set_paragraph_ltr(p)

                    # Apply indentation if present (convert spaces to inches)
                    if leading_spaces > 0:
                        from docx.shared import Inches
                        # Approximately 6 spaces = 0.25 inch, scale accordingly
                        indent_inches = leading_spaces * 0.04  # Rough conversion: 1 space ≈ 0.04 inches
                        p.paragraph_format.left_indent = Inches(indent_inches)

                    # Process markdown formatting in bullet text with explicit font
                    self._process_markdown_formatting(p, bullet_text, set_font=True)
            # Check if this is a markdown heading (###, ##, ####)
            elif line_stripped.startswith('#'):
                # Parse heading level and text
                heading_match = re.match(r'^(#{1,6})\s+(.+)$', line_stripped)
                if heading_match:
                    hashes = heading_match.group(1)
                    heading_text = heading_match.group(2)
                    # Map markdown heading levels to Word heading levels
                    # ## = Heading 2, ### = Heading 3, #### = Heading 4
                    level = len(hashes)
                    word_heading_level = min(level, 4)  # Cap at Heading 4

                    # Add as a Word heading
                    p = self.document.add_heading(heading_text, level=word_heading_level)
                    self._set_paragraph_ltr(p)
                else:
                    # Malformed heading, treat as normal text
                    p = self.document.add_paragraph(style=style)
                    self._set_paragraph_ltr(p)
                    self._process_markdown_formatting(p, line_stripped, set_font=False)
                i += 1
            # Check if this is a block quote
            elif line_stripped.startswith('>'):
                # Collect consecutive block quote lines
                quote_block = []
                while i < len(lines) and lines[i].strip() and lines[i].strip().startswith('>'):
                    # Remove ">" and any leading spaces after it
                    quote_text = lines[i].strip()[1:].lstrip()
                    quote_block.append(quote_text)
                    i += 1

                # Add each quote line as a separate paragraph with indentation and italic
                from docx.shared import Inches
                for quote_text in quote_block:
                    if not quote_text:
                        # Empty quote line - add a blank paragraph
                        self.document.add_paragraph()
                    else:
                        # Non-empty quote line
                        p = self.document.add_paragraph(style=style)
                        p.paragraph_format.left_indent = Inches(0.5)
                        # Explicitly set paragraph to LTR to prevent Word's bidi algorithm from reordering runs
                        self._set_paragraph_ltr(p)
                        # Process markdown formatting in quote text with explicit font
                        self._process_markdown_formatting(p, quote_text, set_font=True)
                        # Make the entire quote italic
                        for run in p.runs:
                            run.italic = True
            else:
                # Collect consecutive non-bullet, non-quote, non-empty lines until we hit an empty line or special formatting
                text_block = []
                is_first_block = (i == 0)
                while i < len(lines):
                    if not lines[i].strip():
                        # We hit an empty line. Let's look ahead to see if both the current collected 
                        # block AND the next block are primarily Hebrew. If they are, we can bridge 
                        # the gap to keep them in the same paragraph with a soft break.
                        if text_block and all(self._is_primarily_hebrew(l) for l in text_block):
                            # Look ahead for next non-empty line
                            next_i = i + 1
                            while next_i < len(lines) and not lines[next_i].strip():
                                next_i += 1
                                
                            if next_i < len(lines):
                                next_line_stripped = lines[next_i].strip()
                                if self._is_primarily_hebrew(next_line_stripped) and not next_line_stripped.startswith('>') and not next_line_stripped.startswith('#') and not next_line_stripped.startswith('- '):
                                    # Bridge the gap! Add a newline to the block and skip the empty lines
                                    # No need to actually add the newline string, as join handles it later. 
                                    # But we do need to advance `i` past the empty lines.
                                    i = next_i
                                    continue
                        # If we didn't bridge the gap, break the block
                        break
                        
                    original_line = lines[i]
                    line_stripped = original_line.strip()
                    # Apply same bullet detection logic
                    is_bullet = (
                        line_stripped.startswith('- ') and
                        len(line_stripped) > 5 and
                        not (
                            line_stripped[2:3].isupper() and
                            (':' in line_stripped[:50] or 'In ' in line_stripped[:10] or 'As ' in line_stripped[:10] or 'For ' in line_stripped[:10])
                        )
                    )

                    # Stop if we hit a bullet, quote, or markdown heading
                    if not is_bullet and not line_stripped.startswith('>') and not line_stripped.startswith('#'):
                        text_block.append(original_line)
                        i += 1
                    else:
                        break

                # Add as a paragraph with soft breaks
                if text_block:
                    self._add_paragraph_with_soft_breaks(
                        '\n'.join(text_block), 
                        style=style, 
                        is_verse_header=(is_verse_commentary and is_first_block)
                    )

    def _process_markdown_formatting(self, paragraph, text, set_font=False):
        """
        Process markdown formatting (bold, italic, etc.) and add runs to the given paragraph.
        This is a helper for adding formatted text to an existing paragraph.

        Args:
            paragraph: The paragraph to add runs to
            text: The text to process
            set_font: If True, explicitly set Aptos 12pt font on all runs (for bullet lists)
        """
        modified_text = self.modifier.modify_text(text)
        # Match bold (**...**) and italic (*...*) patterns
        # Order matters: match ** before * to avoid incorrect splits
        parts = re.split(r'(\*\*.*?\*\*|__.*?__|\*.*?\*|_.*?_|`.*?`)', modified_text)

        for part in parts:
            if not part:
                continue
            if part.startswith('**') and part.endswith('**'):
                # Bold text - recursively process inner content for nested formatting
                inner_content = part[2:-2]
                self._add_formatted_content(paragraph, inner_content, bold=True, italic=False, set_font=set_font)
            elif part.startswith('__') and part.endswith('__'):
                # Bold text - recursively process inner content for nested formatting
                inner_content = part[2:-2]
                self._add_formatted_content(paragraph, inner_content, bold=True, italic=False, set_font=set_font)
            elif (part.startswith('*') and part.endswith('*')) or \
                 (part.startswith('_') and part.endswith('_')):
                self._add_inline_runs(paragraph, part[1:-1], italic=True, set_font=set_font)
            elif part.startswith('`') and part.endswith('`'):
                inner_content = part[1:-1]
                self._add_nested_formatting(paragraph, inner_content, base_italic=True)
                # Note: nested formatting handles its own font settings
            else:
                # Mixed Hebrew/English prose: native RTL runs (logical order, no reversal).
                self._add_inline_runs(paragraph, part, set_font=set_font)

    def _add_formatted_content(self, paragraph, text, bold=False, italic=False, set_font=False):
        """
        Add text content with specified formatting, recursively processing any nested markdown.
        This handles cases like **bold (*italic inside bold*)**.

        Args:
            paragraph: The paragraph to add runs to
            text: The text content to add
            bold: Whether the text should be bold
            italic: Whether the text should be italic
            set_font: If True, explicitly set Aptos 12pt font
        """
        # Check if there's nested formatting to process
        if '*' in text or '_' in text or '`' in text:
            # Split on italic markers to handle nested formatting
            parts = re.split(r'(\*.*?\*|_.*?_|`.*?`)', text)
            for part in parts:
                if not part:
                    continue
                if part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                    # Nested italic
                    self._add_inline_runs(paragraph, part[1:-1], bold=bold, italic=True, set_font=set_font)
                elif part.startswith('_') and part.endswith('_') and not part.startswith('__'):
                    # Nested italic
                    self._add_inline_runs(paragraph, part[1:-1], bold=bold, italic=True, set_font=set_font)
                elif part.startswith('`') and part.endswith('`'):
                    # Nested backtick - process with _add_nested_formatting but apply bold too
                    inner_content = part[1:-1]
                    # For backticks, we need to handle bold + italic
                    nested_parts = re.split(r'(\*\*.*?\*\*)', inner_content)
                    for nested_part in nested_parts:
                        if nested_part.startswith('**') and nested_part.endswith('**'):
                            run = paragraph.add_run(nested_part[2:-2])
                            run.bold = True
                            run.italic = True
                        else:
                            run = paragraph.add_run(nested_part)
                            run.bold = bold
                            run.italic = True
                        if set_font:
                            run.font.name = 'Aptos'
                            run.font.size = Pt(12)
                else:
                    # Regular text with base formatting (native RTL runs, logical order).
                    self._add_inline_runs(paragraph, part, bold=bold, italic=italic, set_font=set_font)
        else:
            # No nested formatting - native RTL runs (logical order).
            self._add_inline_runs(paragraph, text, bold=bold, italic=italic, set_font=set_font)

    def _add_paragraph_with_soft_breaks(self, text: str, style: str = 'Normal', is_verse_header: bool = False):
        """Adds a single paragraph, treating newlines as soft breaks, with nested formatting support."""
        modified_text = self.modifier.modify_text(text)

        # Check for long bare Hebrew segments — render as standalone RTL block
        split = self._split_long_hebrew_block(modified_text)
        if split:
            before, hebrew, after = split
            
            # If 'before' or 'after' are merely orphaned punctuation, whitespace, 
            # or verse references with no English words, merge them back into the Hebrew 
            # block to prevent empty vertical space and orphaned dots on their own lines.
            if before and not re.search(r'[a-zA-Z]', before):
                hebrew = before + hebrew
                before = ""
                
            if after and not re.search(r'[a-zA-Z]', after):
                hebrew = hebrew + after
                after = ""
                
            # Force Aptos exactly when this block is flagged as the core verse header.
            # Otherwise use the default Times New Roman block font for long quotes.
            font_to_use = 'Aptos' if is_verse_header else 'Times New Roman'
                
            if before:
                self._add_paragraph_with_soft_breaks(before, style=style, is_verse_header=is_verse_header)
            self._add_hebrew_block_paragraph(hebrew, style=style, font_name=font_to_use)
            if after:
                self._add_paragraph_with_soft_breaks(after, style=style, is_verse_header=False)
            return

        # Verse-header verse lines, and any standalone primarily-Hebrew line, render as a
        # single native RTL run (one per soft-break line) at 13pt. Keeping the whole line
        # in one RTL run gives correct native placement to cola separators (';'), ketiv/qere
        # [brackets], and the final sof-pasuq; rendered inline in an LTR paragraph, that edge
        # punctuation is pushed to the wrong side and the cola can be reordered.
        heb = len(re.findall(r'[א-ת]', modified_text))
        asc = len(re.findall(r'[a-zA-Z]', modified_text))
        heb_dominant = heb >= 2 and (asc == 0 or heb / (heb + asc) > 0.8)
        if self._is_primarily_hebrew(modified_text) or (is_verse_header and heb_dominant):
            self._add_primarily_hebrew_line(modified_text, style=style, font_size=13)
            return

        p = self.document.add_paragraph(style=style)

        # Explicitly set paragraph to LTR to prevent Word's bidi algorithm from reordering runs
        self._set_paragraph_ltr(p)

        # Split the entire text by markdown markers first
        parts = re.split(r'(\$\$|__.*?__|\*.*?\*|_.*?_|`.*?`)', modified_text)

        for part in parts:
            is_bold = (part.startswith('**') and part.endswith('**')) or \
                      (part.startswith('__') and part.endswith('__'))
            is_backtick = part.startswith('`') and part.endswith('`')
            is_italic = (part.startswith('*') and part.endswith('*')) or \
                        (part.startswith('_') and part.endswith('_')) or \
                        is_backtick

            content = part[2:-2] if is_bold else (part[1:-1] if is_italic else part)

            # Handle backticks with nested bold (for stressed syllables)
            if is_backtick:
                self._add_nested_formatting_with_breaks(p, content, base_italic=True)
            else:
                # Handle soft breaks within the content
                lines = content.split('\n')
                for i, line in enumerate(lines):
                    if line:
                        self._add_inline_runs(p, line, bold=is_bold, italic=is_italic)
                    if i < len(lines) - 1:
                        p.add_run().add_break()

    def _add_nested_formatting(self, paragraph, text: str, base_italic: bool = False):
        """
        Add text with nested formatting (e.g., **BOLD** inside italic context).
        Used for phonetic transcriptions where stressed syllables are in **BOLD CAPS**.

        Args:
            paragraph: The docx paragraph object to add runs to
            text: The text content (e.g., "tə-**HIL**-lāh")
            base_italic: Whether the base text should be italic (True for backtick context)
        """
        # Split by bold markers
        parts = re.split(r'(\$\$|\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Bold text (stressed syllable)
                run = paragraph.add_run(part[2:-2])
                run.bold = True
                run.italic = base_italic  # Maintain italic if in backtick context
            else:
                # Regular text
                run = paragraph.add_run(part)
                run.italic = base_italic

    def _add_nested_formatting_with_breaks(self, paragraph, text: str, base_italic: bool = False):
        """
        Add text with nested formatting AND support for soft breaks.
        Used for phonetic transcriptions in verse commentary.

        Args:
            paragraph: The docx paragraph object to add runs to
            text: The text content with possible newlines
            base_italic: Whether the base text should be italic
        """
        # First split by newlines
        lines = text.split('\n')
        for i, line in enumerate(lines):
            if line:
                # Then handle bold within each line
                parts = re.split(r'(\$\$|\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = paragraph.add_run(part[2:-2])
                        run.bold = True
                        run.italic = base_italic
                    else:
                        run = paragraph.add_run(part)
                        run.italic = base_italic
            if i < len(lines) - 1:
                paragraph.add_run().add_break()

    def _add_bilingual_verse(self, hebrew: str, english: str):
        """Adds a formatted bilingual verse line with Hebrew (RTL) and English (LTR)."""
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Apply divine name modification to the Hebrew text
        modified_hebrew = self.modifier.modify_text(hebrew)

        # Add Hebrew text with RTL properties
        hebrew_run = p.add_run(modified_hebrew)
        hebrew_run.font.name = 'SBL Hebrew'
        hebrew_run.font.size = Pt(14)
        hebrew_run.font.rtl = True

        # Add separator (tabs work well in Word)
        p.add_run('\t\t')

        # Add English text
        english_run = p.add_run(english)
        english_run.font.name = 'Times New Roman'
        english_run.font.size = Pt(12)
        english_run.italic = True

    def _parse_verse_commentary(self, content: str, psalm_text_data: Dict[int, Dict[str, str]]) -> List[Dict[str, str]]:
        """Parses the verse-by-verse commentary file.

        Handles both single verses (Verse 1) and verse ranges (Verses 21-25).
        For ranges, returns a single entry with 'number' as the range string (e.g., '21-25')
        and 'verse_range' as a tuple (start, end).
        """
        verses = []
        # Try multiple formats: "**Verse(s) X(-Y)**" (main format), "### Verse(s) X(-Y)" (college format), "Verse(s) X(-Y)" (fallback)
        # First try the expected bold format - matches both single and range verses
        verse_blocks = re.split(r'(?=^\*\*Verses? [\d\-–]+.*?\*\*\s*\n)', content, flags=re.MULTILINE)

        # If no verses found with bold format, try heading format (college uses this)
        if len(verse_blocks) <= 1:
            verse_blocks = re.split(r'(?=^#{2,} Verses? [\d\-–]+)', content, flags=re.MULTILINE)

        # If still no verses found, try plain format without any markdown
        if len(verse_blocks) <= 1:
            verse_blocks = re.split(r'(?=^Verses? [\d\-–]+)', content, flags=re.MULTILINE)

        for block in verse_blocks:
            block = block.strip()
            if not block.strip():
                continue

            lines = block.strip().split('\n', 1)  # Split only on the first newline

            # Try all formats for verse number matching (single or range)
            # Matches: **Verse 1**, **Verses 21-25**, **Verses 21–25 (description)**, etc.
            verse_num_match = re.match(r'^\*\*Verses? ([\d]+)(?:[\-–]([\d]+))?\s*.*?\*\*', lines[0])  # Bold format
            if not verse_num_match:
                # Heading format: ### Verse 1, ### Verses 21-25 (description), etc.
                verse_num_match = re.match(r'^#{2,} Verses? ([\d]+)(?:[\-–]([\d]+))?\s*', lines[0])
            if not verse_num_match:
                # Plain format: Verse 1, Verses 21-25, etc.
                verse_num_match = re.match(r'^Verses? ([\d]+)(?:[\-–]([\d]+))?\s*', lines[0])

            if not verse_num_match:
                continue

            start_verse = verse_num_match.group(1)
            end_verse = verse_num_match.group(2)  # None if single verse

            if end_verse:
                # It's a range
                verse_num = f"{start_verse}-{end_verse}"
                verse_range = (int(start_verse), int(end_verse))
            else:
                # Single verse
                verse_num = start_verse
                verse_range = None

            commentary_text = '\n'.join(lines[1:]).strip()

            # Get Hebrew/English text from the database data
            # For ranges, we'll get the first verse's text as representative
            verse_info = psalm_text_data.get(int(start_verse), {})
            hebrew_text = verse_info.get('hebrew', '[Hebrew text not found]')
            english_text = verse_info.get('english', '[English text not found]')

            verse_entry = {
                "number": verse_num,
                "hebrew": hebrew_text,
                "english": english_text,
                "commentary": commentary_text
            }

            if verse_range:
                verse_entry["verse_range"] = verse_range

            verses.append(verse_entry)

        return verses

    def _format_psalm_text(self, psalm_num: int, psalm_text_data: Dict[int, Dict[str, str]]):
        """Formats the full psalm text with Hebrew and English side-by-side."""
        self.document.add_heading(f"Psalm {psalm_num}", level=2)
        
        # Create a table with 2 columns
        table = self.document.add_table(rows=0, cols=2)
        table.autofit = True
        
        for verse_num in sorted(psalm_text_data.keys()):
            hebrew = psalm_text_data[verse_num].get('hebrew', '')
            english = psalm_text_data[verse_num].get('english', '')
            # Strip any footnote markers from the English text
            english = strip_sefaria_footnotes(english)
            modified_hebrew = self.modifier.modify_text(hebrew)

            row_cells = table.add_row().cells

            # Hebrew cell (left, but text is RTL)
            p_heb = row_cells[0].paragraphs[0]
            num_run = p_heb.add_run(f"{verse_num}. ")
            num_run.bold = True
            num_run.font.size = Pt(12)
            
            # Explicitly set Complex Script font size (24 half-points = 12pt)
            rPr_num = num_run._element.get_or_add_rPr()
            szCs_num = OxmlElement('w:szCs')
            szCs_num.set(ns.qn('w:val'), '24')
            rPr_num.append(szCs_num)

            heb_run = p_heb.add_run(modified_hebrew)
            heb_run.font.rtl = True
            heb_run.font.size = Pt(12)
            
            # Explicitly set Complex Script font size (24 half-points = 12pt)
            rPr_heb = heb_run._element.get_or_add_rPr()
            szCs_heb = OxmlElement('w:szCs')
            szCs_heb.set(ns.qn('w:val'), '24')
            rPr_heb.append(szCs_heb)

            # English cell (right)
            p_eng = row_cells[1].paragraphs[0]
            run_eng = p_eng.add_run(english)
            run_eng.font.name = 'Cambria'

    def _format_bibliographical_summary(self, stats: Dict[str, Any]) -> str:
        """Formats the stats dictionary into a readable string for the document."""
        # --- Analysis & Research Inputs ---
        analysis_data = stats.get('analysis', {}) or {}
        verse_count = analysis_data.get('verse_count', 0)
        
        # Fallback: if verse_count is 0 or missing (e.g. from a resumed run that 
        # skipped macro analysis), count verses from the database
        if not verse_count:
            try:
                db = TanakhDatabase()
                psalm_data = db.get_psalm(self.psalm_num)
                if psalm_data:
                    verse_count = len(psalm_data.verses)
            except Exception:
                verse_count = 'N/A'
        
        research_data = stats.get('research', {})
        lexicon_count = research_data.get('lexicon_entries_count', 'N/A')
        
        commentaries = research_data.get('commentary_counts', {})
        total_commentaries = sum(commentaries.values()) if commentaries else 'N/A'
        commentary_details = ""
        if commentaries:
            commentary_lines = [f"{c} ({n})" for c, n in sorted(commentaries.items())]
            commentary_details = f" ({'; '.join(commentary_lines)})"

        concordance_results = research_data.get('concordance_results', {}) or {}
        concordance_total = sum(concordance_results.values())

        # Concordance entries breakdown (query -> count)
        concordance_breakdown_str = ""
        # Filter out the legacy 'total_results' key to get per-query entries
        concordance_per_query = {k: v for k, v in concordance_results.items() if k != 'total_results'}
        if concordance_per_query:
            items = [f"{self.modifier.modify_text(q)} ({c})" for q, c in sorted(concordance_per_query.items())]
            concordance_breakdown_str = f" ({'; '.join(items)})"

        figurative_results = research_data.get('figurative_results', {}) or {}
        figurative_total = figurative_results.get('total_instances_used', 0) if isinstance(figurative_results, dict) else 0

        # Figurative parallels breakdown
        figurative_parallels = research_data.get('figurative_parallels_reviewed', {})
        figurative_breakdown_str = ""
        if figurative_parallels:
            items = [f"{v} ({c})" for v, c in sorted(figurative_parallels.items())]
            figurative_breakdown_str = f" ({'; '.join(items)})"

        sacks_count = research_data.get('sacks_references_count', 0)
        related_psalms_count = research_data.get('related_psalms_count', 0)
        related_psalms_list = research_data.get('related_psalms_list', [])

        # Format related psalms display: count and list
        if related_psalms_count > 0 and related_psalms_list:
            psalms_list_str = ', '.join(str(p) for p in related_psalms_list)
            related_psalms_str = f"{related_psalms_count} (Psalms {psalms_list_str})"
        elif related_psalms_count > 0:
            related_psalms_str = str(related_psalms_count)
        else:
            related_psalms_str = 'N/A'

        master_editor_stats = stats.get('steps', {}).get('master_editor', {})
        prompt_chars = master_editor_stats.get('input_char_count', 'N/A')
        if isinstance(prompt_chars, int):
            prompt_chars_str = f"{prompt_chars:,} characters"
        else:
            prompt_chars_str = f"{prompt_chars} characters"

        # Deep Web Research status
        deep_research_available = research_data.get('deep_research_available', False)
        deep_research_included = research_data.get('deep_research_included', False)
        deep_research_removed = research_data.get('deep_research_removed_for_space', False)

        if deep_research_included:
            deep_research_str = "Yes"
        elif deep_research_removed:
            deep_research_str = "No (removed for space)"
        elif deep_research_available:
            deep_research_str = "No (available but not included)"
        else:
            deep_research_str = "No"

        # Literary Echoes Research status
        literary_echoes_included = research_data.get('literary_echoes_included', False)
        literary_echoes_available = research_data.get('literary_echoes_available', False)

        if literary_echoes_included:
            literary_echoes_str = "Yes"
        elif literary_echoes_available:
            literary_echoes_str = "No (available but not included)"
        else:
            literary_echoes_str = "No"

        # Sections trimmed for context length
        sections_trimmed = research_data.get('sections_trimmed', [])
        if sections_trimmed:
            sections_trimmed_str = ", ".join(sections_trimmed)
        else:
            sections_trimmed_str = "None"

        # --- Models Used --- (This section will be built from the markdown in the generate() method)
        models_used_str = "### Models Used"

        summary = f"""
Methodological & Bibliographical Summary

### Research & Data Inputs
**Psalm Verses Analyzed**: {verse_count}
**LXX (Septuagint) Verses Reviewed**: {verse_count}
**Phonetic Transcriptions Generated**: {verse_count}
**Lexicon Entries (BDB/Klein) Reviewed**: {lexicon_count}
**Traditional Commentaries Reviewed**: {total_commentaries}{commentary_details}
**Concordance Entries Reviewed**: {concordance_total if concordance_total > 0 else 'N/A'}{concordance_breakdown_str}
**Figurative Concordance Matches Reviewed**: {figurative_total if figurative_total > 0 else 'N/A'}{figurative_breakdown_str}
**Rabbi Jonathan Sacks References Reviewed**: {sacks_count if sacks_count > 0 else 'N/A'}
**Similar Psalms Analyzed**: {related_psalms_str}
**Deep Web Research**: {deep_research_str}
**Literary Echoes Research**: {literary_echoes_str}
**Sections Trimmed for Context**: {sections_trimmed_str}
**Master Editor Prompt Size**: {prompt_chars_str}

{models_used_str}
        """.strip()
        return summary

    def _add_summary_paragraph(self, line: str):
        """Adds a paragraph to the summary, bolding the label.

        When the value contains Hebrew (e.g. the concordance root breakdown), the
        Hebrew list is rendered in a separate native-RTL paragraph so Word pairs each
        root with its count correctly. An LTR paragraph scrambles a long mixed
        Hebrew/number list, which is the source of the unreadable summary line.
        """
        if ':' in line:
            label, value = line.split(':', 1)
            label = label.strip().strip('**')
            m = re.search(r'[֐-׿]', value)
            if m:
                # Hebrew-containing value = the concordance root breakdown. Render it as a
                # native RTL paragraph of "root — count; ..." entries: an em-dash binds each
                # count to its own root (NBSP so an entry never splits across a line break),
                # and "; " separates entries. The parenthesized inline form is unreadable —
                # the count floats ambiguously between two roots and wraps mid-entry.
                pairs = re.findall(r'([א-ת][֐-׿  ־׳״\-]*?)\s*\((\d+)\)', value)
                mnum = re.match(r'\s*\(*\s*([\d,]+)', value)
                total = mnum.group(1) if mnum else ''

                # Label line (LTR): "Label: total"
                p = self.document.add_paragraph(style='SummaryText')
                self._set_paragraph_ltr(p)
                p.add_run(label).bold = True
                p.add_run(f": {total}" if total else ":")
                for run in p.runs:
                    run.font.name = 'Aptos'
                p.paragraph_format.space_after = Pt(0)  # hug the breakdown line below

                # Breakdown line: native RTL paragraph.
                p2 = self.document.add_paragraph(style='SummaryText')
                pPr = p2._element.get_or_add_pPr()
                bidi = OxmlElement('w:bidi')
                bidi.set(ns.qn('w:val'), '1')
                pPr.append(bidi)
                p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                if pairs:
                    NBSP = ' '
                    for i, (root, count) in enumerate(pairs):
                        last = (i == len(pairs) - 1)
                        # The whole entry "root — count;" is ONE RTL run with only
                        # non-breaking spaces inside, so it can never split across a line
                        # (a root run + a separate LTR "— count" run breaks at that
                        # directional boundary even with an NBSP between them). The single
                        # breakable point is the plain space between entries.
                        entry = (root.strip().replace(' ', NBSP) + NBSP + '—' + NBSP
                                 + count + ('' if last else ';'))
                        run = p2.add_run(entry)
                        self._mark_run_hebrew(run)
                        if not last:
                            sep = p2.add_run(' ')  # breakable: wraps happen only between entries
                            sep.font.name = 'Aptos'
                else:
                    # Fallback: couldn't parse pairs — render the raw Hebrew value as-is.
                    self._add_inline_runs(p2, value.strip())
                return

            p = self.document.add_paragraph(style='SummaryText')
            p.add_run(label).bold = True
            p.add_run(f":{value}")
        else:
            p = self.document.add_paragraph(style='SummaryText')
            p.add_run(line)

        for run in p.runs:
            run.font.name = 'Aptos'

    def generate(self):
        """Main method to generate the .docx file, mirroring print_ready.md structure."""
        # Fetch psalm text data from the database first
        db = TanakhDatabase()
        psalm_data = db.get_psalm(self.psalm_num)
        if not psalm_data:
            raise FileNotFoundError(f"Psalm {self.psalm_num} not found in database.")
        psalm_text_data = {v.verse: {'hebrew': v.hebrew, 'english': v.english} for v in psalm_data.verses}

        # 1. Add Title
        self.document.add_heading(f'Commentary on Psalm {self.psalm_num}', level=1)
        # Adjust spacing after the main title for a "softer" break
        title_style = self.document.styles['Heading 1']
        title_style.paragraph_format.space_after = Pt(12)


        # 2. Add Full Psalm Text
        self._format_psalm_text(self.psalm_num, psalm_text_data)

        # Add a page break after the psalm text table
        self.document.add_paragraph() # Add a paragraph to attach the break to
        self.document.add_page_break()

        # 2b. Add Questions for the Reader (if available)
        if self.reader_questions_path and self.reader_questions_path.exists():
            try:
                questions_data = json.loads(self.reader_questions_path.read_text(encoding='utf-8'))
                questions = questions_data.get('curated_questions', [])
                if questions:
                    self.document.add_heading('Questions for the Reader', level=2)
                    
                    # Add introductory italic text
                    intro_p = self.document.add_paragraph(style='BodySans')
                    intro_run = intro_p.add_run('Before reading this commentary, consider the following questions:')
                    intro_run.italic = True
                    
                    # Add each question as a numbered paragraph
                    for i, question in enumerate(questions, 1):
                        q_p = self.document.add_paragraph(style='BodySans')
                        q_p.add_run(f"{i}. ").bold = True
                        self._process_markdown_formatting(q_p, question, set_font=False)
                    
                    # Add spacing after questions
                    self.document.add_paragraph()
            except Exception as e:
                # Log but don't fail if questions can't be loaded
                print(f"Warning: Could not load reader questions: {e}")

        # 3. Add Introduction
        self.document.add_heading('Introduction', level=2)
        intro_content = self.intro_path.read_text(encoding='utf-8')
        # Use _add_commentary_with_bullets to properly handle blockquotes while preserving headers
        self._process_introduction_content(intro_content.strip(), style='BodySans')

        # 4. Add Verse-by-Verse Commentary
        self.document.add_heading('Verse-by-Verse Commentary', level=2)
        verses_content = self.modifier.modify_text(self.verses_path.read_text(encoding='utf-8'))
        parsed_verses = self._parse_verse_commentary(verses_content, psalm_text_data)

        for verse in parsed_verses:
            self.document.add_heading(f'Verse {verse["number"]}', level=3)

            # Commentary now includes the verse text with punctuation from the LLM
            self._add_commentary_with_bullets(verse["commentary"], style='BodySans', is_verse_commentary=True)

        # 5. Add Methodological Summary
        if self.stats_path.exists():
            self.document.add_page_break()
            p = self.document.add_heading('Methodological & Bibliographical Summary', level=2)
            for r in p.runs:
                r.font.name = 'Aptos'
            stats_data = json.loads(self.stats_path.read_text(encoding='utf-8'))
            summary_text = self._format_bibliographical_summary(stats_data)

            # Manually add the Models Used section text to the summary string
            model_usage = stats_data.get('model_usage', {})
            if model_usage:
                summary_text += f"\n**Structural Analysis (Macro)**: {model_usage.get('macro_analysis', 'N/A')}"
                summary_text += f"\n**Verse Discovery (Micro)**: {model_usage.get('micro_analysis', 'N/A')}"
                
                if 'liturgical_librarian' in model_usage:
                    summary_text += f"\n**Liturgical Librarian**: {model_usage.get('liturgical_librarian', 'N/A')}"
                
                if 'figurative_curator' in model_usage:
                    summary_text += f"\n**Figurative Curator**: {model_usage.get('figurative_curator', 'N/A')}"

                if 'insight_extractor' in model_usage:
                    summary_text += f"\n**Insights Extraction**: {model_usage.get('insight_extractor', 'N/A')}"

                if 'question_curator' in model_usage:
                    summary_text += f"\n**Question Generator**: {model_usage.get('question_curator', 'N/A')}"

                if 'synthesis_discovery' in model_usage:
                    summary_text += f"\n**Cross-Verse Synthesis Discovery**: {model_usage.get('synthesis_discovery', 'N/A')}"

                if 'synthesis' in model_usage:
                    summary_text += f"\n**Commentary Synthesis**: {model_usage.get('synthesis', 'N/A')}"
                    summary_text += f"\n**Editorial Review**: {model_usage.get('master_editor', 'N/A')}"
                else:
                    writer_model = model_usage.get('master_writer') or model_usage.get('master_editor', 'N/A')
                    summary_text += f"\n**Commentary (Master Writer)**: {writer_model}"

                if 'citation_filter' in model_usage:
                    summary_text += f"\n**Citation Verifier Filter**: {model_usage.get('citation_filter', 'N/A')}"

                if 'copy_editor' in model_usage:
                    summary_text += f"\n**Copy Editor**: {model_usage.get('copy_editor', 'N/A')}"

                if 'literary_echoes_pass_1' in model_usage:
                    summary_text += f"\n**Literary Echoes (Passes 1 & 2 — Generation)**: {model_usage.get('literary_echoes_pass_1', 'N/A')}"
                if 'literary_echoes_pass_3' in model_usage:
                    summary_text += f"\n**Literary Echoes (Passes 3 & 4 — Verify + Reconstruct)**: {model_usage.get('literary_echoes_pass_3', 'N/A')}"
            else:
                summary_text += "\nModel attribution data not available."

            # Now, add the Date Produced section heading and data
            summary_text += "\n\n### Date Produced"
            master_editor_stats = stats_data.get('steps', {}).get('master_editor', {})
            completion_date_str = master_editor_stats.get('completion_date')
            if completion_date_str:
                from datetime import datetime
                dt = datetime.fromisoformat(completion_date_str.replace('Z', '+00:00'))
                date_str = dt.strftime('%B %d, %Y')
                summary_text += f"\n{date_str}"
            else:
                summary_text += "\nDate not available."

            # Parse the summary text and add it with basic formatting
            for line in summary_text.split('\n'):
                stripped_line = line.strip()
                if not stripped_line:
                    continue
                if stripped_line.startswith('###'):
                    p = self.document.add_heading(line.replace('###', '').strip(), level=3)
                    for r in p.runs:
                        r.font.name = 'Aptos'
                elif stripped_line.startswith('##'):
                     p = self.document.add_heading(line.replace('##', '').strip(), level=2)
                     for r in p.runs:
                         r.font.name = 'Aptos'
                elif "Methodological & Bibliographical Summary" in stripped_line:
                    continue  # Skip the redundant title line
                else:
                    self._add_summary_paragraph(stripped_line.lstrip('- '))

        # 6. Add Page Numbers to Footer
        section = self.document.sections[0]
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        add_page_number(p)

        # 6. Fix complex-script fonts (Arabic/Persian) and Save Document
        self._join_rtl_runs_across_whitespace()
        self._fix_complex_script_fonts()
        self.document.save(self.output_path)
        print(f"Successfully generated Word document: {self.output_path}")


def main():
    """Command-line interface for the document generator."""
    if sys.platform == 'win32':
        sys.stdout.reconfigure(encoding='utf-8')

    parser = argparse.ArgumentParser(description="Generate a .docx commentary from pipeline outputs.")
    parser.add_argument("--psalm", type=int, required=True, help="Psalm number.")
    parser.add_argument("--intro", type=Path, required=True, help="Path to the edited introduction markdown file.")
    parser.add_argument("--verses", type=Path, required=True, help="Path to the edited verses markdown file.")
    parser.add_argument("--stats", type=Path, required=True, help="Path to the pipeline_stats.json file.")
    parser.add_argument("--output", type=Path, required=True, help="Path to save the final .docx file.")
    args = parser.parse_args()

    if not args.intro.exists():
        print(f"Error: Introduction file not found at {args.intro}")
        sys.exit(1)
    if not args.verses.exists():
        print(f"Error: Verses file not found at {args.verses}")
        sys.exit(1)
    if not args.stats.exists():
        print(f"Warning: Stats file not found at {args.stats}. Bibliographical summary will be skipped.")

    generator = DocumentGenerator(
        psalm_num=args.psalm,
        intro_path=args.intro,
        verses_path=args.verses,
        stats_path=args.stats,
        output_path=args.output
    )
    generator.generate()


if __name__ == "__main__":
    main()
