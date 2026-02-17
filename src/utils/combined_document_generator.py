"""
Generates a combined print-ready Microsoft Word (.docx) document that includes
both main and college commentary versions in a single document.

This script creates a document with:
1. Full psalm text (Hebrew and English)
2. Main introduction
3. College introduction (titled "Introduction - College version")
4. Modern Jewish Liturgical Use section (from main version)
5. Verse-by-verse commentary with both main and college versions
   - Each verse shows the main commentary, then a divider line, then the college commentary
"""

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns

# Handle imports for both module and script usage
if __name__ == '__main__' and __package__ is None:
    sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent))
    from src.data_sources.tanakh_database import TanakhDatabase
    from src.utils.divine_names_modifier import DivineNamesModifier
    from src.data_sources.sefaria_client import strip_sefaria_footnotes
else:
    from ..data_sources.tanakh_database import TanakhDatabase
    from .divine_names_modifier import DivineNamesModifier
    from ..data_sources.sefaria_client import strip_sefaria_footnotes


def add_page_number(paragraph):
    """
    Adds a page number field to the given paragraph.
    The paragraph should be in a footer.
    """
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Create the run that will contain the page number field
    run = paragraph.add_run()

    # Create the complex field for the page number
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(ns.qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(ns.qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'

    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(ns.qn('w:fldCharType'), 'end')

    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_end)

class CombinedDocumentGenerator:
    """Generates a combined .docx with main and college commentaries."""

    def __init__(self, psalm_num: int,
                 main_intro_path: Path, main_verses_path: Path,
                 college_intro_path: Path, college_verses_path: Path,
                 stats_path: Path, output_path: Path,
                 reader_questions_path: Path = None,
                 college_questions_path: Path = None):
        self.psalm_num = psalm_num
        self.main_intro_path = main_intro_path
        self.main_verses_path = main_verses_path
        self.college_intro_path = college_intro_path
        self.college_verses_path = college_verses_path
        self.stats_path = stats_path
        self.output_path = output_path
        self.reader_questions_path = reader_questions_path
        self.college_questions_path = college_questions_path
        self.document = Document()
        self._set_default_styles()
        self.modifier = DivineNamesModifier()


    @staticmethod
    def _split_into_grapheme_clusters(text: str) -> List[str]:
        """
        Split Hebrew text into grapheme clusters (base letter + combining marks).

        A grapheme cluster consists of:
        - A base character (Hebrew letter or space)
        - Followed by zero or more combining characters (nikud, cantillation, shin/sin dot, etc.)

        Combining character ranges:
        - U+0591-U+05BD: Cantillation marks
        - U+05BF: Rafe
        - U+05C1-U+05C2: Shin/Sin dots
        - U+05C4-U+05C7: Other marks
        - U+05B0-U+05BD: Vowel points (nikud)
        """
        # Pattern: base character followed by any combining marks
        # Base: Hebrew letter (U+05D0-U+05EA), maqqef (U+05BE), or space
        # Combining: U+0591-U+05BD, U+05BF, U+05C1-U+05C2, U+05C4-U+05C7
        cluster_pattern = r'[\u05D0-\u05EA\u05BE\s][\u0591-\u05BD\u05BF\u05C1-\u05C2\u05C4-\u05C7]*'
        clusters = re.findall(cluster_pattern, text)
        return clusters

    @staticmethod
    def _reverse_hebrew_by_clusters(hebrew_text: str) -> str:
        """
        Reverse Hebrew text by grapheme clusters (keeping letter+nikud together).

        Example:
        - Input: "שִׁלוֹם" (shalom) = [שִׁ, ל, וֹ, ם]
        - Output: "םוֹלשִׁ" (reversed clusters)

        This prevents combining marks (nikud) from detaching from their base letters,
        which would cause dotted circle placeholders to appear.
        """
        clusters = CombinedDocumentGenerator._split_into_grapheme_clusters(hebrew_text)
        # Reverse the order of clusters (but keep each cluster intact)
        reversed_clusters = clusters[::-1]
        return ''.join(reversed_clusters)

    @staticmethod
    def _is_primarily_hebrew(text: str) -> bool:
        """
        Detect if a line is a standalone Hebrew verse line (not mixed Hebrew/English).
        
        A line is a standalone Hebrew verse if:
        - Contains sof-pasuq (׃) - the Hebrew end-of-verse marker
        - Has very few or no ASCII letters (< 5% of total letters)
        
        This is intentionally strict to avoid triggering on English commentary
        paragraphs that contain many Hebrew quotations.
        """
        if not text.strip():
            return False
        
        # Must have sof-pasuq to be considered a complete Hebrew verse line
        has_sof_pasuq = '׃' in text
        if not has_sof_pasuq:
            return False
        
        # Count Hebrew letters (U+05D0-U+05EA)
        hebrew_letters = len(re.findall(r'[\u05D0-\u05EA]', text))
        # Count ASCII letters
        ascii_letters = len(re.findall(r'[a-zA-Z]', text))
        
        total_letters = hebrew_letters + ascii_letters
        if total_letters == 0:
            return False
        
        # Require very low ASCII ratio (< 5%) to be considered primarily Hebrew
        # This excludes mixed paragraphs like "This is the psalm's theological center..."
        ascii_ratio = ascii_letters / total_letters
        
        return ascii_ratio < 0.05

    @staticmethod
    def _is_hebrew_dominant(text: str) -> bool:
        """
        Check if text is predominantly Hebrew (>80% Hebrew letters, min 2).
        Unlike _is_primarily_hebrew, does NOT require sof-pasuq.
        Safe for content already isolated by markdown formatting markers.
        """
        hebrew_letters = len(re.findall(r'[\u05D0-\u05EA]', text))
        ascii_letters = len(re.findall(r'[a-zA-Z]', text))
        total = hebrew_letters + ascii_letters
        if total == 0:
            return False
        return (hebrew_letters / total) > 0.8 and hebrew_letters >= 2

    def _process_text_rtl(self, text):
        """
        Process text for RTL display. Returns modified text string.

        Handles: primarily-Hebrew lines, Hebrew-dominant text,
        parenthesized/bracketed Hebrew, verse refs, trailing punctuation.
        """
        # Full reversal for primarily Hebrew or Hebrew-dominant text
        if self._is_primarily_hebrew(text) or self._is_hebrew_dominant(text):
            return self._reverse_primarily_hebrew_line(text)

        # Standard RTL processing for mixed content
        modified = text
        LRO = '\u202D'  # LEFT-TO-RIGHT OVERRIDE
        PDF = '\u202C'  # POP DIRECTIONAL FORMATTING

        # Hebrew in parentheses
        hebrew_paren_pattern = r'\(([\u0590-\u05FF\u05B0-\u05BD\u05BF\u05C1-\u05C2\u05C4-\u05C7\s]+)\)'
        if re.search(hebrew_paren_pattern, modified):
            def reverse_hebrew_paren(match):
                hebrew_text = match.group(1)
                reversed_hebrew = self._reverse_hebrew_by_clusters(hebrew_text)
                return f'{LRO}({reversed_hebrew}){PDF}'
            modified = re.sub(hebrew_paren_pattern, reverse_hebrew_paren, modified)

        # Hebrew in square brackets
        hebrew_bracket_pattern = r'\[([\u0590-\u05FF\u05B0-\u05BD\u05BF\u05C1-\u05C2\u05C4-\u05C7\s]+)\]'
        if re.search(hebrew_bracket_pattern, modified):
            def reverse_hebrew_bracket(match):
                hebrew_text = match.group(1)
                reversed_hebrew = self._reverse_hebrew_by_clusters(hebrew_text)
                return f'{LRO}[{reversed_hebrew}]{PDF}'
            modified = re.sub(hebrew_bracket_pattern, reverse_hebrew_bracket, modified)

        # Verse references
        verse_ref_pattern = r'(\(\d+:\d+(?:[–\-]\d+)?\))'
        if re.search(verse_ref_pattern, modified):
            def wrap_verse_ref(match):
                return f'{LRO}{match.group(1)}{PDF}'
            modified = re.sub(verse_ref_pattern, wrap_verse_ref, modified)

        # Trailing punctuation RLM anchor
        RLM = '\u200F'  # RIGHT-TO-LEFT MARK
        hebrew_count = len(re.findall(r'[\u05D0-\u05EA]', modified))
        if hebrew_count >= 3 and re.search(r'[.;:,!?]$', modified):
            modified = modified + RLM

        return modified

    @staticmethod
    def _reverse_primarily_hebrew_line(text: str) -> str:
        """
        Reverse an entire primarily-Hebrew line for correct RTL display in LTR paragraphs.
        
        Handles:
        - Hebrew words (reversed by grapheme clusters)
        - Punctuation mirroring (brackets/parentheses swap directions)
        - Spaces and separators
        - Verse references like (26:6) or (26:6–7) are extracted and appended outside the LRO wrapper
        
        The approach: split by spaces, reverse each Hebrew word, reverse the word order,
        mirror bracket characters.
        """
        LRO = '\u202D'  # LEFT-TO-RIGHT OVERRIDE
        PDF = '\u202C'  # POP DIRECTIONAL FORMATTING
        
        # Extract trailing verse reference (if any) to append OUTSIDE the LRO wrapper
        # Pattern matches trailing space + (N:N) or (N:N–N) or (N:N-N) at end of text
        trailing_ref_pattern = r'(\s*\(\d+:\d+(?:[–\-]\d+)?\)\s*)$'
        trailing_ref_match = re.search(trailing_ref_pattern, text)
        trailing_ref = ''
        if trailing_ref_match:
            trailing_ref = trailing_ref_match.group(1)
            text = text[:trailing_ref_match.start()]  # Remove trailing ref from main text
        
        # Mirror map for brackets/parentheses - these need to swap in RTL
        MIRROR_MAP = {
            '(': ')',
            ')': '(',
            '[': ']',
            ']': '[',
            '{': '}',
            '}': '{',
            '<': '>',
            '>': '<',
        }
        
        # Pattern to split on word boundaries while keeping delimiters
        # This captures spaces, semicolons, parentheses, brackets
        # Note: \[ and \] (single backslash) escape brackets inside the character class
        # Using \\[ would create a literal backslash + unescaped bracket, breaking the class
        tokens = re.split(r'(\s+|[;:,.()\[\]׃])', text)
        
        reversed_tokens = []
        for token in tokens:
            if not token:
                continue
            # If token contains Hebrew letters, reverse by grapheme clusters
            if re.search(r'[\u05D0-\u05EA]', token):
                reversed_token = CombinedDocumentGenerator._reverse_hebrew_by_clusters(token)
                reversed_tokens.append(reversed_token)
            elif token in MIRROR_MAP:
                # Mirror brackets/parentheses for RTL display
                reversed_tokens.append(MIRROR_MAP[token])
            else:
                # Keep other punctuation and spaces as-is
                reversed_tokens.append(token)
        
        # Reverse the entire token list to get RTL word order
        reversed_tokens.reverse()
        
        # Join and wrap with LRO to force display in reversed order
        result = ''.join(reversed_tokens)
        
        # Return with LRO wrapper, then append trailing verse reference OUTSIDE the wrapper
        return f'{LRO}{result}{PDF}{trailing_ref}'

    def _set_default_styles(self):
        """Set default font for the document."""
        # Set default font
        font = self.document.styles['Normal'].font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        # Set complex-script fallback font on Normal style so Arabic/Persian text
        # renders correctly (Aptos and Times New Roman lack Arabic glyphs in some installs)
        normal_rPr = self.document.styles['Normal'].element.get_or_add_rPr()
        normal_rFonts = normal_rPr.find(ns.qn('w:rFonts'))
        if normal_rFonts is None:
            normal_rFonts = OxmlElement('w:rFonts')
            normal_rPr.append(normal_rFonts)
        normal_rFonts.set(ns.qn('w:cs'), 'Times New Roman')

        # Create a new style for the sans-serif body text (intro and commentary)
        style = self.document.styles.add_style('BodySans', 1) # 1 for paragraph style
        style.base_style = self.document.styles['Normal']
        style.font.name = 'Aptos'
        style.font.size = Pt(12) # Aptos 12pt for intro and commentary

        # Set complex-script fallback on BodySans too (Aptos is Latin-only)
        body_rPr = style.element.get_or_add_rPr()
        body_rFonts = body_rPr.find(ns.qn('w:rFonts'))
        if body_rFonts is None:
            body_rFonts = OxmlElement('w:rFonts')
            body_rPr.append(body_rFonts)
        body_rFonts.set(ns.qn('w:cs'), 'Times New Roman')

        # Create a new style for the methodological summary
        summary_style = self.document.styles.add_style('SummaryText', 1) # 1 for paragraph style
        summary_style.base_style = self.document.styles['Normal']
        summary_style.font.name = 'Times New Roman'
        summary_style.font.size = Pt(9) # Times New Roman 9pt for summary
        summary_style.paragraph_format.line_spacing = 0.8 # 70% line spacing

        # Set spacing for all heading levels to be tighter
        for i in range(1, 5):
            style = self.document.styles[f'Heading {i}']
            style.paragraph_format.space_before = Pt(12)
            style.paragraph_format.space_after = Pt(4)

        # Set paragraph spacing to be tighter (soft breaks)
        style = self.document.styles['Normal']
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(8)

        # Set document margins to 1 inch on all sides
        for section in self.document.sections:
            section.top_margin = Pt(72)  # 1 inch = 72 points
            section.bottom_margin = Pt(72)
            section.left_margin = Pt(72)
            section.right_margin = Pt(72)

    def _fix_complex_script_fonts(self):
        """Post-processing pass: fix fonts for Arabic/Persian/Urdu text in all runs.
        
        python-docx's font.name setter overrides the CS (complex script) font attribute,
        causing Arabic text to render as boxes when the Latin font (e.g. Aptos) lacks
        Arabic glyphs. This method iterates all runs and explicitly sets the CS font
        to 'Times New Roman' on any run containing Arabic-range characters.
        """
        import re
        arabic_pattern = re.compile(r'[\u0600-\u06FF\u0750-\u077F\uFB50-\uFDFF\uFE70-\uFEFF]')
        fixed_count = 0
        for paragraph in self.document.paragraphs:
            for run in paragraph.runs:
                if run.text and arabic_pattern.search(run.text):
                    rPr = run._element.get_or_add_rPr()
                    rFonts = rPr.find(ns.qn('w:rFonts'))
                    if rFonts is None:
                        rFonts = OxmlElement('w:rFonts')
                        rPr.append(rFonts)
                    rFonts.set(ns.qn('w:cs'), 'Times New Roman')
                    fixed_count += 1
        # Also check table cells
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.text and arabic_pattern.search(run.text):
                                rPr = run._element.get_or_add_rPr()
                                rFonts = rPr.find(ns.qn('w:rFonts'))
                                if rFonts is None:
                                    rFonts = OxmlElement('w:rFonts')
                                    rPr.append(rFonts)
                                rFonts.set(ns.qn('w:cs'), 'Times New Roman')
                                fixed_count += 1
        if fixed_count > 0:
            print(f"  Fixed complex-script fonts on {fixed_count} run(s) containing Arabic text.")

    def _set_paragraph_ltr(self, paragraph):
        """Set paragraph direction to left-to-right."""
        pPr = paragraph._element.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(ns.qn('w:val'), '0')
        pPr.append(bidi)

    def _add_paragraph_with_markdown(self, text: str, style: str = 'Normal'):
        """
        Add a paragraph with basic markdown formatting (bold, italic).
        Handles Hebrew text with RTL formatting.
        """
        paragraph = self.document.add_paragraph(style=style)
        self._set_paragraph_ltr(paragraph)

        # Process with the centralized formatting method
        self._process_markdown_formatting(paragraph, text, set_font=True)

    def _process_markdown_formatting(self, paragraph, text, set_font=False):
        """
        Process markdown formatting (bold, italic, etc.) and add runs to the given paragraph.
        This handles Hebrew in parentheses correctly using LRO/PDF Unicode control characters.

        Args:
            paragraph: The paragraph to add runs to
            text: The text to process
            set_font: If True, explicitly set Aptos 12pt font on all runs
        """
        modified_text = self.modifier.modify_text(text)
        # Match bold (**...**) and italic (*...*) patterns
        # Order matters: match ** before * to avoid incorrect splits
        parts = re.split(r'(\*\*.*?\*\*|__.*?__|\*.*?\*|_.*?_)', modified_text)

        for part in parts:
            if not part:
                continue

            if part.startswith('**') and part.endswith('**'):
                # Bold text - recursively process inner content for nested formatting
                inner_content = part[2:-2]
                self._add_formatted_content(paragraph, inner_content, bold=True, italic=False, set_font=set_font)
            elif part.startswith('__') and part.endswith('__'):
                # Bold text - recursively process inner content for nested formatting
                inner_content = part[2:-2]
                self._add_formatted_content(paragraph, inner_content, bold=True, italic=False, set_font=set_font)
            elif part.startswith('*') and part.endswith('*'):
                processed = self._process_text_rtl(part[1:-1])
                run = paragraph.add_run(processed)
                run.italic = True
                if set_font:
                    run.font.name = 'Aptos'
                    run.font.size = Pt(12)
            elif part.startswith('_') and part.endswith('_'):
                processed = self._process_text_rtl(part[1:-1])
                run = paragraph.add_run(processed)
                run.italic = True
                if set_font:
                    run.font.name = 'Aptos'
                    run.font.size = Pt(12)
            else:
                # Check if this is a primarily-Hebrew line (like a verse text)
                # If so, apply full-line RTL reversal for correct display
                if self._is_primarily_hebrew(part):
                    modified_part = self._reverse_primarily_hebrew_line(part)
                    run = paragraph.add_run(modified_part)
                else:
                    # Handle parenthesized/bracketed Hebrew with grapheme cluster reversal + LRO
                    # This prevents punctuation from being reversed or moved around
                    hebrew_paren_pattern = r'\(([֐-׿ְ-ֽֿׁ-ׂׄ-ׇ\s]+)\)'
                    hebrew_bracket_pattern = r'\[([\u0590-\u05FF\u05B0-\u05BD\u05BF\u05C1-\u05C2\u05C4-\u05C7\s]+)\]'
                    
                    modified_part = part
                    LRO = '\u202D'  # LEFT-TO-RIGHT OVERRIDE
                    PDF = '\u202C'  # POP DIRECTIONAL FORMATTING
                    
                    if re.search(hebrew_paren_pattern, modified_part):
                        def reverse_hebrew_paren(match):
                            hebrew_text = match.group(1)
                            reversed_hebrew = self._reverse_hebrew_by_clusters(hebrew_text)
                            return f'{LRO}({reversed_hebrew}){PDF}'
                        modified_part = re.sub(hebrew_paren_pattern, reverse_hebrew_paren, modified_part)
                    
                    if re.search(hebrew_bracket_pattern, modified_part):
                        def reverse_hebrew_bracket(match):
                            hebrew_text = match.group(1)
                            reversed_hebrew = self._reverse_hebrew_by_clusters(hebrew_text)
                            return f'{LRO}[{reversed_hebrew}]{PDF}'
                        modified_part = re.sub(hebrew_bracket_pattern, reverse_hebrew_bracket, modified_part)
                    
                    # Wrap verse references with LRO/PDF to prevent Word's bidi algorithm from reversing them
                    # Pattern matches (N:N) or (N:N–N) or (N:N-N) 
                    verse_ref_pattern = r'(\(\d+:\d+(?:[–\-]\d+)?\))'
                    if re.search(verse_ref_pattern, modified_part):
                        def wrap_verse_ref(match):
                            return f'{LRO}{match.group(1)}{PDF}'
                        modified_part = re.sub(verse_ref_pattern, wrap_verse_ref, modified_part)
                    
                    # If text contains significant Hebrew and ends with punctuation,
                    # add RLM (Right-to-Left Mark) to anchor punctuation to the RTL context
                    RLM = '\u200F'  # RIGHT-TO-LEFT MARK
                    hebrew_count = len(re.findall(r'[\u05D0-\u05EA]', modified_part))
                    if hebrew_count >= 3 and re.search(r'[.;:,!?]$', modified_part):
                        modified_part = modified_part + RLM
                    
                    run = paragraph.add_run(modified_part)

                if set_font:
                    run.font.name = 'Aptos'
                    run.font.size = Pt(12)

    def _add_formatted_content(self, paragraph, text, bold=False, italic=False, set_font=False):
        """
        Add text content with specified formatting, recursively processing any nested markdown.
        This handles cases like **bold (*italic inside bold*)**.

        Args:
            paragraph: The paragraph to add runs to
            text: The text content to add
            bold: Whether the text should be bold
            italic: Whether the text should be italic
            set_font: If True, explicitly set Aptos 12pt font
        """
        # Check if there's nested formatting to process
        if '*' in text or '_' in text:
            # Split on italic markers to handle nested formatting
            parts = re.split(r'(\*.*?\*|_.*?_)', text)
            for part in parts:
                if not part:
                    continue
                if part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                    # Nested italic - with RTL processing
                    processed = self._process_text_rtl(part[1:-1])
                    run = paragraph.add_run(processed)
                    run.bold = bold
                    run.italic = True  # Nested italic
                    if set_font:
                        run.font.name = 'Aptos'
                        run.font.size = Pt(12)
                elif part.startswith('_') and part.endswith('_') and not part.startswith('__'):
                    # Nested italic - with RTL processing
                    processed = self._process_text_rtl(part[1:-1])
                    run = paragraph.add_run(processed)
                    run.bold = bold
                    run.italic = True  # Nested italic
                    if set_font:
                        run.font.name = 'Aptos'
                        run.font.size = Pt(12)
                else:
                    # Regular text with base formatting
                    # Check if primarily Hebrew and apply full-line reversal
                    if self._is_primarily_hebrew(part):
                        modified_part = self._reverse_primarily_hebrew_line(part)
                        run = paragraph.add_run(modified_part)
                    else:
                        # Handle parenthesized/bracketed Hebrew with grapheme cluster reversal + LRO
                        hebrew_paren_pattern = r'\(([\u0590-\u05FF\u05B0-\u05BD\u05BF\u05C1-\u05C2\u05C4-\u05C7\s]+)\)'
                        hebrew_bracket_pattern = r'\[([\u0590-\u05FF\u05B0-\u05BD\u05BF\u05C1-\u05C2\u05C4-\u05C7\s]+)\]'
                        
                        modified_part = part
                        LRO = '\u202D'  # LEFT-TO-RIGHT OVERRIDE
                        PDF = '\u202C'  # POP DIRECTIONAL FORMATTING
                        
                        if re.search(hebrew_paren_pattern, modified_part):
                            def reverse_hebrew_paren(match):
                                hebrew_text = match.group(1)
                                reversed_hebrew = self._reverse_hebrew_by_clusters(hebrew_text)
                                return f'{LRO}({reversed_hebrew}){PDF}'
                            modified_part = re.sub(hebrew_paren_pattern, reverse_hebrew_paren, modified_part)
                        
                        if re.search(hebrew_bracket_pattern, modified_part):
                            def reverse_hebrew_bracket(match):
                                hebrew_text = match.group(1)
                                reversed_hebrew = self._reverse_hebrew_by_clusters(hebrew_text)
                                return f'{LRO}[{reversed_hebrew}]{PDF}'
                            modified_part = re.sub(hebrew_bracket_pattern, reverse_hebrew_bracket, modified_part)
                        
                        # Wrap verse references with LRO/PDF
                        verse_ref_pattern = r'(\(\d+:\d+(?:[–\-]\d+)?\))'
                        if re.search(verse_ref_pattern, modified_part):
                            def wrap_verse_ref(match):
                                return f'{LRO}{match.group(1)}{PDF}'
                            modified_part = re.sub(verse_ref_pattern, wrap_verse_ref, modified_part)
                        
                        # Add RLM for Hebrew trailing punctuation
                        RLM = '\u200F'
                        hebrew_count = len(re.findall(r'[\u05D0-\u05EA]', modified_part))
                        if hebrew_count >= 3 and re.search(r'[.;:,!?]$', modified_part):
                            modified_part = modified_part + RLM
                        
                        run = paragraph.add_run(modified_part)
                    run.bold = bold
                    run.italic = italic
                    if set_font:
                        run.font.name = 'Aptos'
                        run.font.size = Pt(12)
        else:
            # No nested formatting - just add with base formatting
            # Check if primarily Hebrew and apply full-line reversal
            if self._is_primarily_hebrew(text):
                modified_text = self._reverse_primarily_hebrew_line(text)
                run = paragraph.add_run(modified_text)
            else:
                # Handle parenthesized/bracketed Hebrew with grapheme cluster reversal + LRO
                hebrew_paren_pattern = r'\(([\u0590-\u05FF\u05B0-\u05BD\u05BF\u05C1-\u05C2\u05C4-\u05C7\s]+)\)'
                hebrew_bracket_pattern = r'\[([\u0590-\u05FF\u05B0-\u05BD\u05BF\u05C1-\u05C2\u05C4-\u05C7\s]+)\]'
                
                modified_text = text
                LRO = '\u202D'  # LEFT-TO-RIGHT OVERRIDE
                PDF = '\u202C'  # POP DIRECTIONAL FORMATTING
                
                if re.search(hebrew_paren_pattern, modified_text):
                    def reverse_hebrew_paren(match):
                        hebrew_text = match.group(1)
                        reversed_hebrew = self._reverse_hebrew_by_clusters(hebrew_text)
                        return f'{LRO}({reversed_hebrew}){PDF}'
                    modified_text = re.sub(hebrew_paren_pattern, reverse_hebrew_paren, modified_text)
                
                if re.search(hebrew_bracket_pattern, modified_text):
                    def reverse_hebrew_bracket(match):
                        hebrew_text = match.group(1)
                        reversed_hebrew = self._reverse_hebrew_by_clusters(hebrew_text)
                        return f'{LRO}[{reversed_hebrew}]{PDF}'
                    modified_text = re.sub(hebrew_bracket_pattern, reverse_hebrew_bracket, modified_text)
                
                # Wrap verse references with LRO/PDF
                verse_ref_pattern = r'(\(\d+:\d+(?:[–\-]\d+)?\))'
                if re.search(verse_ref_pattern, modified_text):
                    def wrap_verse_ref(match):
                        return f'{LRO}{match.group(1)}{PDF}'
                    modified_text = re.sub(verse_ref_pattern, wrap_verse_ref, modified_text)
                
                # Add RLM for Hebrew trailing punctuation
                RLM = '\u200F'
                hebrew_count = len(re.findall(r'[\u05D0-\u05EA]', modified_text))
                if hebrew_count >= 3 and re.search(r'[.;:,!?]$', modified_text):
                    modified_text = modified_text + RLM
                
                run = paragraph.add_run(modified_text)
            run.bold = bold
            run.italic = italic
            if set_font:
                run.font.name = 'Aptos'
                run.font.size = Pt(12)

    def _add_commentary_with_bullets(self, text: str, style: str = 'Normal'):
        """
        Add commentary text with proper bullet formatting.
        Handles both regular text and bullet points.
        """
        lines = text.split('\n')

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Skip horizontal rule markers (---, ___, ***)
            if line in ('---', '___', '***', '—', '–'):
                continue

            # Check if this is a markdown heading (###, ##, ####)
            if line.startswith('#'):
                # Parse heading level and text
                heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
                if heading_match:
                    hashes = heading_match.group(1)
                    heading_text = heading_match.group(2)
                    # Map markdown heading levels to Word heading levels
                    level = len(hashes)
                    word_heading_level = min(level, 4)  # Cap at Heading 4

                    # Add as a Word heading
                    p = self.document.add_heading(heading_text, level=word_heading_level)
                    self._set_paragraph_ltr(p)
                else:
                    # Malformed heading, treat as normal text
                    self._add_paragraph_with_markdown(line, style=style)
            # Check if this is a bullet point
            elif line.startswith('- '):
                # Create a bulleted paragraph
                paragraph = self.document.add_paragraph(style='List Bullet')
                self._set_paragraph_ltr(paragraph)
                line_without_bullet = line[2:]  # Remove "- "
                # Process with proper formatting
                self._process_markdown_formatting(paragraph, line_without_bullet, set_font=True)
            # Check if this is a block quote
            elif line.startswith('>'):
                # Check if it's an empty quote line (just ">" with optional spaces)
                quote_text = line[1:].lstrip() if len(line) > 1 else ''  # Remove ">" and any leading spaces
                if not quote_text:
                    # Empty quote line - add a blank paragraph
                    self.document.add_paragraph()
                else:
                    # Non-empty quote line - create a block quote paragraph with indentation and italic
                    paragraph = self.document.add_paragraph(style=style)
                    self._set_paragraph_ltr(paragraph)
                    paragraph.paragraph_format.left_indent = Inches(0.5)
                    # Process with proper formatting
                    self._process_markdown_formatting(paragraph, quote_text, set_font=True)
                    # Make the entire quote italic
                    for run in paragraph.runs:
                        run.italic = True
            else:
                # Regular paragraph
                self._add_paragraph_with_markdown(line, style=style)

    def _add_summary_paragraph(self, text: str):
        """Add a summary paragraph with the SummaryText style."""
        self._add_paragraph_with_markdown(text, style='SummaryText')

    def _parse_verse_commentary(self, content: str, psalm_text_data: Dict[int, Dict[str, str]]) -> List[Dict[str, str]]:
        """Parses the verse-by-verse commentary file.

        Handles both single verses (Verse 1) and verse ranges (Verses 21-25).
        For ranges, returns a single entry with 'number' as the range string (e.g., '21-25')
        and 'verse_range' as a tuple (start, end).
        """
        verses = []
        # Try multiple formats: "**Verse(s) X(-Y)**" (main format), "### Verse(s) X(-Y)" (college format)
        # Match both singular "Verse" and plural "Verses" with optional range
        # First try the expected bold format
        verse_blocks = re.split(r'(?=^\*\*Verses? [\d\-–]+.*?\*\*\s*\n)', content, flags=re.MULTILINE)

        # If no verses found with bold format, try heading format (college uses this)
        if len(verse_blocks) <= 1:
            verse_blocks = re.split(r'(?=^#{2,} Verses? [\d\-–]+)', content, flags=re.MULTILINE)

        # If still no verses found, try plain format without any markdown
        if len(verse_blocks) <= 1:
            verse_blocks = re.split(r'(?=^Verses? [\d\-–]+)', content, flags=re.MULTILINE)

        for block in verse_blocks:
            block = block.strip()
            if not block.strip():
                continue

            lines = block.strip().split('\n', 1)  # Split only on the first newline

            # Try all formats for verse number matching (single or range)
            # Matches: **Verse 1**, **Verses 21-25**, **Verses 21–25 (description)**, etc.
            verse_num_match = re.match(r'^\*\*Verses? ([\d]+)(?:[\-–]([\d]+))?\s*.*?\*\*', lines[0])  # Bold format
            if not verse_num_match:
                # Heading format: ### Verse 1, ### Verses 21-25 (description), etc.
                verse_num_match = re.match(r'^#{2,} Verses? ([\d]+)(?:[\-–]([\d]+))?\s*', lines[0])
            if not verse_num_match:
                # Plain format: Verse 1, Verses 21-25, etc.
                verse_num_match = re.match(r'^Verses? ([\d]+)(?:[\-–]([\d]+))?\s*', lines[0])

            if not verse_num_match:
                continue

            start_verse = verse_num_match.group(1)
            end_verse = verse_num_match.group(2)  # None if single verse

            if end_verse:
                # It's a range
                verse_num = f"{start_verse}-{end_verse}"
                verse_range = (int(start_verse), int(end_verse))
            else:
                # Single verse
                verse_num = start_verse
                verse_range = None

            commentary_text = '\n'.join(lines[1:]).strip()

            # Get Hebrew/English text from the database data
            verse_info = psalm_text_data.get(int(start_verse), {})
            hebrew_text = verse_info.get('hebrew', '[Hebrew text not found]')
            english_text = verse_info.get('english', '[English text not found]')

            verse_entry = {
                "number": verse_num,
                "hebrew": hebrew_text,
                "english": english_text,
                "commentary": commentary_text
            }
            if verse_range:
                verse_entry["verse_range"] = verse_range

            verses.append(verse_entry)
        return verses

    def _format_psalm_text(self, psalm_num: int, psalm_text_data: Dict[int, Dict[str, str]]):
        """Formats the full psalm text with Hebrew and English side-by-side."""
        self.document.add_heading(f"Psalm {psalm_num}", level=2)

        # Create a table with 2 columns
        table = self.document.add_table(rows=0, cols=2)
        table.autofit = True

        for verse_num in sorted(psalm_text_data.keys()):
            hebrew = psalm_text_data[verse_num].get('hebrew', '')
            english = psalm_text_data[verse_num].get('english', '')
            # Strip any footnote markers from the English text
            english = strip_sefaria_footnotes(english)
            modified_hebrew = self.modifier.modify_text(hebrew)

            row_cells = table.add_row().cells

            # Hebrew cell (left, but text is RTL)
            p_heb = row_cells[0].paragraphs[0]
            p_heb.add_run(f"{verse_num}. ").bold = True
            p_heb.add_run(modified_hebrew).font.rtl = True

            # English cell (right)
            p_eng = row_cells[1].paragraphs[0]
            run_eng = p_eng.add_run(english)
            run_eng.font.name = 'Cambria'

    def _extract_liturgical_section(self, intro_content: str) -> str:
        """
        Extracts the 'Modern Jewish Liturgical Use' section from the intro content.
        Returns the section content or empty string if not found.

        Handles both:
        - Header-based: ## Modern Jewish Liturgical Use
        - Marker-based: ---LITURGICAL-SECTION-START---
        """
        # First try the header-based format
        match = re.search(r'## Modern Jewish Liturgical Use\s*\n(.*)', intro_content, re.DOTALL | re.IGNORECASE)
        if match:
            return "## Modern Jewish Liturgical Use\n" + match.group(1).strip()

        # Try the marker-based format (used by V2 prompts)
        match = re.search(r'---LITURGICAL-SECTION-START---\s*\n(.*)', intro_content, re.DOTALL)
        if match:
            return "## Modern Jewish Liturgical Use\n" + match.group(1).strip()

        return ""

    def _format_bibliographical_summary(self, stats: Dict[str, Any]) -> str:
        """Formats the stats dictionary into a readable string for the document."""
        # --- Analysis & Research Inputs ---
        analysis_data = stats.get('analysis', {}) or {}
        verse_count = analysis_data.get('verse_count', 'N/A')

        research_data = stats.get('research', {})
        ugaritic_count = len(research_data.get('ugaritic_parallels', []))
        lexicon_count = research_data.get('lexicon_entries_count', 'N/A')

        commentaries = research_data.get('commentary_counts', {})
        total_commentaries = sum(commentaries.values()) if commentaries else 'N/A'
        commentary_details = ""
        if commentaries:
            commentary_lines = [f"{c} ({n})" for c, n in sorted(commentaries.items())]
            commentary_details = f" ({'; '.join(commentary_lines)})"

        concordance_total = sum((research_data.get('concordance_results', {}) or {}).values())

        figurative_results = research_data.get('figurative_results', {}) or {}
        figurative_total = figurative_results.get('total_instances_used', 0) if isinstance(figurative_results, dict) else 0

        # Figurative parallels breakdown
        figurative_parallels = research_data.get('figurative_parallels_reviewed', {})
        figurative_breakdown_str = ""
        if figurative_parallels:
            items = [f"{v} ({c})" for v, c in sorted(figurative_parallels.items())]
            figurative_breakdown_str = f" ({'; '.join(items)})"

        sacks_count = research_data.get('sacks_references_count', 0)
        related_psalms_count = research_data.get('related_psalms_count', 0)
        related_psalms_list = research_data.get('related_psalms_list', [])

        # Format related psalms display: count and list
        if related_psalms_count > 0 and related_psalms_list:
            psalms_list_str = ', '.join(str(p) for p in related_psalms_list)
            related_psalms_str = f"{related_psalms_count} (Psalms {psalms_list_str})"
        elif related_psalms_count > 0:
            related_psalms_str = str(related_psalms_count)
        else:
            related_psalms_str = 'N/A'

        master_editor_stats = stats.get('steps', {}).get('master_editor', {})
        prompt_chars = master_editor_stats.get('input_char_count', 'N/A')
        if isinstance(prompt_chars, int):
            prompt_chars_str = f"{prompt_chars:,} characters"
        else:
            prompt_chars_str = f"{prompt_chars} characters"

        # Deep Web Research status
        deep_research_available = research_data.get('deep_research_available', False)
        deep_research_included = research_data.get('deep_research_included', False)
        deep_research_removed = research_data.get('deep_research_removed_for_space', False)

        if deep_research_included:
            deep_research_str = "Yes"
        elif deep_research_removed:
            deep_research_str = "No (removed for space)"
        elif deep_research_available:
            deep_research_str = "No (available but not included)"
        else:
            deep_research_str = "No"

        # Literary Echoes Research status
        literary_echoes_included = research_data.get('literary_echoes_included', False)
        literary_echoes_available = research_data.get('literary_echoes_available', False)

        if literary_echoes_included:
            literary_echoes_str = "Yes"
        elif literary_echoes_available:
            literary_echoes_str = "No (available but not included)"
        else:
            literary_echoes_str = "No"

        # --- Models Used --- (This section will be built from the markdown in the generate() method)
        models_used_str = "### Models Used"

        summary = f"""
Methodological & Bibliographical Summary

### Research & Data Inputs
**Psalm Verses Analyzed**: {verse_count}
**LXX (Septuagint) Verses Reviewed**: {verse_count}
**Phonetic Transcriptions Generated**: {verse_count}
**Ugaritic Parallels Reviewed**: {ugaritic_count}
**Lexicon Entries (BDB\\Klein) Reviewed**: {lexicon_count}
**Traditional Commentaries Reviewed**: {total_commentaries}{commentary_details}
**Concordance Entries Reviewed**: {concordance_total if concordance_total > 0 else 'N/A'}
**Figurative Concordance Matches Reviewed**: {figurative_total if figurative_total > 0 else 'N/A'}{figurative_breakdown_str}
**Rabbi Jonathan Sacks References Reviewed**: {sacks_count if sacks_count > 0 else 'N/A'}
**Similar Psalms Analyzed**: {related_psalms_str}
**Deep Web Research**: {deep_research_str}
**Literary Echoes Research**: {literary_echoes_str}
**Master Editor Prompt Size**: {prompt_chars_str}

{models_used_str}
        """.strip()
        return summary

    def _add_em_dash_separator(self):
        """
        Adds an em dash separator between main and college commentary.
        """
        paragraph = self.document.add_paragraph(style='BodySans')
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run('—')  # Em dash
        run.font.name = 'Aptos'
        run.font.size = Pt(12)

    def _add_horizontal_border(self):
        """
        Adds a horizontal border line between verses.
        """
        paragraph = self.document.add_paragraph()
        pPr = paragraph._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')

        # Add bottom border
        bottom = OxmlElement('w:bottom')
        bottom.set(ns.qn('w:val'), 'single')
        bottom.set(ns.qn('w:sz'), '12')  # Size in 1/8 pt
        bottom.set(ns.qn('w:space'), '1')
        bottom.set(ns.qn('w:color'), '000000')

        pBdr.append(bottom)
        pPr.append(pBdr)

    def generate(self):
        """Main method to generate the combined .docx file."""
        # Fetch psalm text data from the database first
        db = TanakhDatabase()
        psalm_data = db.get_psalm(self.psalm_num)
        if not psalm_data:
            raise FileNotFoundError(f"Psalm {self.psalm_num} not found in database.")
        psalm_text_data = {v.verse: {'hebrew': v.hebrew, 'english': v.english} for v in psalm_data.verses}

        # 1. Add Title
        self.document.add_heading(f'Commentary on Psalm {self.psalm_num}', level=1)
        # Adjust spacing after the main title for a "softer" break
        title_style = self.document.styles['Heading 1']
        title_style.paragraph_format.space_after = Pt(12)

        # 2. Add Full Psalm Text
        self._format_psalm_text(self.psalm_num, psalm_text_data)

        # Add a page break after the psalm text table
        self.document.add_paragraph() # Add a paragraph to attach the break to
        self.document.add_page_break()

        # 2b. Add Questions for the Reader (if available)
        if self.reader_questions_path and self.reader_questions_path.exists():
            try:
                questions_data = json.loads(self.reader_questions_path.read_text(encoding='utf-8'))
                questions = questions_data.get('curated_questions', [])
                if questions:
                    self.document.add_heading('Questions for the Reader', level=2)
                    
                    # Add introductory italic text
                    intro_p = self.document.add_paragraph(style='BodySans')
                    intro_run = intro_p.add_run('Before reading this commentary, consider the following questions:')
                    intro_run.italic = True
                    
                    # Add each question as a numbered paragraph
                    for i, question in enumerate(questions, 1):
                        q_p = self.document.add_paragraph(style='BodySans')
                        q_p.add_run(f"{i}. ").bold = True
                        self._process_markdown_formatting(q_p, question, set_font=False)
                    
                    # Add spacing after questions
                    self.document.add_paragraph()
            except Exception as e:
                # Log but don't fail if questions can't be loaded
                print(f"Warning: Could not load reader questions: {e}")

        # 3. Add Main Introduction
        self.document.add_heading('Introduction', level=2)
        main_intro_content = self.main_intro_path.read_text(encoding='utf-8')

        # Extract liturgical section from main intro before processing
        liturgical_section = self._extract_liturgical_section(main_intro_content)

        # Remove liturgical section from main intro
        # Handle both header-based (## Modern Jewish Liturgical Use) and marker-based (---LITURGICAL-SECTION-START---)
        main_intro_without_liturgical = re.sub(
            r'(## Modern Jewish Liturgical Use.*|---LITURGICAL-SECTION-START---.*)',
            '',
            main_intro_content,
            flags=re.DOTALL | re.IGNORECASE
        ).strip()

        # Add main intro paragraphs
        for para in main_intro_without_liturgical.strip().split('\n'):
            if para.strip():
                # Handle subheadings
                if para.strip().startswith('####'):
                    self.document.add_heading(para.strip().replace('####', '').strip(), level=4)
                elif para.strip().startswith('###'):
                    self.document.add_heading(para.strip().replace('###', '').strip(), level=3)
                # Handle bullet lists
                elif para.strip().startswith('- '):
                    paragraph = self.document.add_paragraph(style='List Bullet')
                    self._set_paragraph_ltr(paragraph)
                    self._process_markdown_formatting(paragraph, para.strip()[2:], set_font=True)
                # Handle block quotes
                elif para.strip().startswith('>'):
                    # Check if it's an empty quote line (just ">" with optional spaces)
                    quote_text = para.strip()[1:].lstrip()  # Remove ">" and any leading spaces after it
                    if not quote_text:
                        # Empty quote line - add a blank paragraph
                        self.document.add_paragraph()
                    else:
                        # Non-empty quote line
                        paragraph = self.document.add_paragraph(style='BodySans')
                        self._set_paragraph_ltr(paragraph)
                        # Add left indentation and italic formatting for quotes
                        paragraph.paragraph_format.left_indent = Inches(0.5)
                        self._process_markdown_formatting(paragraph, quote_text, set_font=True)
                        # Make the entire quote italic
                        for run in paragraph.runs:
                            run.italic = True
                else:
                    self._add_paragraph_with_markdown(para, style='BodySans')

        # 4. Add College Questions (if available) - before college intro
        if self.college_questions_path and self.college_questions_path.exists():
            try:
                questions_data = json.loads(self.college_questions_path.read_text(encoding='utf-8'))
                questions = questions_data.get('curated_questions', [])
                if questions:
                    # Add college questions heading with green "College" label
                    q_heading = self.document.add_heading('', level=2)
                    q_heading.add_run('Questions for the Reader - ')
                    college_label = q_heading.add_run('College')
                    college_label.font.color.rgb = RGBColor(0, 128, 0)  # Green color
                    q_heading.add_run(' version')
                    
                    # Add introductory italic text
                    intro_p = self.document.add_paragraph(style='BodySans')
                    intro_run = intro_p.add_run('Before reading this commentary, consider the following questions:')
                    intro_run.italic = True
                    
                    # Add each question as a numbered paragraph
                    for i, question in enumerate(questions, 1):
                        q_p = self.document.add_paragraph(style='BodySans')
                        q_p.add_run(f"{i}. ").bold = True
                        self._process_markdown_formatting(q_p, question, set_font=False)
                    
                    # Add spacing after questions
                    self.document.add_paragraph()
            except Exception as e:
                # Log but don't fail if questions can't be loaded
                print(f"Warning: Could not load college reader questions: {e}")

        # 5. Add College Introduction with green "College" in heading
        college_heading = self.document.add_heading('', level=2)
        college_heading.add_run('Introduction - ')
        college_run = college_heading.add_run('College')
        college_run.font.color.rgb = RGBColor(0, 128, 0)  # Green color
        college_heading.add_run(' version')

        college_intro_content = self.college_intro_path.read_text(encoding='utf-8')

        # Remove liturgical section from college intro if present
        # Handle both header-based (## Modern Jewish Liturgical Use) and marker-based (---LITURGICAL-SECTION-START---)
        college_intro_without_liturgical = re.sub(
            r'(## Modern Jewish Liturgical Use.*|---LITURGICAL-SECTION-START---.*)',
            '',
            college_intro_content,
            flags=re.DOTALL | re.IGNORECASE
        ).strip()
        # Also remove the horizontal rule before liturgical section if present
        college_intro_without_liturgical = re.sub(r'\n---\n*$', '', college_intro_without_liturgical).strip()

        # Add college intro paragraphs
        for para in college_intro_without_liturgical.strip().split('\n'):
            if para.strip():
                # Handle subheadings
                if para.strip().startswith('####'):
                    self.document.add_heading(para.strip().replace('####', '').strip(), level=4)
                elif para.strip().startswith('###'):
                    self.document.add_heading(para.strip().replace('###', '').strip(), level=3)
                # Handle bullet lists
                elif para.strip().startswith('- '):
                    paragraph = self.document.add_paragraph(style='List Bullet')
                    self._set_paragraph_ltr(paragraph)
                    self._process_markdown_formatting(paragraph, para.strip()[2:], set_font=True)
                # Handle block quotes
                elif para.strip().startswith('>'):
                    # Check if it's an empty quote line (just ">" with optional spaces)
                    quote_text = para.strip()[1:].lstrip()  # Remove ">" and any leading spaces after it
                    if not quote_text:
                        # Empty quote line - add a blank paragraph
                        self.document.add_paragraph()
                    else:
                        # Non-empty quote line
                        paragraph = self.document.add_paragraph(style='BodySans')
                        self._set_paragraph_ltr(paragraph)
                        # Add left indentation and italic formatting for quotes
                        paragraph.paragraph_format.left_indent = Inches(0.5)
                        self._process_markdown_formatting(paragraph, quote_text, set_font=True)
                        # Make the entire quote italic
                        for run in paragraph.runs:
                            run.italic = True
                else:
                    self._add_paragraph_with_markdown(para, style='BodySans')

        # 5. Add Modern Jewish Liturgical Use section (from main version)
        if liturgical_section:
            self.document.add_heading('Modern Jewish Liturgical Use', level=2)
            # Remove ALL occurrences of the heading from the section content (handles duplicates)
            liturgical_content = re.sub(r'^## Modern Jewish Liturgical Use\s*\n', '', liturgical_section, flags=re.MULTILINE).strip()
            liturgical_content = re.sub(r'## Modern Jewish Liturgical Use\s*', '', liturgical_content).strip()
            for para in liturgical_content.split('\n'):
                if para.strip():
                    # Handle subheadings
                    if para.strip().startswith('####'):
                        self.document.add_heading(para.strip().replace('####', '').strip(), level=4)
                    elif para.strip().startswith('###'):
                        self.document.add_heading(para.strip().replace('###', '').strip(), level=3)
                    # Handle bullet lists
                    elif para.strip().startswith('- '):
                        paragraph = self.document.add_paragraph(style='List Bullet')
                        self._set_paragraph_ltr(paragraph)
                        self._process_markdown_formatting(paragraph, para.strip()[2:], set_font=True)
                    # Handle block quotes
                    elif para.strip().startswith('>'):
                        # Check if it's an empty quote line (just ">" with optional spaces)
                        quote_text = para.strip()[1:].lstrip()  # Remove ">" and any leading spaces after it
                        if not quote_text:
                            # Empty quote line - add a blank paragraph
                            self.document.add_paragraph()
                        else:
                            # Non-empty quote line
                            paragraph = self.document.add_paragraph(style='BodySans')
                            self._set_paragraph_ltr(paragraph)
                            # Add left indentation and italic formatting for quotes
                            paragraph.paragraph_format.left_indent = Inches(0.5)
                            self._process_markdown_formatting(paragraph, quote_text, set_font=True)
                            # Make the entire quote italic
                            for run in paragraph.runs:
                                run.italic = True
                    else:
                        self._add_paragraph_with_markdown(para, style='BodySans')

        # 6. Add Verse-by-Verse Commentary with both versions
        verse_heading = self.document.add_heading('', level=2)
        verse_heading.add_run('Verse-by-verse commentary: Main and ')
        college_word = verse_heading.add_run('College')
        college_word.font.color.rgb = RGBColor(0, 128, 0)  # Green color
        verse_heading.add_run(' versions')

        # Parse both verse files
        main_verses_content = self.modifier.modify_text(self.main_verses_path.read_text(encoding='utf-8'))
        college_verses_content = self.modifier.modify_text(self.college_verses_path.read_text(encoding='utf-8'))

        main_parsed_verses = self._parse_verse_commentary(main_verses_content, psalm_text_data)
        college_parsed_verses = self._parse_verse_commentary(college_verses_content, psalm_text_data)

        # Create mappings for easy lookup
        # For college, we need to handle both exact matches and range matches
        college_verse_map = {}  # Direct number -> verse mapping
        college_range_map = {}  # For ranges: stores {end_verse: verse_entry}
        for v in college_parsed_verses:
            college_verse_map[v['number']] = v
            if 'verse_range' in v:
                # Store by end verse for range matching
                start, end = v['verse_range']
                college_range_map[end] = v

        # Track which college ranges have been displayed
        displayed_college_ranges = set()

        for i, main_verse in enumerate(main_parsed_verses):
            verse_num = main_verse['number']
            verse_num_int = int(verse_num) if verse_num.isdigit() else int(verse_num.split('-')[0])

            # Add verse heading
            self.document.add_heading(f'Verse {verse_num}', level=3)

            # Add main commentary
            self._add_commentary_with_bullets(main_verse["commentary"], style='BodySans')

            # Check for college commentary
            # First try exact match
            college_verse = None
            college_key = None

            if verse_num in college_verse_map:
                college_verse = college_verse_map[verse_num]
                college_key = verse_num
            else:
                # Check if this main verse is the END of a college range
                # This ensures the college block appears after the last verse in the range
                if verse_num_int in college_range_map:
                    college_verse = college_range_map[verse_num_int]
                    college_key = college_verse['number']

            # Only display if we found a match and haven't displayed this range yet
            if college_verse and college_key not in displayed_college_ranges:
                # Add em dash separator between main and college
                self._add_em_dash_separator()
                displayed_college_ranges.add(college_key)
                college_commentary = college_verse["commentary"]

                # Remove any leading Hebrew verse lines and English translation
                # College commentary often starts with the Hebrew verse followed by English translation
                # We want to skip both and start with the actual commentary
                lines = college_commentary.split('\n')
                commentary_lines = []
                in_verse_section = True  # We're in the verse section (Hebrew + English translation)
                in_english_quote = False  # Track if we're inside the English translation quote block

                for line in lines:
                    line_stripped = line.strip()
                    if not line_stripped:
                        if not in_verse_section:
                            commentary_lines.append(line)
                        continue

                    if in_verse_section:
                        # Check for Hebrew line (with or without ** markdown or other formatting)
                        # Remove markdown formatting before checking for Hebrew
                        cleaned_line = re.sub(r'^\*+|\*+$', '', line_stripped)  # Remove leading/trailing asterisks
                        if re.search(r'[\u0590-\u05FF]', cleaned_line[:10]):  # Check first 10 chars for Hebrew
                            continue  # Skip Hebrew verse lines

                        # Check for English translation quote block
                        if line_stripped.startswith('"') or line_stripped.startswith("'"):
                            in_english_quote = True
                            continue

                        # If we're inside the English quote block, check for closing quote
                        if in_english_quote:
                            if line_stripped.endswith('"') or line_stripped.endswith("'"):
                                in_english_quote = False
                            continue

                        # If we're here, we've exited the verse section
                        in_verse_section = False

                    # Add the line to commentary
                    commentary_lines.append(line)

                # Join the English commentary lines
                english_commentary = '\n'.join(commentary_lines).strip()

                if english_commentary:
                    # Now add the commentary with green/bold first word
                    lines = english_commentary.split('\n')
                    first_word_done = False

                    for line in lines:
                        line = line.strip()
                        if not line:
                            continue

                        # Skip horizontal rule markers (---, ___, ***)
                        if line in ('---', '___', '***', '—', '–'):
                            continue

                        # Process first non-empty line specially
                        if not first_word_done:
                            # Handle bullet, block quote, or regular line
                            is_bullet = line.startswith('- ')
                            is_quote = line.startswith('>')

                            if is_bullet:
                                line = line[2:]  # Remove bullet marker
                            elif is_quote:
                                line = line[1:].lstrip()  # Remove ">" and any leading spaces
                                if not line:
                                    # Empty quote line - skip and mark as done
                                    self.document.add_paragraph()
                                    first_word_done = True
                                    continue

                            # Split on first whitespace to get first word
                            words = line.split(None, 1)
                            if words:
                                first_word = words[0]
                                rest_of_line = words[1] if len(words) > 1 else ""

                                # Strip markdown markers from first word (**, *, __, _)
                                first_word_clean = re.sub(r'^\*\*|\*\*$|^\*|\*$|^__$|__$|^_$|_$', '', first_word)

                                # Create paragraph
                                if is_bullet:
                                    paragraph = self.document.add_paragraph(style='List Bullet')
                                else:
                                    paragraph = self.document.add_paragraph(style='BodySans')
                                    if is_quote:
                                        paragraph.paragraph_format.left_indent = Inches(0.5)

                                self._set_paragraph_ltr(paragraph)

                                # Add first word with green and bold (using cleaned word)
                                first_run = paragraph.add_run(first_word_clean + (' ' if rest_of_line else ''))
                                first_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
                                first_run.bold = True
                                first_run.font.name = 'Aptos'
                                first_run.font.size = Pt(12)

                                # Add rest of line with normal formatting (processes markdown)
                                if rest_of_line:
                                    self._process_markdown_formatting(paragraph, rest_of_line, set_font=True)

                                # Make entire quote italic if it's a block quote
                                if is_quote:
                                    for run in paragraph.runs:
                                        run.italic = True

                                first_word_done = True
                            else:
                                # Shouldn't happen, but add normally if it does
                                self._add_paragraph_with_markdown(line, style='BodySans')
                                first_word_done = True
                        else:
                            # Not first line, add normally
                            if line.startswith('- '):
                                paragraph = self.document.add_paragraph(style='List Bullet')
                                self._set_paragraph_ltr(paragraph)
                                self._process_markdown_formatting(paragraph, line[2:], set_font=True)
                            elif line.startswith('>'):
                                # Handle block quote
                                quote_text = line[1:].lstrip()
                                if not quote_text:
                                    # Empty quote line
                                    self.document.add_paragraph()
                                else:
                                    paragraph = self.document.add_paragraph(style='BodySans')
                                    paragraph.paragraph_format.left_indent = Inches(0.5)
                                    self._set_paragraph_ltr(paragraph)
                                    self._process_markdown_formatting(paragraph, quote_text, set_font=True)
                                    # Make the entire quote italic
                                    for run in paragraph.runs:
                                        run.italic = True
                            else:
                                self._add_paragraph_with_markdown(line, style='BodySans')

            # Add horizontal border between verses (except after the last verse)
            if i < len(main_parsed_verses) - 1:
                self._add_horizontal_border()

        # 7. Add Methodological Summary
        if self.stats_path.exists():
            self.document.add_page_break()
            self.document.add_heading('Methodological & Bibliographical Summary', level=2)
            stats_data = json.loads(self.stats_path.read_text(encoding='utf-8'))
            summary_text = self._format_bibliographical_summary(stats_data)

            # Manually add the Models Used section text to the summary string
            model_usage = stats_data.get('model_usage', {})
            if model_usage:
                summary_text += f"\n**Structural Analysis (Macro)**: {model_usage.get('macro_analysis', 'N/A')}"
                summary_text += f"\n**Verse Discovery (Micro)**: {model_usage.get('micro_analysis', 'N/A')}"
                
                if 'liturgical_librarian' in model_usage:
                    summary_text += f"\n**Liturgical Librarian**: {model_usage.get('liturgical_librarian', 'N/A')}"
                
                if 'figurative_curator' in model_usage:
                    summary_text += f"\n**Figurative Curator**: {model_usage.get('figurative_curator', 'N/A')}"
                
                if 'question_curator' in model_usage:
                    summary_text += f"\n**Question Generation**: {model_usage.get('question_curator', 'N/A')}"
                
                if 'insight_extractor' in model_usage:
                    summary_text += f"\n**Insights Extraction**: {model_usage.get('insight_extractor', 'N/A')}"

                if 'master_writer' in model_usage:
                    # Single-pass pipeline (no synthesis writer)
                    summary_text += f"\n**Commentary (Master Writer)**: {model_usage.get('master_writer', 'N/A')}"
                else:
                    # Standard two-pass pipeline (synthesis + editor)
                    summary_text += f"\n**Commentary Synthesis**: {model_usage.get('synthesis', 'N/A')}"
                    summary_text += f"\n**Editorial Review**: {model_usage.get('master_editor', 'N/A')}"
            else:
                summary_text += "\nModel attribution data not available."

            # Now, add the Date Produced section heading and data
            summary_text += "\n\n### Date Produced"
            master_editor_stats = stats_data.get('steps', {}).get('master_editor', {})
            completion_date_str = master_editor_stats.get('completion_date')
            if completion_date_str:
                from datetime import datetime
                dt = datetime.fromisoformat(completion_date_str.replace('Z', '+00:00'))
                date_str = dt.strftime('%B %d, %Y')
                summary_text += f"\n{date_str}"
            else:
                summary_text += "\nDate not available."

            # Parse the summary text and add it with basic formatting
            for line in summary_text.split('\n'):
                stripped_line = line.strip()
                if not stripped_line:
                    continue
                if stripped_line.startswith('###'):
                    self.document.add_heading(line.replace('###', '').strip(), level=3)
                elif stripped_line.startswith('##'):
                     self.document.add_heading(line.replace('##', '').strip(), level=2)
                elif "Methodological & Bibliographical Summary" in stripped_line:
                    continue  # Skip the redundant title line
                elif stripped_line.startswith('- '): # Old format
                    self._add_summary_paragraph(stripped_line.lstrip('- '))
                elif stripped_line:
                    self._add_summary_paragraph(stripped_line)

        # 8. Add Page Numbers to Footer
        section = self.document.sections[0]
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        add_page_number(p)

        # 9. Fix complex-script fonts (Arabic/Persian) and Save Document
        self._fix_complex_script_fonts()
        self.document.save(self.output_path)
        print(f"Successfully generated combined Word document: {self.output_path}")


def main():
    """Command-line interface for the combined document generator."""
    parser = argparse.ArgumentParser(
        description='Generate a combined .docx file with main and college commentaries.'
    )
    parser.add_argument('psalm_number', type=int, help='Psalm number to generate')
    parser.add_argument(
        '--output-dir',
        type=Path,
        default=Path('output'),
        help='Output directory (default: output/)'
    )

    args = parser.parse_args()
    psalm_num = args.psalm_number
    psalm_num_str = f"{psalm_num:03d}"

    # Construct paths
    psalm_dir = args.output_dir / f"psalm_{psalm_num}"
    main_intro_path = psalm_dir / f"psalm_{psalm_num_str}_edited_intro.md"
    main_verses_path = psalm_dir / f"psalm_{psalm_num_str}_edited_verses.md"
    college_intro_path = psalm_dir / f"psalm_{psalm_num_str}_edited_intro_college.md"
    college_verses_path = psalm_dir / f"psalm_{psalm_num_str}_edited_verses_college.md"
    stats_path = psalm_dir / f"psalm_{psalm_num_str}_pipeline_stats.json"
    output_path = psalm_dir / f"psalm_{psalm_num_str}_commentary_combined.docx"

    # Check that all required files exist
    required_files = [
        main_intro_path, main_verses_path,
        college_intro_path, college_verses_path
    ]
    missing_files = [f for f in required_files if not f.exists()]

    if missing_files:
        print(f"Error: Missing required files:")
        for f in missing_files:
            print(f"  - {f}")
        return 1

    # Generate the combined document
    generator = CombinedDocumentGenerator(
        psalm_num=psalm_num,
        main_intro_path=main_intro_path,
        main_verses_path=main_verses_path,
        college_intro_path=college_intro_path,
        college_verses_path=college_verses_path,
        stats_path=stats_path,
        output_path=output_path
    )

    generator.generate()
    return 0


if __name__ == '__main__':
    sys.exit(main())
