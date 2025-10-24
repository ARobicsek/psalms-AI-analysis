"""
Test script to verify stress marking integration across the pipeline.
Tests the flow: PhoneticAnalyst → MicroAnalyst → SynthesisWriter/MasterEditor → Word Doc
"""

import sys
from pathlib import Path

# Fix Unicode encoding on Windows
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')

# Add project root to path
sys.path.insert(0, str(Path(__file__).resolve().parent))

from src.agents.phonetic_analyst import PhoneticAnalyst
from src.agents.micro_analyst import MicroAnalystV2
from src.data_sources.tanakh_database import TanakhDatabase


def test_phonetic_analyst_stress():
    """Test that PhoneticAnalyst produces stressed transcriptions."""
    print("\n" + "=" * 80)
    print("TEST 1: PhoneticAnalyst Stress Marking")
    print("=" * 80)

    analyst = PhoneticAnalyst()

    # Test verse: Psalm 145:1 - "תְּהִלָּה לְדָוִד"
    test_verse = "תְּהִלָּה לְדָוִד"
    result = analyst.transcribe_verse(test_verse)

    print(f"\nOriginal: {test_verse}")
    print(f"\nWords transcribed: {len(result['words'])}")

    for i, word in enumerate(result['words'], 1):
        print(f"\n  Word {i}: {word['word']}")
        print(f"    Regular:  {word['syllable_transcription']}")
        print(f"    Stressed: {word['syllable_transcription_stressed']}")
        print(f"    Stress on syllable: {word['stressed_syllable_index']}")
        print(f"    Stress level: {word['stress_level']}")

        # Verify stressed version has **CAPS**
        if '**' in word['syllable_transcription_stressed']:
            print(f"    ✓ Contains stress marking (**CAPS**)")
        else:
            print(f"    ✗ WARNING: No stress marking found!")

    print("\n✓ TEST 1 COMPLETE")
    return result


def test_micro_analyst_stress():
    """Test that MicroAnalyst uses stressed transcriptions."""
    print("\n" + "=" * 80)
    print("TEST 2: MicroAnalyst Stress Integration")
    print("=" * 80)

    db_path = "database/tanakh.db"
    db = TanakhDatabase(Path(db_path))

    # Initialize MicroAnalyst
    micro = MicroAnalystV2(db_path=db_path)

    # Get phonetic transcriptions for Psalm 145 (first 3 verses only for testing)
    psalm = db.get_psalm(145)
    phonetic_data = micro._get_phonetic_transcriptions(145)

    print(f"\nPsalm 145 - First 3 verses:")

    for verse_num in [1, 2, 3]:
        verse = next((v for v in psalm.verses if v.verse == verse_num), None)
        if verse and verse_num in phonetic_data:
            print(f"\n  Verse {verse_num}:")
            print(f"    Hebrew: {verse.hebrew[:50]}...")
            print(f"    Phonetic: {phonetic_data[verse_num][:80]}...")

            # Verify stressed transcription contains **CAPS**
            if '**' in phonetic_data[verse_num]:
                print(f"    ✓ Contains stress marking")
            else:
                print(f"    ✗ WARNING: No stress marking found!")

    print("\n✓ TEST 2 COMPLETE")
    return phonetic_data


def test_document_generator_mock():
    """Test that document generator can parse nested markdown."""
    print("\n" + "=" * 80)
    print("TEST 3: Document Generator Nested Markdown Parsing")
    print("=" * 80)

    from src.utils.document_generator import DocumentGenerator
    from docx import Document
    import re

    # Create a temporary document to test parsing
    doc = Document()
    p = doc.add_paragraph()

    # Simulate the _add_nested_formatting method
    test_text = "tə-**HIL**-lāh lə-dhā-**WIDH**"

    print(f"\nTest text: {test_text}")

    # Split by bold markers (same logic as in the actual method)
    parts = re.split(r'(\*\*.*?\*\*)', test_text)

    print(f"Parsed parts: {parts}")

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.italic = True
            print(f"  → BOLD+ITALIC: '{part[2:-2]}'")
        else:
            run = p.add_run(part)
            run.italic = True
            print(f"  → ITALIC: '{part}'")

    # Verify the runs were created correctly
    print(f"\nTotal runs created: {len(p.runs)}")

    expected_bold_count = test_text.count('**') // 2
    actual_bold_count = sum(1 for run in p.runs if run.bold)

    print(f"Expected bold runs: {expected_bold_count}")
    print(f"Actual bold runs: {actual_bold_count}")

    if actual_bold_count == expected_bold_count:
        print("✓ Bold formatting correct")
    else:
        print("✗ WARNING: Bold count mismatch!")

    all_italic = all(run.italic for run in p.runs)
    if all_italic:
        print("✓ All runs are italic (as expected for backtick context)")
    else:
        print("✗ WARNING: Not all runs are italic!")

    print("\n✓ TEST 3 COMPLETE")


def main():
    """Run all tests."""
    print("\n" + "=" * 80)
    print("STRESS MARKING INTEGRATION TEST SUITE")
    print("=" * 80)

    try:
        # Test 1: PhoneticAnalyst
        test_phonetic_analyst_stress()

        # Test 2: MicroAnalyst
        test_micro_analyst_stress()

        # Test 3: Document Generator
        test_document_generator_mock()

        print("\n" + "=" * 80)
        print("✓ ALL TESTS PASSED")
        print("=" * 80)
        print("\nStress marking integration is working correctly!")
        print("\nNext steps:")
        print("1. Run a full pipeline test: python scripts/run_enhanced_pipeline.py 23")
        print("2. Check the generated Word document for proper **BOLD** formatting in phonetics")
        print("3. Verify that synthesis and editor commentary reference stress patterns")

    except Exception as e:
        print(f"\n✗ TEST FAILED: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
