"""
Minimal test document to debug why the transformation isn't working in the full document.
"""
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement, ns
from pathlib import Path
import re
import sys

# Fix Windows console encoding
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')

# Import the DocumentGenerator class to use its methods
from src.utils.document_generator import DocumentGenerator

def set_paragraph_ltr(paragraph):
    """Set paragraph direction to LTR at XML level."""
    pPr = paragraph._element.get_or_add_pPr()
    bidi_elem = pPr.find(ns.qn('w:bidi'))
    if bidi_elem is not None:
        pPr.remove(bidi_elem)
    bidi_elem = OxmlElement('w:bidi')
    bidi_elem.set(ns.qn('w:val'), '0')
    pPr.append(bidi_elem)

# Create document
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Aptos'
style.font.size = Pt(12)

doc.add_heading('Minimal Test: Problematic Paragraph', level=1)

# The exact problematic text from the intro
problematic_text = 'In the Psalter, bones are often the deepest register of a person\'s condition: "All my bones shall say, \'YHWH, who is like you?\'" (Ps 35:10); "There is no wholeness in my bones because of my sin" (Ps 38:4); "I can count all my bones" (Ps 22:18). In Psalm 6, the fear runs to the frame. The next line intensifies the reach: "And my נֶפֶשׁ is deeply terrified" (וְנַפְשִׁי נִבְהֲלָה מְאֹד). נֶפֶשׁ is not a detachable soul; it is the living self, breath and appetite. Terror has progressed from skeleton to self.'

print("Testing problematic paragraph...")
print(f"Text length: {len(problematic_text)} chars")
print()

# Method 1: Add as single paragraph (what we expect to happen)
doc.add_heading('Method 1: Single Paragraph via _add_paragraph_with_soft_breaks', level=2)

# Create a paragraph
p1 = doc.add_paragraph(style='Normal')
set_paragraph_ltr(p1)

# Apply the transformation manually
LRO = '\u202D'
PDF = '\u202C'

def apply_transformation(text):
    hebrew_paren_pattern = r'\(([\u0590-\u05FF\u05B0-\u05BD\u05BF\u05C1-\u05C2\u05C4-\u05C7\s]+)\)'

    def reverse_hebrew_match(match):
        hebrew_text = match.group(1)
        reversed_hebrew = DocumentGenerator._reverse_hebrew_by_clusters(hebrew_text)
        return f'{LRO}({reversed_hebrew}){PDF}'

    return re.sub(hebrew_paren_pattern, reverse_hebrew_match, text)

transformed_text = apply_transformation(problematic_text)
print(f"Transformed text contains {transformed_text.count(LRO)} LRO characters")
print()

# Find all matches
pattern = r'\(([\u0590-\u05FF\u05B0-\u05BD\u05BF\u05C1-\u05C2\u05C4-\u05C7\s]+)\)'
matches = list(re.finditer(pattern, problematic_text))
print(f"Found {len(matches)} Hebrew parentheses in original text:")
for i, match in enumerate(matches):
    print(f"  {i+1}. '{match.group(0)}'")
print()

# Add the transformed text
p1.add_run(transformed_text)

doc.add_paragraph()

# Method 2: Simulate what _add_paragraph_with_markdown does
doc.add_heading('Method 2: Via _add_paragraph_with_markdown Simulation', level=2)

# Split by markdown (simulating the regex split)
parts = re.split(r'(\*\*|\*|_|`)', problematic_text)
print(f"Split into {len(parts)} parts by markdown delimiters")

p2 = doc.add_paragraph(style='Normal')
set_paragraph_ltr(p2)

for part in parts:
    # Apply transformation to each part
    transformed_part = apply_transformation(part)
    p2.add_run(transformed_part)

# Save
output_path = Path('output/bidi_tests/minimal_test_debug.docx')
doc.save(output_path)

print(f"\nDocument saved: {output_path}")
print("\nPlease open this document in Word and check:")
print("1. Does Method 1 render correctly?")
print("2. Does Method 2 render correctly?")
print("3. If one works and the other doesn't, that tells us where the bug is!")
