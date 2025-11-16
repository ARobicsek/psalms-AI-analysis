"""
Test script to systematically try 4 creative solutions for the bidirectional text rendering bug.

This script generates 5 test documents:
1. Control (current LRI/PDI approach - known to fail)
2. Solution 1: Hebrew Ornate Parentheses (U+FD3E/U+FD3F)
3. Solution 2: Zero-Width Joiner (U+200D)
4. Solution 3: LEFT-TO-RIGHT OVERRIDE (U+202D)
5. Solution 4: Pre-mirrored parentheses

Each document contains the same test text with parenthesized Hebrew.
Open each in Word to see which rendering is correct.
"""

from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement, ns
from pathlib import Path


def set_paragraph_ltr(paragraph):
    """Set paragraph direction to LTR at XML level."""
    pPr = paragraph._element.get_or_add_pPr()
    bidi_elem = pPr.find(ns.qn('w:bidi'))
    if bidi_elem is not None:
        pPr.remove(bidi_elem)
    bidi_elem = OxmlElement('w:bidi')
    bidi_elem.set(ns.qn('w:val'), '0')
    pPr.append(bidi_elem)


def create_test_document(output_path, solution_name, transform_func):
    """
    Create a test document with a specific bidi solution applied.

    Args:
        output_path: Path to save the document
        solution_name: Name of the solution being tested
        transform_func: Function that transforms the text with parenthesized Hebrew
    """
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Aptos'
    style.font.size = Pt(12)

    # Add title
    doc.add_heading(f'Bidi Test: {solution_name}', level=1)

    # Add explanation
    p_explain = doc.add_paragraph()
    p_explain.add_run('Expected: ').bold = True
    p_explain.add_run('Parenthesized Hebrew should appear as ')
    p_explain.add_run('(וְנַפְשִׁי נִבְהֲלָה מְאֹד)').bold = True
    p_explain.add_run(' with Hebrew reading right-to-left inside the parentheses.')

    doc.add_paragraph()

    # Test cases with increasing complexity
    test_cases = [
        {
            'label': 'Test 1: Simple parenthesized Hebrew',
            'text': 'And my נֶפֶשׁ is deeply terrified (וְנַפְשִׁי נִבְהֲלָה מְאֹד). נֶפֶשׁ is not a detachable soul.'
        },
        {
            'label': 'Test 2: Multiple instances',
            'text': 'The word אַשְׁרֵי (ashrei) means blessed (אַשְׁרֵי הָאִישׁ), and תְּהִלָּה (tehillah) means praise (תְּהִלָּה לְדָוִד).'
        },
        {
            'label': 'Test 3: Hebrew with spaces',
            'text': 'The phrase "Happy is the man" (אַשְׁרֵי הָאִישׁ אֲשֶׁר) is the opening of Psalm 1.'
        },
        {
            'label': 'Test 4: Mixed nikud',
            'text': 'The root ברך appears in בָּרוּךְ (blessed) and בְּרָכָה (blessing).'
        },
        {
            'label': 'Test 5: Long Hebrew text',
            'text': 'The entire first verse (אַשְׁרֵי הָאִישׁ אֲשֶׁר לֹא הָלַךְ בַּעֲצַת רְשָׁעִים) sets the tone for the psalm.'
        }
    ]

    for test_case in test_cases:
        # Add test case label
        p_label = doc.add_paragraph()
        set_paragraph_ltr(p_label)
        p_label.add_run(test_case['label']).bold = True

        # Add test case with transformation applied
        p_test = doc.add_paragraph()
        set_paragraph_ltr(p_test)

        # Apply the transformation function
        transformed_text = transform_func(test_case['text'])
        p_test.add_run(transformed_text)

        # Add spacing
        doc.add_paragraph()

    # Save document
    doc.save(output_path)
    print(f"OK: Generated {output_path.name}")


# Solution transformation functions

def control_lri_pdi(text):
    """Control: Current approach using LRI/PDI isolates (known to fail)."""
    import re
    LRI = '\u2066'  # Left-to-Right Isolate
    PDI = '\u2069'  # Pop Directional Isolate

    hebrew_paren_pattern = r'([\(][\u0590-\u05FF\u05B0-\u05BD\u05BF\u05C1-\u05C2\u05C4-\u05C7\s]+[\)])'
    return re.sub(
        hebrew_paren_pattern,
        lambda m: f'{LRI}{m.group(0)}{PDI}',
        text
    )


def solution1_ornate_parens(text):
    """Solution 1: Use Hebrew ornate parentheses (U+FD3E/U+FD3F) - RTL characters."""
    import re

    # Hebrew ornate parentheses - these are explicitly RTL characters
    ORNATE_OPEN = '\uFD3E'  # ﴾
    ORNATE_CLOSE = '\uFD3F'  # ﴿

    # Replace ASCII parentheses around Hebrew with ornate parentheses
    hebrew_paren_pattern = r'\(([\u0590-\u05FF\u05B0-\u05BD\u05BF\u05C1-\u05C2\u05C4-\u05C7\s]+)\)'
    return re.sub(
        hebrew_paren_pattern,
        lambda m: f'{ORNATE_OPEN}{m.group(1)}{ORNATE_CLOSE}',
        text
    )


def solution2_zwj(text):
    """Solution 2: Zero-Width Joiner to create grapheme cluster."""
    import re

    ZWJ = '\u200D'  # Zero-Width Joiner

    # Insert ZWJ immediately after opening paren and before closing paren
    hebrew_paren_pattern = r'\(([\u0590-\u05FF\u05B0-\u05BD\u05BF\u05C1-\u05C2\u05C4-\u05C7\s]+)\)'
    return re.sub(
        hebrew_paren_pattern,
        lambda m: f'({ZWJ}{m.group(1)}{ZWJ})',
        text
    )


def solution3_lro(text):
    """Solution 3: LEFT-TO-RIGHT OVERRIDE (stronger than embedding)."""
    import re

    LRO = '\u202D'  # LEFT-TO-RIGHT OVERRIDE
    PDF = '\u202C'  # POP DIRECTIONAL FORMATTING

    # Wrap entire (Hebrew) with LRO...PDF
    hebrew_paren_pattern = r'([\(][\u0590-\u05FF\u05B0-\u05BD\u05BF\u05C1-\u05C2\u05C4-\u05C7\s]+[\)])'
    return re.sub(
        hebrew_paren_pattern,
        lambda m: f'{LRO}{m.group(0)}{PDF}',
        text
    )


def solution4_pre_mirror(text):
    """Solution 4: Pre-mirror the parentheses (reverse engineering)."""
    import re

    # Insert parentheses in reverse: )Hebrew(
    # Hypothesis: Word's bidi algorithm will mirror them back to (Hebrew)
    hebrew_paren_pattern = r'\(([\u0590-\u05FF\u05B0-\u05BD\u05BF\u05C1-\u05C2\u05C4-\u05C7\s]+)\)'
    return re.sub(
        hebrew_paren_pattern,
        lambda m: f'){m.group(1)}(',
        text
    )


def main():
    """Generate all test documents."""
    output_dir = Path('output/bidi_tests')
    output_dir.mkdir(parents=True, exist_ok=True)

    print("Generating bidirectional text test documents...\n")

    # Test configurations
    tests = [
        ('control_lri_pdi', 'Control (LRI/PDI - Current)', control_lri_pdi),
        ('solution1_ornate_parens', 'Solution 1: Hebrew Ornate Parentheses (U+FD3E/FD3F)', solution1_ornate_parens),
        ('solution2_zwj', 'Solution 2: Zero-Width Joiner (U+200D)', solution2_zwj),
        ('solution3_lro', 'Solution 3: LEFT-TO-RIGHT OVERRIDE (U+202D)', solution3_lro),
        ('solution4_pre_mirror', 'Solution 4: Pre-Mirrored Parentheses', solution4_pre_mirror),
    ]

    for filename, solution_name, transform_func in tests:
        output_path = output_dir / f'{filename}.docx'
        create_test_document(output_path, solution_name, transform_func)

    print(f"\n{'='*70}")
    print("Test documents generated successfully!")
    print(f"{'='*70}")
    print(f"\nLocation: {output_dir.absolute()}")
    print("\nNext steps:")
    print("1. Open each .docx file in Microsoft Word")
    print("2. Check if parenthesized Hebrew renders correctly")
    print("3. Note which solution(s) work correctly")
    print("4. Report back which file works best!")
    print("\nExpected rendering:")
    print("  - Hebrew text should be INSIDE the parentheses")
    print("  - Hebrew should read right-to-left")
    print("  - No duplication or splitting of text")
    print("\nFiles generated:")
    print("  - control_lri_pdi.docx (current approach - known to fail)")
    print("  - solution1_ornate_parens.docx (Hebrew ornate parentheses)")
    print("  - solution2_zwj.docx (Zero-Width Joiner)")
    print("  - solution3_lro.docx (LEFT-TO-RIGHT OVERRIDE)")
    print("  - solution4_pre_mirror.docx (Pre-mirrored parentheses)")


if __name__ == '__main__':
    main()
