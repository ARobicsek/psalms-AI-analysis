#!/usr/bin/env python3
"""
Comprehensive verification against FORMATTING_GUIDE.md checklist.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

def hex_to_rgb(hex_color):
    """Convert hex color to RGB tuple."""
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

def get_run_color_hex(run):
    """Get run color as hex string."""
    if run.font.color and run.font.color.rgb:
        rgb = run.font.color.rgb
        return f"#{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    return None

def check_heading_format(para, expected_font, expected_size, expected_color, expected_bold, expected_italic=None):
    """Check if heading has correct formatting."""
    if not para.runs:
        return {"error": "No runs found"}

    run = para.runs[0]
    issues = []

    # Font name
    actual_font = run.font.name
    if actual_font and actual_font != expected_font:
        issues.append(f"Font: {actual_font} (expected {expected_font})")

    # Font size
    actual_size = run.font.size.pt if run.font.size else None
    if actual_size and abs(actual_size - expected_size) > 0.1:
        issues.append(f"Size: {actual_size}pt (expected {expected_size}pt)")

    # Color
    actual_color = get_run_color_hex(run)
    if actual_color and actual_color.upper() != f"#{expected_color.upper()}":
        issues.append(f"Color: {actual_color} (expected #{expected_color})")

    # Bold
    actual_bold = run.font.bold
    if actual_bold is not None and actual_bold != expected_bold:
        issues.append(f"Bold: {actual_bold} (expected {expected_bold})")

    # Italic (for H4)
    if expected_italic is not None:
        actual_italic = run.font.italic
        if actual_italic is not None and actual_italic != expected_italic:
            issues.append(f"Italic: {actual_italic} (expected {expected_italic})")

    return {
        "text": para.text[:50],
        "issues": issues,
        "ok": len(issues) == 0
    }

def verify_document(doc_path):
    """Run comprehensive verification."""
    doc = Document(doc_path)

    print("=" * 80)
    print("COMPREHENSIVE FORMATTING VERIFICATION")
    print("=" * 80)

    # Heading specifications
    heading_specs = {
        'Heading 1': {'font': 'Times New Roman', 'size': 17, 'bold': True, 'color': '2E74B5'},
        'Heading 2': {'font': 'Times New Roman', 'size': 14, 'bold': True, 'color': '2E74B5'},
        'Heading 3': {'font': 'Times New Roman', 'size': 13, 'bold': True, 'color': '1F4D78'},
        'Heading 4': {'font': 'Times New Roman', 'size': 11, 'bold': False, 'italic': True, 'color': '2E74B5'},
    }

    # Check each heading type
    for heading_type, specs in heading_specs.items():
        print(f"\n### {heading_type.upper()} ###")

        headings = [p for p in doc.paragraphs if p.style and p.style.name == heading_type]
        if not headings:
            print(f"  ⚠️  No {heading_type} paragraphs found!")
            continue

        print(f"  Found {len(headings)} paragraphs")

        # Check first 3 samples
        issues_count = 0
        for i, para in enumerate(headings[:3]):
            result = check_heading_format(
                para,
                specs['font'],
                specs['size'],
                specs['color'],
                specs['bold'],
                specs.get('italic')
            )

            if not result['ok']:
                issues_count += 1
                print(f"\n  Sample {i+1}: '{result['text']}'")
                for issue in result['issues']:
                    print(f"    ❌ {issue}")
            elif i == 0:  # Always show first sample
                print(f"  ✓ Sample 1 looks correct: '{result['text']}'")

        if issues_count == 0:
            print(f"  ✓ All samples checked OK")
        else:
            print(f"  ⚠️  Found {issues_count} samples with issues (check all {len(headings)} paragraphs)")

    # Check body text
    print("\n### BODY TEXT (NON-HEADING, NON-LIST) ###")
    body_paras = [p for p in doc.paragraphs
                  if p.style and p.style.name not in ['Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'List Paragraph']]
    print(f"  Found {len(body_paras)} body paragraphs")

    if body_paras:
        # Check first few samples
        body_issues = 0
        for i, para in enumerate(body_paras[:3]):
            if not para.runs:
                continue

            run = para.runs[0]
            actual_font = run.font.name
            actual_size = run.font.size.pt if run.font.size else None

            issues = []
            if actual_font and actual_font != 'Arial':
                issues.append(f"Font: {actual_font} (expected Arial)")
            if actual_size and abs(actual_size - 11) > 0.1:
                issues.append(f"Size: {actual_size}pt (expected 11pt)")

            if issues:
                body_issues += 1
                print(f"  Sample {i+1}: '{para.text[:50]}'")
                for issue in issues:
                    print(f"    ❌ {issue}")

        if body_issues == 0:
            print(f"  ✓ Body text samples look correct")

    # Check List Paragraph
    print("\n### LIST PARAGRAPH (BULLETS) ###")
    list_paras = [p for p in doc.paragraphs if p.style and p.style.name == 'List Paragraph']
    print(f"  Found {len(list_paras)} list paragraphs")

    if list_paras:
        list_issues = 0
        for i, para in enumerate(list_paras[:3]):
            if not para.runs:
                continue

            run = para.runs[0]
            actual_font = run.font.name
            actual_size = run.font.size.pt if run.font.size else None

            issues = []
            if actual_font and actual_font != 'Arial':
                issues.append(f"Font: {actual_font} (expected Arial)")
            if actual_size and abs(actual_size - 11) > 0.1:
                issues.append(f"Size: {actual_size}pt (expected 11pt)")

            if issues:
                list_issues += 1
                print(f"  Sample {i+1}: '{para.text[:50]}'")
                for issue in issues:
                    print(f"    ❌ {issue}")

        if list_issues == 0:
            print(f"  ✓ List paragraph samples look correct")

    # Document structure
    print("\n### DOCUMENT STRUCTURE ###")
    h1_paras = [p for p in doc.paragraphs if p.style and p.style.name == 'Heading 1']
    print(f"  Heading 1 count: {len(h1_paras)} (expected 3-4)")
    for h1 in h1_paras:
        print(f"    - {h1.text}")

    h2_paras = [p for p in doc.paragraphs if p.style and p.style.name == 'Heading 2']
    print(f"  Heading 2 count: {len(h2_paras)} (expected ~23)")

    # Check bullet character
    print("\n### BULLET CHARACTER ###")
    if hasattr(doc, '_part') and hasattr(doc._part, 'numbering_part'):
        numbering_part = doc._part.numbering_part
        if numbering_part:
            numbering_xml = numbering_part._element
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            abstract_nums = numbering_xml.findall('.//w:abstractNum', ns)

            for abs_num in abstract_nums:
                abs_num_id = abs_num.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}abstractNumId')
                levels = abs_num.findall('.//w:lvl', ns)
                lvl = levels[0] if levels else None
                if lvl:
                    lvl_text = lvl.find('.//w:lvlText', ns)
                    if lvl_text is not None:
                        bullet_char = lvl_text.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                        if bullet_char == '•':
                            print(f"  ✓ Abstract Num {abs_num_id}, Level 0: Using '•' (correct small bullet)")
                        elif bullet_char == '●':
                            print(f"  ❌ Abstract Num {abs_num_id}, Level 0: Using '●' (large filled circle - should be '•')")
                        else:
                            print(f"  ℹ️  Abstract Num {abs_num_id}, Level 0: Using '{bullet_char}'")

    print("\n" + "=" * 80)

if __name__ == "__main__":
    doc_path = "/home/user/psalms-AI-analysis/docs/Cordance/Cordance_insights_bank_draft.docx"
    verify_document(doc_path)
