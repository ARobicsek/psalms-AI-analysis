#!/usr/bin/env python3
"""
Fix header formatting with correct specifications:
- Heading 1: Times New Roman 17pt, #2E74B5
- Heading 2: Times New Roman 14pt, #2E74B5
- Heading 3: Times New Roman 13pt, #1F4D78
- Heading 4: Times New Roman italic 11pt, #2E74B5
- Body text remains Arial 11pt
"""

from docx import Document
from docx.shared import Pt, RGBColor

def fix_header_formatting():
    """Apply correct header formatting to Primary Care document."""

    print("Loading Primary Care document...")
    doc = Document('docs/Cordance/Cordance_Health_Insights_Bank_Primary_Care.docx')

    # Define correct header specifications
    header_specs = {
        'Heading 1': {
            'font': 'Times New Roman',
            'size': 17,
            'color': RGBColor(0x2E, 0x74, 0xB5),
            'bold': True,
            'italic': False
        },
        'Heading 2': {
            'font': 'Times New Roman',
            'size': 14,
            'color': RGBColor(0x2E, 0x74, 0xB5),
            'bold': True,
            'italic': False
        },
        'Heading 3': {
            'font': 'Times New Roman',
            'size': 13,
            'color': RGBColor(0x1F, 0x4D, 0x78),
            'bold': True,
            'italic': False
        },
        'Heading 4': {
            'font': 'Times New Roman',
            'size': 11,
            'color': RGBColor(0x2E, 0x74, 0xB5),
            'bold': False,
            'italic': True
        },
    }

    print(f"\nProcessing {len(doc.paragraphs)} paragraphs...")

    headers_updated = 0

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else "None"

        # Update headers with correct formatting
        if style_name in header_specs:
            spec = header_specs[style_name]
            for run in para.runs:
                run.font.name = spec['font']
                run.font.size = Pt(spec['size'])
                run.font.color.rgb = spec['color']
                run.bold = spec['bold']
                run.italic = spec['italic']
            headers_updated += 1

    print(f"\n✓ Updated {headers_updated} header paragraphs with correct formatting:")
    for style, spec in header_specs.items():
        style_desc = f"{spec['font']} {spec['size']}pt"
        if spec['bold']:
            style_desc += " bold"
        if spec['italic']:
            style_desc += " italic"
        # Extract RGB values directly from RGBColor object
        r, g, b = spec['color']
        color_hex = f"#{r:02X}{g:02X}{b:02X}"
        print(f"  - {style}: {style_desc}, {color_hex}")

    # Save the document
    output_path = 'docs/Cordance/Cordance_Health_Insights_Bank_Primary_Care.docx'
    print(f"\nSaving to {output_path}...")
    doc.save(output_path)

    print("\n✅ Header formatting corrected!")

if __name__ == '__main__':
    fix_header_formatting()
