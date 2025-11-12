"""
Solution 5: Reverse Hebrew + LEFT-TO-RIGHT OVERRIDE (Hybrid Approach)

Based on user testing, Solution 3 (LRO) kept text inside parentheses without duplication,
but displayed the Hebrew backwards. This solution reverses the Hebrew text FIRST,
then applies LRO, so the backwards text displayed LTR results in correct RTL appearance.
"""

from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement, ns
from pathlib import Path
import re


def set_paragraph_ltr(paragraph):
    """Set paragraph direction to LTR at XML level."""
    pPr = paragraph._element.get_or_add_pPr()
    bidi_elem = pPr.find(ns.qn('w:bidi'))
    if bidi_elem is not None:
        pPr.remove(bidi_elem)
    bidi_elem = OxmlElement('w:bidi')
    bidi_elem.set(ns.qn('w:val'), '0')
    pPr.append(bidi_elem)


def solution5_reverse_hebrew_plus_lro(text):
    """
    Solution 5: Pre-reverse Hebrew text + LEFT-TO-RIGHT OVERRIDE

    Strategy:
    1. Extract Hebrew text from inside parentheses
    2. Reverse the Hebrew character order (e.g., "ABC" -> "CBA")
    3. Wrap with LRO...PDF to force left-to-right display
    4. Result: Reversed Hebrew displayed LTR = correct RTL visual appearance
    """
    LRO = '\u202D'  # LEFT-TO-RIGHT OVERRIDE
    PDF = '\u202C'  # POP DIRECTIONAL FORMATTING

    def reverse_hebrew_in_match(match):
        """Extract Hebrew, reverse it, wrap with LRO."""
        hebrew_text = match.group(1)

        # Reverse the Hebrew text character by character
        reversed_hebrew = hebrew_text[::-1]

        # Wrap the entire (reversed_hebrew) with LRO...PDF
        return f'{LRO}({reversed_hebrew}){PDF}'

    # Pattern to match (Hebrew text with nikud and spaces)
    hebrew_paren_pattern = r'\(([\u0590-\u05FF\u05B0-\u05BD\u05BF\u05C1-\u05C2\u05C4-\u05C7\s]+)\)'

    return re.sub(hebrew_paren_pattern, reverse_hebrew_in_match, text)


def create_test_document():
    """Create test document with Solution 5."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Aptos'
    style.font.size = Pt(12)

    # Add title
    doc.add_heading('Solution 5: Reversed Hebrew + LRO (HYBRID)', level=1)

    # Add explanation
    p_explain = doc.add_paragraph()
    p_explain.add_run('Strategy: ').bold = True
    p_explain.add_run('Pre-reverse Hebrew text, then apply LEFT-TO-RIGHT OVERRIDE. ')
    p_explain.add_run('Result: ').bold = True
    p_explain.add_run('Reversed text displayed LTR = correct RTL appearance!')

    doc.add_paragraph()

    # Test cases
    test_cases = [
        {
            'label': 'Test 1: Simple parenthesized Hebrew',
            'text': 'And my נֶפֶשׁ is deeply terrified (וְנַפְשִׁי נִבְהֲלָה מְאֹד). נֶפֶשׁ is not a detachable soul.'
        },
        {
            'label': 'Test 2: Multiple instances',
            'text': 'The word אַשְׁרֵי (ashrei) means blessed (אַשְׁרֵי הָאִישׁ), and תְּהִלָּה (tehillah) means praise (תְּהִלָּה לְדָוִד).'
        },
        {
            'label': 'Test 3: Hebrew with spaces',
            'text': 'The phrase "Happy is the man" (אַשְׁרֵי הָאִישׁ אֲשֶׁר) is the opening of Psalm 1.'
        },
        {
            'label': 'Test 4: Mixed nikud',
            'text': 'The root ברך appears in בָּרוּךְ (blessed) and בְּרָכָה (blessing).'
        },
        {
            'label': 'Test 5: Long Hebrew text',
            'text': 'The entire first verse (אַשְׁרֵי הָאִישׁ אֲשֶׁר לֹא הָלַךְ בַּעֲצַת רְשָׁעִים) sets the tone for the psalm.'
        }
    ]

    for test_case in test_cases:
        # Add test case label
        p_label = doc.add_paragraph()
        set_paragraph_ltr(p_label)
        p_label.add_run(test_case['label']).bold = True

        # Add test case with Solution 5 applied
        p_test = doc.add_paragraph()
        set_paragraph_ltr(p_test)

        transformed_text = solution5_reverse_hebrew_plus_lro(test_case['text'])
        p_test.add_run(transformed_text)

        # Add spacing
        doc.add_paragraph()

    # Add comparison section
    doc.add_heading('How This Works', level=2)

    p_how = doc.add_paragraph()
    p_how.add_run('Original Hebrew: ').bold = True
    p_how.add_run('וְנַפְשִׁי נִבְהֲלָה מְאֹד (reads right-to-left)')

    p_reversed = doc.add_paragraph()
    p_reversed.add_run('Reversed Hebrew: ').bold = True
    p_reversed.add_run('דֹאמ הָלֲהבְנִ יִשְׁפַנְוֶ (character order flipped)')

    p_lro = doc.add_paragraph()
    p_lro.add_run('LRO forces LTR: ').bold = True
    p_lro.add_run('Reversed text displayed left-to-right = correct RTL visual result!')

    # Save
    output_path = Path('output/bidi_tests/solution5_reversed_hebrew_lro.docx')
    doc.save(output_path)
    print(f"OK: Generated {output_path.name}")
    print(f"\nLocation: {output_path.absolute()}")
    print("\nPlease open this document in Word and check if ALL test cases render correctly!")
    print("Expected: Hebrew text inside parentheses, reading right-to-left, no duplication")


if __name__ == '__main__':
    create_test_document()
