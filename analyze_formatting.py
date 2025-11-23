#!/usr/bin/env python3
"""
Analyze the Cordance document formatting to identify issues.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from collections import defaultdict

def rgb_to_hex(rgb_color):
    """Convert RGBColor to hex string."""
    if rgb_color is None:
        return "None"
    return f"#{rgb_color[0]:02X}{rgb_color[1]:02X}{rgb_color[2]:02X}"

def analyze_document(doc_path):
    """Analyze document formatting."""
    doc = Document(doc_path)

    print("=" * 80)
    print("DOCUMENT FORMATTING ANALYSIS")
    print("=" * 80)

    # Track heading counts
    heading_counts = defaultdict(int)
    bullet_issues = []
    heading_formatting = defaultdict(list)

    # Analyze each paragraph
    for idx, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else "No Style"

        # Count headings
        if 'Heading' in style_name:
            heading_counts[style_name] += 1

            # Check heading formatting
            if para.runs:
                run = para.runs[0]
                font_name = run.font.name
                font_size = run.font.size.pt if run.font.size else "Unknown"
                font_bold = run.font.bold
                font_color = rgb_to_hex(run.font.color.rgb) if run.font.color and run.font.color.rgb else "None"

                heading_formatting[style_name].append({
                    'text': para.text[:50],
                    'font': font_name,
                    'size': font_size,
                    'bold': font_bold,
                    'color': font_color
                })

        # Check bullet formatting
        if style_name == 'List Paragraph' or para.text.strip().startswith('•'):
            if para.runs:
                # Check for oversized bullets
                for run_idx, run in enumerate(para.runs):
                    if run.text.strip().startswith('•'):
                        font_size = run.font.size.pt if run.font.size else "Unknown"
                        if isinstance(font_size, (int, float)) and font_size > 12:
                            bullet_issues.append({
                                'para_idx': idx,
                                'text': para.text[:60],
                                'bullet_size': font_size,
                                'style': style_name
                            })

    # Print heading counts
    print("\n### HEADING COUNTS ###")
    for heading_type in ['Heading 1', 'Heading 2', 'Heading 3', 'Heading 4']:
        count = heading_counts[heading_type]
        print(f"{heading_type}: {count}")

    # Print heading formatting samples
    print("\n### HEADING FORMATTING SAMPLES ###")
    for heading_type in ['Heading 1', 'Heading 2', 'Heading 3', 'Heading 4']:
        if heading_formatting[heading_type]:
            print(f"\n{heading_type}:")
            for i, sample in enumerate(heading_formatting[heading_type][:2]):  # First 2 samples
                print(f"  Sample {i+1}: '{sample['text']}'")
                print(f"    Font: {sample['font']}, Size: {sample['size']}, Bold: {sample['bold']}, Color: {sample['color']}")

    # Print bullet issues
    print("\n### BULLET POINT ISSUES ###")
    if bullet_issues:
        print(f"Found {len(bullet_issues)} paragraphs with oversized bullets:")
        for issue in bullet_issues[:10]:  # Show first 10
            print(f"  Para {issue['para_idx']}: '{issue['text']}...'")
            print(f"    Bullet size: {issue['bullet_size']}pt (should be ~11pt)")
    else:
        print("No oversized bullet issues detected in run formatting.")

    # Check for sections
    print("\n### DOCUMENT SECTIONS ###")
    h1_sections = [p.text for p in doc.paragraphs if p.style and p.style.name == 'Heading 1']
    for section in h1_sections:
        print(f"  - {section}")

    print("\n" + "=" * 80)

    return {
        'heading_counts': dict(heading_counts),
        'bullet_issues': bullet_issues,
        'h1_sections': h1_sections
    }

if __name__ == "__main__":
    doc_path = "/home/user/psalms-AI-analysis/docs/Cordance/Cordance_insights_bank_draft.docx"
    analyze_document(doc_path)
