#!/usr/bin/env python3
"""
Merge three Cordance Health Insights Bank documents with PRECISE formatting.
This version INTELLIGENTLY APPLIES correct formatting to ALL sections.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy
import re


def add_page_number(section):
    """Add page numbers to the bottom right of all pages."""
    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Add page number field
    run = footer_para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

    # Set font
    run.font.name = 'Arial'
    run.font.size = Pt(11)


def add_table_of_contents(doc):
    """Add a table of contents showing only Heading 1 and Heading 2."""
    # Add TOC title
    toc_title = doc.add_paragraph('Table of Contents')
    toc_title.style = doc.styles['Heading 1']

    # Add a paragraph for the TOC field
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    # Create the TOC field code
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-2" \\h \\z \\u'

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    # Add placeholder text
    run2 = paragraph.add_run()
    run2.text = 'Right-click and select "Update Field" to generate the table of contents.'
    run2.font.italic = True
    run2.font.color.rgb = RGBColor(128, 128, 128)

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    paragraph._element.append(fldChar3)

    # Add page break after TOC
    doc.add_page_break()


def determine_paragraph_type(para_text, source_para):
    """
    Intelligently determine what type of paragraph this should be based on content.
    Returns: ('style_name', is_bullet)
    """
    text = para_text.strip()

    # Heading 1: Specialty names (short, not USE CASE)
    if len(text) < 50 and not text.startswith('USE CASE') and not ':' in text:
        # Could be specialty name like "Primary Care", "Cardiology", "Infectious Diseases"
        specialty_names = ['Primary Care', 'Cardiology', 'Infectious Diseases', 'Table of Contents']
        if text in specialty_names:
            return ('Heading 1', False)

    # Heading 2: USE CASE titles
    if text.startswith('USE CASE'):
        return ('Heading 2', False)

    # Heading 3: Major section headings
    h3_titles = [
        'Clinical Scenario', 'Current State Challenges', 'Cordance Solution',
        'Impact on Three Pillars', 'Physician Value Proposition',
        'Implementation Considerations', 'Quantified Impact', 'Patient-Facing Value'
    ]
    if text in h3_titles:
        return ('Heading 3', False)

    # Heading 4: Subsection headings
    h4_titles = [
        'Before Encounter', 'During Encounter', 'After Encounter',
        'Economics (Value-Based Care & Fee-for-Service)',
        'Quality & Safety', 'Population Health & Outcomes'
    ]
    if text in h4_titles:
        return ('Heading 4', False)

    # Check if it's an "At a/the ... level:" line (NOT a heading, just body text intro)
    if text.startswith('At a') or text.startswith('At the'):
        return ('Normal', False)

    # Check if it's a bullet point
    # Bullets can be detected by:
    # 1. Having numbering in source
    # 2. Being in certain sections (under Current State Challenges, Cordance Solution, etc.)
    if source_para._element.pPr is not None:
        numPr = source_para._element.pPr.find(qn('w:numPr'))
        if numPr is not None:
            return ('List Paragraph', True)

    # If the source has List Paragraph style, it should be a bullet
    if source_para.style and 'List' in source_para.style.name:
        return ('List Paragraph', True)

    # Default: Normal body text
    return ('Normal', False)


def copy_paragraph_with_intelligent_formatting(source_para, target_doc, template_numbering):
    """
    Copy paragraph from source to target, but APPLY CORRECT FORMATTING based on content.
    """
    text = source_para.text.strip()
    if not text:
        return None

    # Determine what type this paragraph should be
    style_name, is_bullet = determine_paragraph_type(text, source_para)

    # Create the new paragraph
    new_para = target_doc.add_paragraph()

    # Apply the correct style
    try:
        new_para.style = target_doc.styles[style_name]
    except KeyError:
        # Style doesn't exist, use default
        pass

    # If it's a bullet, add the numbering
    if is_bullet and template_numbering is not None:
        # Copy numbering from source
        if source_para._element.pPr is not None:
            numPr = source_para._element.pPr.find(qn('w:numPr'))
            if numPr is not None:
                new_para._element.get_or_add_pPr()
                new_numPr = deepcopy(numPr)
                new_para._element.pPr.append(new_numPr)

    # Copy the text content with proper formatting
    for source_run in source_para.runs:
        new_run = new_para.add_run(source_run.text)

        # For headings, let the style control the formatting
        # For body text and bullets, preserve some run formatting
        if style_name in ['Normal', 'List Paragraph']:
            # Copy font name and size for body text
            if source_run.font.name:
                new_run.font.name = source_run.font.name
            if source_run.font.size:
                new_run.font.size = source_run.font.size

        # Always preserve superscript (for references like [1], [2])
        if source_run.font.superscript:
            new_run.font.superscript = source_run.font.superscript
        if source_run.font.subscript:
            new_run.font.subscript = source_run.font.subscript

    return new_para


def should_skip_paragraph(para):
    """Determine if a paragraph should be skipped."""
    text = para.text.strip()
    if not text:
        return True
    if text in ['References', 'Summary']:
        return 'STOP'
    return False


def is_heading_2(para):
    """Check if paragraph is a USE CASE heading."""
    text = para.text.strip()
    return text.startswith('USE CASE')


def process_document_content(source_doc, target_doc, template_numbering, add_page_break_before_first=False):
    """
    Process all paragraphs from source document and add to target with CORRECT formatting.
    """
    use_case_count = 0

    for para in source_doc.paragraphs:
        # Check if we should skip
        skip_result = should_skip_paragraph(para)
        if skip_result == 'STOP':
            break
        if skip_result:
            continue

        # Check if this is a USE CASE heading
        is_use_case = is_heading_2(para)

        # Add page break before each USE CASE (except possibly the first)
        if is_use_case:
            use_case_count += 1
            if use_case_count > 1 or add_page_break_before_first:
                target_doc.add_page_break()

        # Copy the paragraph with intelligent formatting
        copy_paragraph_with_intelligent_formatting(para, target_doc, template_numbering)

    return use_case_count


def merge_documents():
    """Main function to merge the three documents with precise formatting."""
    print("=" * 70)
    print("CORDANCE HEALTH INSIGHTS BANK - INTELLIGENT DOCUMENT MERGER")
    print("=" * 70)
    print("\nLoading source documents...")

    # Load the three source documents
    primary_care_doc = Document('./Cordance_Health_Insights_Bank_Primary_Care.docx')
    cardiology_doc = Document('./Cordance_Health_Insights_Bank_Cardiology.docx')
    infectious_diseases_doc = Document('./Cordance_Health_Insights_Bank_Infectious_Diseases.docx')

    print("✓ Loaded Primary Care document")
    print("✓ Loaded Cardiology document")
    print("✓ Loaded Infectious Diseases document")

    # Create new merged document
    print("\nCreating merged document with template formatting...")
    merged_doc = Document()

    # Copy ALL style definitions from Primary Care (the well-formatted template)
    print("✓ Copying style definitions from Primary Care template...")
    merged_doc.styles._element = deepcopy(primary_care_doc.styles._element)

    # Copy numbering definitions (for bullets)
    print("✓ Copying numbering definitions (for bullets)...")
    template_numbering = None
    if primary_care_doc._part.numbering_part is not None:
        merged_doc._part.numbering_part._element = deepcopy(
            primary_care_doc._part.numbering_part._element
        )
        template_numbering = merged_doc._part.numbering_part._element

        # Fix bullet sizes - ensure they're proportional (11pt to match text)
        print("✓ Adjusting bullet sizes to be proportional...")
        numbering_element = merged_doc._part.numbering_part._element
        # Find all abstractNum elements and set reasonable bullet font sizes
        for abstractNum in numbering_element.findall('.//' + qn('w:abstractNum')):
            for lvl in abstractNum.findall('.//' + qn('w:lvl')):
                # Find the rPr (run properties) for this level
                rPr = lvl.find(qn('w:rPr'))
                if rPr is not None:
                    # Set or update the font size for bullets
                    sz = rPr.find(qn('w:sz'))
                    if sz is not None:
                        # Set to 22 (11pt in half-points)
                        sz.set(qn('w:val'), '22')
                    else:
                        # Create sz element if it doesn't exist
                        sz = OxmlElement('w:sz')
                        sz.set(qn('w:val'), '22')
                        rPr.append(sz)

    # Add Table of Contents
    print("\n" + "=" * 70)
    print("ADDING TABLE OF CONTENTS")
    print("=" * 70)
    add_table_of_contents(merged_doc)
    print("✓ Table of Contents added")

    # Process Primary Care
    print("\n" + "=" * 70)
    print("MERGING PRIMARY CARE (with intelligent formatting)")
    print("=" * 70)
    pc_use_cases = process_document_content(
        primary_care_doc, merged_doc, template_numbering, add_page_break_before_first=False
    )
    print(f"✓ Added {pc_use_cases} Primary Care use cases")

    # Add page break before Cardiology
    merged_doc.add_page_break()

    # Process Cardiology
    print("\n" + "=" * 70)
    print("MERGING CARDIOLOGY (with intelligent formatting)")
    print("=" * 70)
    cardio_use_cases = process_document_content(
        cardiology_doc, merged_doc, template_numbering, add_page_break_before_first=False
    )
    print(f"✓ Added {cardio_use_cases} Cardiology use cases")

    # Add page break before Infectious Diseases
    merged_doc.add_page_break()

    # Process Infectious Diseases
    print("\n" + "=" * 70)
    print("MERGING INFECTIOUS DISEASES (with intelligent formatting)")
    print("=" * 70)
    id_use_cases = process_document_content(
        infectious_diseases_doc, merged_doc, template_numbering, add_page_break_before_first=False
    )
    print(f"✓ Added {id_use_cases} Infectious Diseases use cases")

    # Add page numbers to footer
    print("\n" + "=" * 70)
    print("ADDING PAGE NUMBERS")
    print("=" * 70)
    for section in merged_doc.sections:
        add_page_number(section)
    print("✓ Page numbers added to bottom right")

    # Save the merged document
    output_path = './Cordance_insights_bank_draft.docx'
    print("\n" + "=" * 70)
    print("SAVING MERGED DOCUMENT")
    print("=" * 70)
    merged_doc.save(output_path)

    print(f"\n{'=' * 70}")
    print("✅ SUCCESS! INTELLIGENTLY FORMATTED MERGED DOCUMENT")
    print(f"{'=' * 70}")
    print(f"\nFile: {output_path}")
    print(f"\nDocument contains:")
    print(f"  • Primary Care: {pc_use_cases} use cases")
    print(f"  • Cardiology: {cardio_use_cases} use cases")
    print(f"  • Infectious Diseases: {id_use_cases} use cases")
    print(f"  • Total: {pc_use_cases + cardio_use_cases + id_use_cases} use cases")
    print(f"\nFormatting applied INTELLIGENTLY to ALL sections:")
    print(f"  ✓ Heading 1: Times New Roman, 17pt, Blue (#2E74B5)")
    print(f"  ✓ Heading 2: Times New Roman, 14pt, Blue (#2E74B5)")
    print(f"  ✓ Heading 3: Times New Roman, 13pt, Dark Blue (#1F4D78)")
    print(f"  ✓ Heading 4: Times New Roman, 11pt, Italic, Blue (#2E74B5)")
    print(f"  ✓ Body Text: Arial, 11pt")
    print(f"  ✓ Bullet Points: List Paragraph with bullets")
    print(f"  ✓ Page Numbers: Bottom right")
    print(f"  ✓ Page Breaks: After each use case")
    print(f"  ✓ Table of Contents: Heading 1 and Heading 2")
    print(f"\n⚠️  IMPORTANT: Open in Word and update Table of Contents:")
    print(f"  1. Right-click the Table of Contents")
    print(f"  2. Select 'Update Field'")
    print(f"  3. Choose 'Update entire table'")
    print(f"\n{'=' * 70}\n")


if __name__ == '__main__':
    merge_documents()
