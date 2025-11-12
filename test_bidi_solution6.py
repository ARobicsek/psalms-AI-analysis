"""
Solution 6: Reverse Hebrew by GRAPHEME CLUSTERS + LEFT-TO-RIGHT OVERRIDE

Solution 5 was close but had issues with combining characters (nikud/vowel points).
This solution properly reverses Hebrew by keeping base letter + combining marks together.
"""

from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement, ns
from pathlib import Path
import re


def set_paragraph_ltr(paragraph):
    """Set paragraph direction to LTR at XML level."""
    pPr = paragraph._element.get_or_add_pPr()
    bidi_elem = pPr.find(ns.qn('w:bidi'))
    if bidi_elem is not None:
        pPr.remove(bidi_elem)
    bidi_elem = OxmlElement('w:bidi')
    bidi_elem.set(ns.qn('w:val'), '0')
    pPr.append(bidi_elem)


def split_into_grapheme_clusters(text):
    """
    Split Hebrew text into grapheme clusters (base letter + combining marks).

    A grapheme cluster consists of:
    - A base character (Hebrew letter or space)
    - Followed by zero or more combining characters (nikud, cantillation, shin/sin dot, etc.)

    Combining character ranges:
    - U+0591-U+05BD: Cantillation marks
    - U+05BF: Rafe
    - U+05C1-U+05C2: Shin/Sin dots
    - U+05C4-U+05C7: Other marks
    - U+05B0-U+05BD: Vowel points (nikud)
    """
    # Pattern: base character followed by any combining marks
    # Base: Hebrew letter (U+05D0-U+05EA) or space
    # Combining: U+0591-U+05BD, U+05BF, U+05C1-U+05C2, U+05C4-U+05C7
    cluster_pattern = r'[\u05D0-\u05EA\s][\u0591-\u05BD\u05BF\u05C1-\u05C2\u05C4-\u05C7]*'

    clusters = re.findall(cluster_pattern, text)
    return clusters


def reverse_hebrew_by_clusters(hebrew_text):
    """
    Reverse Hebrew text by grapheme clusters (keeping letter+nikud together).

    Example:
    - Input: "שִׁלוֹם" (shalom) = [שִׁ, ל, וֹ, ם]
    - Output: "םוֹלשִׁ" (reversed clusters)
    """
    clusters = split_into_grapheme_clusters(hebrew_text)

    # Reverse the order of clusters (but keep each cluster intact)
    reversed_clusters = clusters[::-1]

    return ''.join(reversed_clusters)


def solution6_cluster_reverse_plus_lro(text):
    """
    Solution 6: Reverse Hebrew by grapheme clusters + LEFT-TO-RIGHT OVERRIDE

    Strategy:
    1. Extract Hebrew text from inside parentheses
    2. Split Hebrew into grapheme clusters (letter+nikud units)
    3. Reverse the ORDER of clusters (but keep each cluster intact)
    4. Wrap with LRO...PDF to force left-to-right display
    5. Result: Properly formed Hebrew displayed with correct RTL appearance
    """
    LRO = '\u202D'  # LEFT-TO-RIGHT OVERRIDE
    PDF = '\u202C'  # POP DIRECTIONAL FORMATTING

    def reverse_hebrew_in_match(match):
        """Extract Hebrew, reverse by clusters, wrap with LRO."""
        hebrew_text = match.group(1)

        # Reverse by grapheme clusters (keeps combining marks attached)
        reversed_hebrew = reverse_hebrew_by_clusters(hebrew_text)

        # Wrap the entire (reversed_hebrew) with LRO...PDF
        return f'{LRO}({reversed_hebrew}){PDF}'

    # Pattern to match (Hebrew text with nikud and spaces)
    hebrew_paren_pattern = r'\(([\u0590-\u05FF\u05B0-\u05BD\u05BF\u05C1-\u05C2\u05C4-\u05C7\s]+)\)'

    return re.sub(hebrew_paren_pattern, reverse_hebrew_in_match, text)


def create_test_document():
    """Create test document with Solution 6."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Aptos'
    style.font.size = Pt(12)

    # Add title
    doc.add_heading('Solution 6: Grapheme Cluster Reversal + LRO', level=1)

    # Add explanation
    p_explain = doc.add_paragraph()
    p_explain.add_run('Strategy: ').bold = True
    p_explain.add_run('Split Hebrew into grapheme clusters (letter+nikud), reverse cluster order, then apply LRO. ')
    p_explain.add_run('Result: ').bold = True
    p_explain.add_run('Combining marks stay attached to base letters, no dotted circles!')

    doc.add_paragraph()

    # Test cases
    test_cases = [
        {
            'label': 'Test 1: Simple parenthesized Hebrew',
            'text': 'And my נֶפֶשׁ is deeply terrified (וְנַפְשִׁי נִבְהֲלָה מְאֹד). נֶפֶשׁ is not a detachable soul.'
        },
        {
            'label': 'Test 2: Multiple instances',
            'text': 'The word אַשְׁרֵי (ashrei) means blessed (אַשְׁרֵי הָאִישׁ), and תְּהִלָּה (tehillah) means praise (תְּהִלָּה לְדָוִד).'
        },
        {
            'label': 'Test 3: Hebrew with spaces',
            'text': 'The phrase "Happy is the man" (אַשְׁרֵי הָאִישׁ אֲשֶׁר) is the opening of Psalm 1.'
        },
        {
            'label': 'Test 4: Mixed nikud',
            'text': 'The root ברך appears in בָּרוּךְ (blessed) and בְּרָכָה (blessing).'
        },
        {
            'label': 'Test 5: Long Hebrew text',
            'text': 'The entire first verse (אַשְׁרֵי הָאִישׁ אֲשֶׁר לֹא הָלַךְ בַּעֲצַת רְשָׁעִים) sets the tone for the psalm.'
        },
        {
            'label': 'Test 6: Complex nikud (stress test)',
            'text': 'The word שִׁמְעוּ (hear) has multiple combining marks on the shin.'
        }
    ]

    for test_case in test_cases:
        # Add test case label
        p_label = doc.add_paragraph()
        set_paragraph_ltr(p_label)
        p_label.add_run(test_case['label']).bold = True

        # Add test case with Solution 6 applied
        p_test = doc.add_paragraph()
        set_paragraph_ltr(p_test)

        transformed_text = solution6_cluster_reverse_plus_lro(test_case['text'])
        p_test.add_run(transformed_text)

        # Add spacing
        doc.add_paragraph()

    # Add technical explanation
    doc.add_heading('Technical Details', level=2)

    p_tech1 = doc.add_paragraph()
    p_tech1.add_run('Grapheme Cluster: ').bold = True
    p_tech1.add_run('A base letter plus all its combining marks (nikud, shin/sin dot, dagesh, etc.)')

    p_tech2 = doc.add_paragraph()
    p_tech2.add_run('Example: ').bold = True
    p_tech2.add_run('The cluster "שִׁ" = [ש + ִ + ׁ] stays together when reversed')

    p_tech3 = doc.add_paragraph()
    p_tech3.add_run('Why this works: ').bold = True
    p_tech3.add_run('Combining marks remain attached to their base letter, preventing dotted circle placeholders')

    # Save
    output_path = Path('output/bidi_tests/solution6_cluster_reverse_lro.docx')
    doc.save(output_path)
    print(f"OK: Generated {output_path.name}")
    print(f"\nLocation: {output_path.absolute()}")
    print("\nThis solution properly handles Hebrew combining characters!")
    print("Please check if the dotted circles are gone and nikud displays correctly.")


if __name__ == '__main__':
    create_test_document()
