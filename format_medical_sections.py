#!/usr/bin/env python3
"""
Format Cardiology and Infectious Diseases sections to match Primary Care formatting.
This script ensures all sections follow the exact formatting specifications from FORMATTING_GUIDE.md
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy
import sys

def rgb_to_hex(rgb_color):
    """Convert RGBColor to hex string for comparison."""
    if rgb_color is None:
        return None
    return f"#{rgb_color.rgb:06X}" if hasattr(rgb_color, 'rgb') else None

def get_paragraph_style_info(para):
    """Extract detailed style information from a paragraph."""
    info = {
        'text': para.text[:100],  # First 100 chars
        'style_name': para.style.name if para.style else 'None',
        'font_name': None,
        'font_size': None,
        'bold': None,
        'italic': None,
        'color': None,
    }

    # Check runs for direct formatting
    if para.runs:
        first_run = para.runs[0]
        if first_run.font.name:
            info['font_name'] = first_run.font.name
        if first_run.font.size:
            info['font_size'] = first_run.font.size.pt
        if first_run.font.bold is not None:
            info['bold'] = first_run.font.bold
        if first_run.font.italic is not None:
            info['italic'] = first_run.font.italic
        if first_run.font.color and first_run.font.color.rgb:
            info['color'] = rgb_to_hex(first_run.font.color)

    return info

def remove_direct_formatting(para):
    """Remove direct formatting from paragraph runs to rely on style inheritance."""
    for run in para.runs:
        # Clear direct formatting
        run.font.name = None
        run.font.size = None
        run.font.bold = None
        run.font.italic = None
        run.font.color.rgb = None

def apply_style_only(para, style_name):
    """
    Apply ONLY the style to a paragraph, removing any direct formatting.
    This follows the guide's directive to use style inheritance, not direct formatting.
    """
    # Apply the style
    para.style = style_name

    # Remove any direct formatting from all runs
    for run in para.runs:
        # Clear all direct formatting to rely on style inheritance
        run.font.name = None
        run.font.size = None
        run.font.bold = None
        run.font.italic = None
        run.font.color.rgb = None


def is_heading_level(para, level):
    """Check if paragraph is a specific heading level."""
    if not para.style:
        return False
    style_name = para.style.name
    return (f'Heading {level}' == style_name or
            f'Heading{level}' == style_name.replace(' ', ''))

def find_section_boundaries(doc):
    """Find paragraph indices for each specialty section."""
    sections = {
        'Primary Care': {'start': None, 'end': None},
        'Cardiology': {'start': None, 'end': None},
        'Infectious Diseases': {'start': None, 'end': None}
    }

    section_order = []

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        if is_heading_level(para, 1):
            if 'Primary Care' in text:
                sections['Primary Care']['start'] = i
                section_order.append('Primary Care')
            elif 'Cardiology' in text:
                sections['Cardiology']['start'] = i
                section_order.append('Cardiology')
            elif 'Infectious Diseases' in text:
                sections['Infectious Diseases']['start'] = i
                section_order.append('Infectious Diseases')

    # Set end boundaries
    for i, section_name in enumerate(section_order):
        if i < len(section_order) - 1:
            next_section = section_order[i + 1]
            sections[section_name]['end'] = sections[next_section]['start']
        else:
            sections[section_name]['end'] = len(doc.paragraphs)

    return sections

def format_section(doc, start_idx, end_idx):
    """
    Format all paragraphs in a section according to the formatting guide.
    Uses STYLE INHERITANCE ONLY - no direct formatting.
    """
    for i in range(start_idx, end_idx):
        if i >= len(doc.paragraphs):
            break

        para = doc.paragraphs[i]
        text = para.text.strip()

        if not text:
            continue

        # Determine paragraph type and apply appropriate style (NO direct formatting)
        if is_heading_level(para, 1):
            # Heading 1: Specialty name
            apply_style_only(para, 'Heading 1')

        elif is_heading_level(para, 2):
            # Heading 2: USE CASE titles
            apply_style_only(para, 'Heading 2')

        elif is_heading_level(para, 3):
            # Heading 3: Major sections (Clinical Scenario, etc.)
            apply_style_only(para, 'Heading 3')

        elif is_heading_level(para, 4):
            # Heading 4: Subsections (Before Encounter, etc.)
            apply_style_only(para, 'Heading 4')

        elif para.style and 'List' in para.style.name:
            # List/bullet paragraphs
            apply_style_only(para, 'List Paragraph')

        else:
            # Regular body text
            apply_style_only(para, 'Normal')

def main():
    doc_path = '/home/user/psalms-AI-analysis/docs/Cordance/Cordance_insights_bank_draft.docx'

    print("Loading document...")
    doc = Document(doc_path)

    print("\nFinding section boundaries...")
    sections = find_section_boundaries(doc)

    for section_name, bounds in sections.items():
        if bounds['start'] is not None:
            print(f"  {section_name}: paragraphs {bounds['start']} to {bounds['end']}")
        else:
            print(f"  {section_name}: NOT FOUND")

    # Verify Primary Care section exists
    if sections['Primary Care']['start'] is None:
        print("\nERROR: Primary Care section not found!")
        return 1

    print("\n" + "="*80)
    print("ANALYZING PRIMARY CARE FORMATTING (Reference Section)")
    print("="*80)

    # Analyze a sample of Primary Care formatting
    pc_start = sections['Primary Care']['start']
    pc_end = min(pc_start + 50, sections['Primary Care']['end'])

    sample_paragraphs = []
    for i in range(pc_start, pc_end):
        para = doc.paragraphs[i]
        if para.text.strip():
            info = get_paragraph_style_info(para)
            if info['style_name'].startswith('Heading'):
                sample_paragraphs.append(info)
                if len(sample_paragraphs) >= 10:
                    break

    for info in sample_paragraphs:
        print(f"\nStyle: {info['style_name']}")
        print(f"  Text: {info['text'][:50]}...")
        print(f"  Font: {info['font_name']}, Size: {info['font_size']}pt")
        print(f"  Bold: {info['bold']}, Italic: {info['italic']}, Color: {info['color']}")

    # Format Cardiology section
    if sections['Cardiology']['start'] is not None:
        print("\n" + "="*80)
        print("FORMATTING CARDIOLOGY SECTION")
        print("="*80)
        format_section(
            doc,
            sections['Cardiology']['start'],
            sections['Cardiology']['end']
        )
        print("✓ Cardiology section formatted")
    else:
        print("\nWARNING: Cardiology section not found, skipping")

    # Format Infectious Diseases section
    if sections['Infectious Diseases']['start'] is not None:
        print("\n" + "="*80)
        print("FORMATTING INFECTIOUS DISEASES SECTION")
        print("="*80)
        format_section(
            doc,
            sections['Infectious Diseases']['start'],
            sections['Infectious Diseases']['end']
        )
        print("✓ Infectious Diseases section formatted")
    else:
        print("\nWARNING: Infectious Diseases section not found, skipping")

    # Save the formatted document
    output_path = doc_path.replace('.docx', '_formatted.docx')
    print(f"\nSaving formatted document to: {output_path}")
    doc.save(output_path)

    print("\n" + "="*80)
    print("FORMATTING COMPLETE!")
    print("="*80)
    print(f"\nFormatted document saved as: {output_path}")
    print("\nNext steps:")
    print("1. Open the document in Microsoft Word")
    print("2. Review the Navigation Pane to verify heading hierarchy")
    print("3. Check that all sections match Primary Care formatting")
    print("4. Update the Table of Contents if present (right-click → Update Field)")

    return 0

if __name__ == '__main__':
    sys.exit(main())
